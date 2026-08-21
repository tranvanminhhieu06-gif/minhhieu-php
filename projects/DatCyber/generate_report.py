# -*- coding: utf-8 -*-
"""
Script to generate high-quality academic and technical document BaoCao.docx
for DatCyber E-commerce Project adhering strictly to mucluc.txt.
"""

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_styled(doc, text, level):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(17)
        run.font.color.rgb = RGBColor(2, 100, 160)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 3:
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_body_p(doc, text, bold_prefix=None, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
        
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.italic = italic
    r.font.color.rgb = RGBColor(30, 41, 59)
    return p

def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0284C7")
        set_cell_margins(hdr_cells[i], top=150, bottom=150, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = str(cell_value)
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(30, 41, 59)

    if col_widths:
        for row in table.rows:
            for col_idx, width in enumerate(col_widths):
                row.cells[col_idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table

def generate_doc():
    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.0)

    # ================= COVER PAGE =================
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_inst.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG\n-------------------------***-------------------------\n\n\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(51, 65, 85)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_title.add_run("BÁO CÁO ĐỒ ÁN CHUYÊN NGÀNH\n")
    r_sub.font.name = 'Times New Roman'
    r_sub.font.size = Pt(15)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(2, 132, 199)

    r_main = p_title.add_run("NGHIÊN CỨU, PHÂN TÍCH, THIẾT KẾ VÀ XÂY DỰNG WEBSITE BÁN ĐỒ GIA DỤNG THÔNG MINH DATCYBER SỬ DỤNG PHP & MYSQL\n\n\n")
    r_main.font.name = 'Times New Roman'
    r_main.font.size = Pt(20)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(15, 23, 42)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.left_indent = Inches(1.5)
    p_meta.paragraph_format.line_spacing = 1.4
    
    r_m = p_meta.add_run("Giảng viên hướng dẫn: \tTS. Nguyễn Văn Hướng\nSinh viên thực hiện: \tTrần Văn Minh Hiếu\nMã số sinh viên: \t\t20268899\nChuyên ngành: \t\tCông nghệ Phần mềm & Lập trình Web\nKhóa học: \t\t\t2022 - 2026\n\n\n\n")
    r_m.font.name = 'Times New Roman'
    r_m.font.size = Pt(12.5)
    r_m.font.color.rgb = RGBColor(30, 41, 59)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_foot.add_run("Hà Nội, Năm 2026")
    r_f.font.name = 'Times New Roman'
    r_f.font.size = Pt(12)
    r_f.font.bold = True
    r_f.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # ================= MỤC LỤC =================
    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_toc = p_toc_title.add_run("MỤC LỤC BÁO CÁO")
    r_toc.font.name = 'Times New Roman'
    r_toc.font.size = Pt(18)
    r_toc.font.bold = True
    r_toc.font.color.rgb = RGBColor(2, 132, 199)

    toc_items = [
        ("LỜI NÓI ĐẦU", "Trang 3"),
        ("Chương 1. Tổng quan lập trình web", "Trang 4"),
        ("   1.1 Ngôn ngữ lập trình PHP", "Trang 4"),
        ("   1.2 Hệ quản trị cơ sở dữ liệu MySQL", "Trang 7"),
        ("   1.3 Cài đặt máy chủ", "Trang 9"),
        ("Chương 2. Phân tích và thiết kế website", "Trang 11"),
        ("   2.1 Chức năng (usecase)", "Trang 11"),
        ("   2.2 Cơ sở dữ liệu", "Trang 15"),
        ("Chương 3. Chương trình thử nghiệm", "Trang 18"),
        ("   3.1 Giao diện", "Trang 18"),
        ("   3.2 Kết luận", "Trang 22"),
        ("TÀI LIỆU THAM KHẢO", "Trang 24")
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (title, page) in enumerate(toc_items):
        cell_t = toc_table.rows[idx].cells[0]
        cell_p = toc_table.rows[idx].cells[1]
        cell_t.text = title
        cell_p.text = page
        cell_p.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for c in (cell_t, cell_p):
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11.5)
                if "Chương" in title or "LỜI" in title or "TÀI LIỆU" in title:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(2, 100, 160)
                else:
                    r.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_page_break()

    # ================= LỜI NÓI ĐẦU =================
    add_header_styled(doc, "LỜI NÓI ĐẦU", level=1)
    add_body_p(doc, "Trong kỷ nguyên bùng nổ của cuộc Cách mạng Công nghiệp 4.0 và chuyển đổi số toàn diện, thương mại điện tử (E-commerce) đã trở thành một phần tất yếu trong đời sống kinh tế xã hội. Người tiêu dùng ngày nay không chỉ tìm kiếm các sản phẩm thiết yếu với mức giá cạnh tranh, mà còn đặc biệt quan tâm đến trải nghiệm mua sắm trực tuyến: tốc độ truy cập mượt mà, giao diện trực quan, hiệu ứng tương tác hiện đại, hình ảnh sản phẩm chân thực cùng sự minh bạch về nguồn gốc và chính sách bảo hành.")
    add_body_p(doc, "Ngành hàng đồ gia dụng và thiết bị thông minh phục vụ đời sống gia đình (nồi chiên không dầu, robot hút bụi, máy ép chậm, máy lọc không khí, ấm siêu tốc...) đang chứng kiến tốc độ tăng trưởng vượt bậc tại thị trường Việt Nam. Nhận thức được nhu cầu cấp thiết đó, đồ án chuyên ngành này tập trung nghiên cứu, thiết kế và phát triển toàn diện hệ thống website thương mại điện tử chuyên biệt mang tên 'DatCyber' - Thiết Bị Gia Dụng Thông Minh Cao Cấp.")
    add_body_p(doc, "Đề tài được triển khai bằng việc ứng dụng ngôn ngữ lập trình phía máy chủ PHP phiên bản 8.x kết hợp cùng hệ quản trị cơ sở dữ liệu quan hệ MySQL. Đồng thời, website tích hợp bộ giải pháp hình ảnh sản phẩm chất lượng cao sinh bởi trí tuệ nhân tạo (AI Image Generation), hệ thống hiệu ứng chuyển động mượt mà (CSS Transitions & Keyframe Animations) cùng kiến trúc quản trị Admin mạnh mẽ, đáp ứng đầy đủ các tiêu chuẩn kỹ thuật của một ứng dụng web thương mại điện tử chuẩn mực.")

    # ================= CHƯƠNG 1 =================
    add_header_styled(doc, "Chương 1. Tổng quan lập trình web", level=1)
    
    # 1.1 PHP
    add_header_styled(doc, "1.1 Ngôn ngữ lập trình PHP", level=2)
    add_body_p(doc, "PHP (viết tắt đệ quy của Hypertext Preprocessor) là một trong những ngôn ngữ lập trình kịch bản phía máy chủ (Server-side Scripting Language) phổ biến và thành công nhất trong lịch sử phát triển của mạng Internet toàn cầu. Được tạo ra ban đầu bởi Rasmus Lerdorf vào năm 1994, trải qua hơn ba thập kỷ liên tục phát triển và hoàn thiện, PHP hiện đang là nền tảng vận hành cho hơn 75% các trang web có sử dụng ngôn ngữ lập trình máy chủ trên thế giới (theo thống kê của W3Techs).")
    
    add_body_p(doc, "1. Cơ chế hoạt động và kiến trúc Server-side của PHP:", bold_prefix="a. ")
    add_body_p(doc, "Khi một người dùng (Client) gửi yêu cầu HTTP Request thông qua trình duyệt web tới máy chủ chứa mã nguồn PHP, Web Server (như Apache hoặc Nginx) sẽ chuyển tiếp tệp `.php` tới bộ xử lý PHP Zend Engine. Tại đây, mã PHP được phân tích cú pháp (Parsing), biên dịch thành các mã chỉ lệnh (Opcode) và thực thi trong môi trường máy chủ. Kết quả đầu ra sau khi xử lý xong các tác vụ nghiệp vụ, truy vấn cơ sở dữ liệu MySQL và dựng giao diện sẽ là chuỗi văn bản thuần HTML/CSS/JavaScript hoặc JSON, được đóng gói vào gói phản hồi HTTP Response gửi ngược lại cho trình duyệt hiển thị. Điều này giúp mã nguồn và logic nghiệp vụ được bảo vệ an toàn tuyệt đối tại máy chủ, người dùng không thể can thiệp vào logic xử lý dữ liệu nội bộ.")

    add_body_p(doc, "2. Những đột phá mạnh mẽ trên phiên bản PHP 8.x:", bold_prefix="b. ")
    add_body_p(doc, "Dự án DatCyber được phát triển trên nền tảng PHP 8.2 với những cải tiến công nghệ vượt bậc so với các phiên bản PHP 5.x hay 7.x cũ:", italic=False)
    
    php_feats = [
        ["Tính năng PHP 8.x", "Mô tả chi tiết và ý nghĩa kỹ thuật", "Ứng dụng trong dự án DatCyber"],
        ["JIT (Just-In-Time) Compiler", "Biên dịch mã opcode trực tiếp sang mã máy x86/ARM trong thời gian chạy, tăng hiệu năng tính toán lên 2-3 lần.", "Tối ưu tốc độ dựng trang và xử lý dữ liệu giỏ hàng lớn."],
        ["Named Arguments", "Cho phép truyền tham số vào hàm dựa trên tên định danh thay vì thứ tự vị trí bắt buộc.", "Tăng tính linh hoạt và dễ bảo trì trong các hàm tiện ích (`functions.php`)."],
        ["Constructor Property Promotion", "Rút ngắn cú pháp khai báo và gán thuộc tính trong hàm khởi tạo của Class.", "Viết code OOP tinh gọn, chuẩn mực và dễ đọc."],
        ["Match Expression", "Cấu trúc lựa chọn giá trị trả về an toàn kiểu (Type-safe) và ngắn gọn hơn switch-case truyền thống.", "Phân loại trạng thái đơn hàng và tính toán tỷ lệ khuyến mãi."],
        ["PDO (PHP Data Objects)", "Tầng trừu tượng hóa kết nối cơ sở dữ liệu, hỗ trợ Prepared Statements chống SQL Injection.", "Đảm bảo an toàn 100% cho mọi truy vấn dữ liệu vào MySQL."]
    ]
    create_styled_table(doc, php_feats[0], php_feats[1:], [Inches(1.8), Inches(3.2), Inches(1.8)])

    add_body_p(doc, "3. Ưu điểm nổi bật của PHP trong phát triển ứng dụng thương mại điện tử:", bold_prefix="c. ")
    add_body_p(doc, "- Tính sẵn sàng cao và dễ dàng triển khai: PHP có thể chạy liền mạch trên hầu hết các nền tảng hệ điều hành (Windows, Linux, macOS) và tương thích hoàn hảo với các web server phổ biến nhất.")
    add_body_p(doc, "- Hệ sinh thái phong phú và cộng đồng lớn: Cung cấp hàng triệu thư viện hỗ trợ xử lý chuỗi, mã hóa mật khẩu (`password_hash`), xử lý ảnh, tạo hóa đơn và giao tiếp API RESTful.")
    add_body_p(doc, "- Chi phí vận hành thấp: Mã nguồn mở miễn phí, tài nguyên tiêu thụ tối ưu, phù hợp cho cả các doanh nghiệp khởi nghiệp và hệ thống quy mô lớn.")

    # 1.2 MySQL
    add_header_styled(doc, "1.2 Hệ quản trị cơ sở dữ liệu MySQL", level=2)
    add_body_p(doc, "MySQL là hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở (RDBMS - Relational Database Management System) phổ biến nhất hiện nay. MySQL lưu trữ dữ liệu dưới dạng các bảng có cấu trúc (Tables) gồm các hàng (Rows/Records) và các cột (Columns/Attributes), cho phép thiết lập các mối quan hệ chặt chẽ (Relationships) giữa các thực thể dữ liệu thông qua cơ chế khóa chính (Primary Key) và khóa ngoại (Foreign Key).")

    add_body_p(doc, "1. Các đặc tính cốt lõi của MySQL được ứng dụng trong DatCyber:", bold_prefix="a. ")
    add_body_p(doc, "- Động cơ lưu trữ InnoDB: Mặc định trên MySQL 8.x, hỗ trợ đầy đủ tiêu chuẩn ACID (Atomicity, Consistency, Isolation, Durability) cho các giao dịch (Transactions). Trong luồng thanh toán giỏ hàng của website DatCyber, thao tác trừ tồn kho sản phẩm và ghi nhận đơn hàng được bọc trong khối giao dịch `beginTransaction()` - `commit()` - `rollBack()`, ngăn chặn triệt để tình trạng xung đột hoặc sai lệch số lượng hàng tồn.")
    add_body_p(doc, "- Khóa ngoại và tính toàn vẹn tham chiếu: Các bảng `order_items`, `products`, `reviews` đều được thiết lập khóa ngoại liên kết với `ON DELETE CASCADE` hoặc `ON DELETE SET NULL`, bảo đảm khi danh mục hoặc sản phẩm thay đổi thì tính toàn vẹn dữ liệu luôn được duy trì.")
    add_body_p(doc, "- Bộ giải mã ký tự UTF8MB4: Cơ sở dữ liệu `datcyber_appliances_db` được cấu hình sử dụng `utf8mb4_unicode_ci`, hỗ trợ đầy đủ và chuẩn xác các ký tự tiếng Việt có dấu, ký hiệu tiền tệ (₫) cũng như các ký tự biểu cảm emoji hiện đại.")

    mysql_props = [
        ["Tiêu chuẩn ACID", "Nội dung nguyên lý", "Ý nghĩa đối với hệ thống DatCyber"],
        ["Atomicity (Nguyên tử)", "Mọi tác vụ trong một giao dịch hoặc cùng thành công hoặc cùng bị hủy bỏ.", "Khi tạo đơn hàng, nếu lưu chi tiết sản phẩm thất bại thì toàn bộ đơn hàng sẽ tự động rollback."],
        ["Consistency (Nhất quán)", "Dữ liệu chuyển từ trạng thái hợp lệ này sang trạng thái hợp lệ khác đúng ràng buộc.", "Tồn kho không bao giờ bị âm, đơn giá luôn đúng chuẩn định dạng số."],
        ["Isolation (Cô lập)", "Các giao dịch chạy đồng thời không can thiệp hay làm sai lệch dữ liệu của nhau.", "Nhiều khách hàng cùng đặt mua một sản phẩm cùng lúc không bị tranh chấp dữ liệu."],
        ["Durability (Bền vững)", "Dữ liệu sau khi đã commit sẽ được lưu trữ an toàn ngay cả khi máy chủ gặp sự cố nguồn điện.", "Đơn hàng đã xác nhận thành công sẽ được ghi vĩnh viễn vào ổ đĩa."]
    ]
    create_styled_table(doc, mysql_props[0], mysql_props[1:], [Inches(1.5), Inches(3.0), Inches(2.3)])

    # 1.3 Cài đặt máy chủ
    add_header_styled(doc, "1.3 Cài đặt máy chủ", level=2)
    add_body_p(doc, "Để triển khai và vận hành hệ thống website PHP & MySQL trong môi trường thực nghiệm, mô hình máy chủ web cục bộ (Local Web Server) được lựa chọn thông qua gói phần mềm tích hợp XAMPP (Apache + MariaDB/MySQL + PHP + Perl).")

    add_body_p(doc, "1. Kiến trúc mô hình máy chủ Client - Server:", bold_prefix="a. ")
    add_body_p(doc, "Mô hình Client - Server phân chia hệ thống thành hai phần độc lập nhưng tương tác chặt chẽ qua giao thức truyền thông tiêu chuẩn:")
    add_body_p(doc, "- Phía Client (Trình duyệt web của người dùng): Đảm nhận nhiệm vụ gửi yêu cầu tương tác và hiển thị giao diện đồ họa. Sử dụng HTML5, CSS3 và JavaScript thuần để render hình ảnh sản phẩm, thực hiện các hiệu ứng hoạt họa và bắt sự kiện người dùng.")
    add_body_p(doc, "- Phía Server (Máy chủ Apache & PHP): Lắng nghe tại cổng mạng (Port 80/443 hoặc cổng tùy chỉnh), nhận yêu cầu, gọi chương trình PHP biên dịch logic, truy vấn cơ sở dữ liệu MySQL qua cổng 3306 và phản hồi dữ liệu về cho Client.")

    add_body_p(doc, "2. Quy trình cài đặt và thiết lập máy chủ XAMPP cho dự án DatCyber:", bold_prefix="b. ")
    add_body_p(doc, "Bước 1: Tải và cài đặt phiên bản XAMPP mới nhất hỗ trợ PHP 8.2 và MariaDB 10.4 trên hệ điều hành Windows tại thư mục `C:\\xampp`.")
    add_body_p(doc, "Bước 2: Kích hoạt các dịch vụ Apache HTTP Server và MySQL Server từ bảng điều khiển XAMPP Control Panel hoặc thông qua script quản lý tự động.")
    add_body_p(doc, "Bước 3: Cấu hình tệp `php.ini` để kích hoạt các extension quan trọng như `extension=pdo_mysql`, `extension=mbstring`, `extension=fileinfo` và thiết lập giới hạn bộ nhớ `memory_limit = 512M`.")
    add_body_p(doc, "Bước 4: Thiết lập cơ sở dữ liệu `datcyber_appliances_db` và nạp toàn bộ cấu trúc bảng cùng dữ liệu mẫu (Seed data) thông qua tập lệnh SQL `database.sql`.")
    add_body_p(doc, "Bước 5: Khởi chạy máy chủ ảo hoặc PHP Built-in Web Server tại thư mục dự án và truy cập kiểm thử qua địa chỉ `http://localhost/` hoặc `http://127.0.0.1:8000`.")

    # ================= CHƯƠNG 2 =================
    add_header_styled(doc, "Chương 2. Phân tích và thiết kế website", level=1)
    
    # 2.1 Chức năng (usecase)
    add_header_styled(doc, "2.1 Chức năng ( usecase )", level=2)
    add_body_p(doc, "Phân tích chức năng hệ thống là giai đoạn then chốt nhằm xác định rõ ràng các nhóm tác nhân (Actors), ranh giới hệ thống cũng như các trường hợp sử dụng (Use Cases) mà website DatCyber cần đáp ứng để phục vụ người dùng cuối và đội ngũ quản trị.")

    add_body_p(doc, "1. Xác định các tác nhân (Actors) trong hệ thống:", bold_prefix="a. ")
    add_body_p(doc, "- Khách vãng lai (Guest): Người dùng truy cập website nhưng chưa đăng nhập tài khoản. Có quyền duyệt danh mục, tìm kiếm sản phẩm, xem chi tiết thông số kỹ thuật, xem đánh giá, thêm hàng vào giỏ và đặt hàng nhanh.")
    add_body_p(doc, "- Khách hàng thành viên (Customer): Người dùng có tài khoản đã đăng ký, có thể đăng nhập để quản lý lịch sử đơn hàng cá nhân, gửi nhận xét và nhận các mã ưu đãi độc quyền.")
    add_body_p(doc, "- Quản trị viên (Admin): Người quản lý toàn quyền hệ thống. Truy cập phân hệ Admin để theo dõi biểu đồ doanh thu, quản lý danh mục, thêm/sửa/xóa sản phẩm, cập nhật tồn kho và thay đổi trạng thái xử lý đơn hàng.")

    add_body_p(doc, "2. Bảng phân rã chi tiết các Use Case chính trong hệ thống DatCyber:", bold_prefix="b. ")
    
    uc_data = [
        ["Mã UC", "Tên Use Case", "Tác nhân chính", "Mô tả tóm tắt luồng xử lý"],
        ["UC01", "Xem & Tìm kiếm sản phẩm", "Khách hàng, Khách vãng lai", "Tìm kiếm theo từ khóa tên sản phẩm, lọc theo danh mục, lọc theo khoảng giá, sắp xếp theo giá tăng/giảm."],
        ["UC02", "Xem chi tiết sản phẩm", "Khách hàng, Khách vãng lai", "Xem ảnh AI phóng to, thông số kỹ thuật, giá khuyến mãi, tình trạng tồn kho, chính sách bảo hành 24T."],
        ["UC03", "Quản lý giỏ hàng thông minh", "Khách hàng, Khách vãng lai", "Thêm sản phẩm qua AJAX không cần load lại trang, mở giỏ hàng trượt (Drawer), tăng giảm số lượng, xóa món."],
        ["UC04", "Áp dụng mã giảm giá (Voucher)", "Khách hàng, Khách vãng lai", "Nhập mã voucher (DATCYBER10, FREESHIP), hệ thống tự động kiểm tra điều kiện và trừ tiền trực tiếp."],
        ["UC05", "Đặt hàng & Thanh toán", "Khách hàng, Khách vãng lai", "Điền thông tin nhận hàng, chọn hình thức thanh toán (COD, VietQR Banking, Ví điện tử), xác nhận đơn."],
        ["UC06", "Theo dõi tiến trình đơn hàng", "Khách hàng, Khách vãng lai", "Tra cứu đơn hàng qua mã code, hiển thị thanh tiến trình 4 giai đoạn từ đóng gói đến giao thành công."],
        ["UC07", "Đánh giá & Nhận xét sản phẩm", "Khách hàng", "Gửi số sao (1-5 sao) và cảm nhận đánh giá; hệ thống tự động tính lại điểm trung bình cho sản phẩm."],
        ["UC08", "Quản trị sản phẩm (CRUD)", "Quản trị viên (Admin)", "Thêm sản phẩm mới, cập nhật giá bán, cập nhật ảnh AI, chỉnh sửa tồn kho, thiết lập Flash Sale / Best Seller."],
        ["UC09", "Quản trị đơn hàng", "Quản trị viên (Admin)", "Xem danh sách đơn đặt hàng, chi tiết người nhận và chuyển đổi trạng thái: Chờ xử lý -> Đang giao -> Hoàn thành."],
        ["UC10", "Thống kê doanh thu & Dashboard", "Quản trị viên (Admin)", "Xem tổng doanh thu tích lũy, tổng số lượng đơn, số lượng danh mục và các sản phẩm bán chạy nhất."]
    ]
    create_styled_table(doc, uc_data[0], uc_data[1:], [Inches(0.9), Inches(1.8), Inches(1.6), Inches(2.5)])

    add_body_p(doc, "3. Đặc tả chi tiết luồng nghiệp vụ Đặt hàng & Thanh toán (UC05):", bold_prefix="c. ")
    add_body_p(doc, "- Điều kiện tiên quyết: Giỏ hàng phải có ít nhất 01 sản phẩm.")
    add_body_p(doc, "- Luồng sự kiện chính:")
    add_body_p(doc, "  + Bước 1: Người dùng nhấn 'Tiến hành đặt hàng' từ giỏ hàng hoặc nhấn 'Mua ngay' từ trang chi tiết sản phẩm.")
    add_body_p(doc, "  + Bước 2: Hệ thống chuyển hướng tới giao diện `checkout.php` và hiển thị bảng tóm tắt đơn giá, phí vận chuyển và số tiền giảm giá.")
    add_body_p(doc, "  + Bước 3: Người dùng nhập thông tin người nhận (Họ tên, SĐT, Địa chỉ chi tiết, Ghi chú) và lựa chọn phương thức thanh toán.")
    add_body_p(doc, "  + Bước 4: Người dùng nhấn nút 'Xác nhận đặt hàng'. Hệ thống mở Transaction kiểm tra số lượng tồn kho trong database.")
    add_body_p(doc, "  + Bước 5: Hệ thống tạo bản ghi mới trong bảng `orders` với mã định danh duy nhất (VD: `DC-20260820-A1B2C3`), đồng thời sao chép toàn bộ các sản phẩm từ Session Cart vào bảng `order_items` và trừ số lượng tương ứng trong bảng `products`.")
    add_body_p(doc, "  + Bước 6: Transaction được Commit thành công, hệ thống xóa rỗng giỏ hàng và chuyển hướng người dùng tới trang `order-success.php` kèm hóa đơn chi tiết.")

    # 2.2 Cơ sở dữ liệu
    add_header_styled(doc, "2.2 Cơ sở dữ liệu", level=2)
    add_body_p(doc, "Cơ sở dữ liệu của hệ thống DatCyber được thiết kế theo mô hình quan hệ chuẩn hóa 3NF (Third Normal Form), loại bỏ hoàn toàn các dị thường dư thừa dữ liệu (Redundancy) và đảm bảo hiệu năng truy vấn tối ưu.")

    add_body_p(doc, "1. Danh mục các bảng trong cơ sở dữ liệu `datcyber_appliances_db`:", bold_prefix="a. ")
    
    tbl_summary = [
        ["Tên bảng (Table)", "Mục đích sử dụng & Ý nghĩa nghiệp vụ", "Số trường", "Khóa ngoại liên kết"],
        ["categories", "Lưu trữ các nhóm danh mục thiết bị gia dụng (Nhà bếp, Robot, Lọc khí...)", "5 trường", "Không"],
        ["products", "Lưu thông tin chi tiết từng sản phẩm gia dụng, giá, ảnh AI, tồn kho, thông số", "17 trường", "category_id -> categories(id)"],
        ["users", "Lưu trữ thông tin tài khoản người dùng và tài khoản quản trị Admin", "8 trường", "Không"],
        ["orders", "Lưu trữ thông tin tổng quan các đơn đặt hàng, người nhận, tổng tiền, trạng thái", "14 trường", "user_id -> users(id)"],
        ["order_items", "Lưu trữ chi tiết từng mặt hàng và số lượng trong một đơn đặt hàng", "8 trường", "order_id -> orders(id), product_id -> products(id)"],
        ["reviews", "Lưu trữ các bài đánh giá, điểm số sao và bình luận của khách hàng", "6 trường", "product_id -> products(id)"],
        ["coupons", "Lưu trữ các mã khuyến mãi giảm giá, điều kiện tối thiểu và hạn dùng", "7 trường", "Không"]
    ]
    create_styled_table(doc, tbl_summary[0], tbl_summary[1:], [Inches(1.3), Inches(3.2), Inches(1.0), Inches(1.5)])

    add_body_p(doc, "2. Chi tiết cấu trúc dữ liệu của các bảng trọng tâm:", bold_prefix="b. ")
    
    add_body_p(doc, "- Cấu trúc bảng `products` (Sản phẩm gia dụng):", bold_prefix="* ")
    prod_fields = [
        ["Tên cột (Field)", "Kiểu dữ liệu", "Ràng buộc (Constraint)", "Mô tả ý nghĩa"],
        ["id", "INT", "PRIMARY KEY, AUTO_INCREMENT", "Mã định danh duy nhất của sản phẩm"],
        ["category_id", "INT", "NOT NULL, FOREIGN KEY", "Khóa ngoại tham chiếu đến bảng categories"],
        ["name", "VARCHAR(255)", "NOT NULL", "Tên đầy đủ của thiết bị gia dụng"],
        ["slug", "VARCHAR(255)", "NOT NULL, UNIQUE", "Đường dẫn URL thân thiện chuẩn SEO"],
        ["price", "DECIMAL(12,0)", "NOT NULL", "Giá bán thực tế hiện tại (VNĐ)"],
        ["old_price", "DECIMAL(12,0)", "DEFAULT NULL", "Giá niêm yết gốc trước khi giảm giá"],
        ["image", "VARCHAR(255)", "NOT NULL", "Tên tệp hình ảnh sản phẩm tạo bởi AI"],
        ["short_description", "TEXT", "DEFAULT NULL", "Mô tả tóm tắt tính năng nổi bật"],
        ["description", "LONGTEXT", "DEFAULT NULL", "Nội dung bài viết chi tiết sản phẩm"],
        ["specs", "TEXT", "DEFAULT NULL", "Bảng thông số kỹ thuật (công suất, dung tích...)"],
        ["stock", "INT", "NOT NULL, DEFAULT 50", "Số lượng hàng hóa còn trong kho"],
        ["rating", "DECIMAL(2,1)", "DEFAULT 5.0", "Điểm đánh giá trung bình (1.0 - 5.0 sao)"],
        ["review_count", "INT", "DEFAULT 0", "Tổng số lượt đánh giá của khách hàng"],
        ["is_flash_sale", "TINYINT(1)", "DEFAULT 0", "Đánh dấu sản phẩm thuộc đợt Flash Sale"],
        ["is_best_seller", "TINYINT(1)", "DEFAULT 0", "Đánh dấu sản phẩm bán chạy nhất"]
    ]
    create_styled_table(doc, prod_fields[0], prod_fields[1:], [Inches(1.5), Inches(1.3), Inches(1.8), Inches(2.2)])

    add_body_p(doc, "- Cấu trúc bảng `orders` (Đơn đặt hàng):", bold_prefix="* ")
    order_fields = [
        ["Tên cột (Field)", "Kiểu dữ liệu", "Ràng buộc (Constraint)", "Mô tả ý nghĩa"],
        ["id", "INT", "PRIMARY KEY, AUTO_INCREMENT", "Mã khóa chính đơn hàng"],
        ["order_code", "VARCHAR(50)", "NOT NULL, UNIQUE", "Mã đơn hàng hiển thị cho khách (VD: HM-20260820-01)"],
        ["customer_name", "VARCHAR(150)", "NOT NULL", "Họ và tên người nhận hàng"],
        ["customer_phone", "VARCHAR(30)", "NOT NULL", "Số điện thoại liên hệ giao hàng"],
        ["customer_address", "TEXT", "NOT NULL", "Địa chỉ nhận hàng chi tiết"],
        ["payment_method", "VARCHAR(50)", "DEFAULT 'cod'", "Hình thức: cod, banking, momo"],
        ["total_amount", "DECIMAL(12,0)", "NOT NULL", "Tổng tiền hàng trước giảm giá"],
        ["discount_amount", "DECIMAL(12,0)", "DEFAULT 0", "Số tiền được giảm trừ qua mã Voucher"],
        ["shipping_fee", "DECIMAL(12,0)", "DEFAULT 0", "Phí vận chuyển (0đ nếu trên 500k)"],
        ["final_amount", "DECIMAL(12,0)", "NOT NULL", "Số tiền cuối cùng khách phải thanh toán"],
        ["status", "ENUM", "pending, processing, shipping, completed, cancelled", "Trạng thái xử lý của đơn hàng"]
    ]
    create_styled_table(doc, order_fields[0], order_fields[1:], [Inches(1.5), Inches(1.3), Inches(1.8), Inches(2.2)])

    # ================= CHƯƠNG 3 =================
    add_header_styled(doc, "Chương 3. Chương trình thử nghiệm", level=1)
    
    # 3.1 Giao diện
    add_header_styled(doc, "3.1 Giao diện", level=2)
    add_body_p(doc, "Giao diện người dùng của website DatCyber được thiết kế và xây dựng tuân thủ nghiêm ngặt các nguyên lý thiết kế UI/UX hiện đại bậc nhất, mang đến trải nghiệm thị giác ấn tượng (Visual Wow Effect), cảm giác tương tác mượt mà và khả năng sử dụng trực quan, tiện lợi.")

    add_body_p(doc, "1. Ngôn ngữ thiết kế và Hệ thống nhận diện thị giác (Design System):", bold_prefix="a. ")
    add_body_p(doc, "- Bảng màu chủ đạo (Color Palette): Sử dụng gam màu xanh công nghệ thanh lịch (Ocean Blue #0284c7) kết hợp nền xanh đen hiện đại (Dark Slate #0f172a) và điểm xuyết các sắc cam năng động (Vibrant Orange #f97316) cho các nút hành động, thẻ giảm giá và đếm ngược Flash Sale.")
    add_body_p(doc, "- Nghệ thuật chữ (Typography): Sử dụng font chữ Google Fonts `Plus Jakarta Sans` với các trọng số từ 400 đến 800, mang lại vẻ đẹp sắc nét, hiện đại và tối ưu khả năng đọc trên mọi kích thước màn hình.")
    add_body_p(doc, "- Hiệu ứng kính mờ (Glassmorphism): Thanh điều hướng Header Sticky và các khối thẻ nổi được áp dụng kỹ thuật `backdrop-filter: blur(16px)` với độ trong suốt tinh tế, tạo chiều sâu không gian đa tầng.")

    add_body_p(doc, "2. Hình ảnh sản phẩm chất lượng cao sinh bởi Trí tuệ Nhân tạo (AI Image Generation):", bold_prefix="b. ")
    add_body_p(doc, "Toàn bộ hình ảnh sản phẩm trong hệ thống DatCyber đều được khởi tạo trực tiếp bằng công nghệ AI sinh ảnh chuyên nghiệp, đảm bảo độ phân giải cao 1:1, góc chụp studio thương mại chuẩn mực, ánh sáng tự nhiên ấm áp và độ sắc nét chân thực. Danh mục hình ảnh bao gồm:")
    add_body_p(doc, "+ Nồi chiên không dầu điện tử `air_fryer.jpg`: Thân máy đen mờ viền đồng Rose Gold, màn hình kính cảm ứng hiển thị nhiệt độ sống động.")
    add_body_p(doc, "+ Robot hút bụi lau nhà thông minh `robot_vacuum.jpg`: Trạm sạc tự giặt sấy giẻ, cảm biến laser LiDAR phát sáng xanh công nghệ cao.")
    add_body_p(doc, "+ Máy ép chậm nguyên quả `slow_juicer.jpg`: Trục xoay inox và ly nước ép cam tươi nguyên chất bắt mắt.")
    add_body_p(doc, "+ Máy lọc không khí thông minh `air_purifier.jpg`: Tháp tròn tối giản với màn hình tròn OLED hiển thị chỉ số bụi mịn PM2.5.")
    add_body_p(doc, "+ Ấm siêu tốc giữ nhiệt `electric_kettle.jpg`: Thân thủy tinh Borosilicate chịu nhiệt phát sáng LED xanh dương khi sôi.")
    add_body_p(doc, "+ Máy pha cà phê Espresso `coffee_machine.jpg`: Vỏ kim loại bóng bẩy, đồng hồ đo áp suất 20 Bar cổ điển và lớp bọt Crema mịn màng.")
    add_body_p(doc, "+ Máy rửa bát để bàn `countertop_dishwasher.jpg`: Cửa kính cường lực sang trọng nhìn rõ bát đĩa sáng bóng.")

    add_body_p(doc, "3. Hiệu ứng chuyển động (Transitions & Keyframe Animations):", bold_prefix="c. ")
    add_body_p(doc, "Để website luôn mang lại cảm giác sống động và phản hồi tức thì với thao tác người dùng, các hiệu ứng chuyển động CSS được tích hợp đồng bộ:")
    add_body_p(doc, "- Hiệu ứng lơ lửng Hero (`@keyframes floatHero`): Hình ảnh sản phẩm chủ đạo chuyển động nhịp nhàng theo chu kỳ mượt mà, thu hút ánh nhìn đầu tiên của khách hàng.")
    add_body_p(doc, "- Hiệu ứng phóng to thẻ và nâng bóng (`Card Hover Lift & Zoom`): Khi rê chuột vào thẻ sản phẩm, hình ảnh phóng to nhẹ 108% với đường cong chuyển động `cubic-bezier(0.16, 1, 0.3, 1)`, đồng thời các nút xem nhanh và thêm giỏ hàng trượt lên êm ái.")
    add_body_p(doc, "- Giỏ hàng trượt (Slide-out Cart Drawer): Bảng giỏ hàng trượt từ cạnh phải màn hình với lớp màn mờ nền (Backdrop Blur), cập nhật số lượng và tổng tiền tức thì qua AJAX mà không cần tải lại trang.")
    add_body_p(doc, "- Thông báo Toast (`Toast Notification Engine`): Khi người dùng thêm hàng vào giỏ hoặc gửi nhận xét, thông báo dạng thẻ nổi trượt từ góc phải dưới màn hình kèm thanh trạng thái tự động biến mất sau 3.5 giây.")
    add_body_p(doc, "- Đồng hồ đếm ngược giờ vàng (Flash Sale Countdown): Hiển thị thời gian Giờ : Phút : Giây cập nhật theo từng giây với hiệu ứng số rung nhẹ.")

    add_body_p(doc, "4. Phân tích các phân hệ màn hình chức năng:", bold_prefix="d. ")
    add_body_p(doc, "- Trang chủ (`index.php`): Gồm Hero Banner giới thiệu sản phẩm nổi bật, Lưới danh mục với biểu tượng xoay 360 độ khi hover, Khối Flash Sale giờ vàng, Sản phẩm bán chạy nhất, Khối banner mã giảm giá độc quyền và Đánh giá từ khách hàng thực tế.")
    add_body_p(doc, "- Trang danh mục sản phẩm (`products.php`): Tích hợp bộ lọc đa năng tại cột bên trái (lọc theo danh mục con, lọc theo mức giá từ dưới 1 triệu đến trên 6 triệu), ô tìm kiếm trực tiếp và thanh công cụ sắp xếp giá tăng/giảm.")
    add_body_p(doc, "- Trang chi tiết sản phẩm (`product-detail.php`): Thư viện ảnh phóng to khi hover, thông tin giá khuyến mãi kèm phần trăm tiết kiệm, bộ chọn số lượng tăng giảm mượt mà, bảng thông số kỹ thuật phân dòng chi tiết, danh sách đánh giá của khách hàng và form gửi nhận xét mới.")
    add_body_p(doc, "- Trang giỏ hàng & thanh toán (`cart.php`, `checkout.php`): Bảng danh sách mặt hàng, tính năng nhập mã giảm giá Voucher có kiểm tra điều kiện áp dụng, lựa chọn phương thức thanh toán linh hoạt (COD, VietQR, Ví điện tử).")
    add_body_p(doc, "- Trang hóa đơn & theo dõi đơn hàng (`order-success.php`): Hiển thị biểu tượng xác nhận đơn thành công, mã đơn hàng ngẫu nhiên duy nhất, tiến trình vận chuyển 4 bước mô phỏng và bảng chi tiết các món hàng.")
    add_body_p(doc, "- Phân hệ Quản trị Admin (`admin/index.php`, `admin/products.php`, `admin/orders.php`, `admin/categories.php`): Giao diện quản trị hiện đại với 4 thẻ chỉ số doanh thu/đơn hàng, bảng cập nhật trạng thái đơn hàng (Chờ xử lý, Đang giao, Hoàn thành), form thêm/sửa sản phẩm và quản lý danh mục.")

    # 3.2 Kết luận
    add_header_styled(doc, "3.2 Kết luận", level=2)
    add_body_p(doc, "Sau quá trình nghiên cứu lý thuyết chuyên sâu, phân tích yêu cầu kỹ thuật và hiện thực hóa mã nguồn, dự án xây dựng website thương mại điện tử đồ gia dụng DatCyber bằng ngôn ngữ PHP và hệ quản trị cơ sở dữ liệu MySQL đã hoàn thành xuất sắc toàn bộ các mục tiêu đề ra.")

    add_body_p(doc, "1. Các kết quả nổi bật đạt được:", bold_prefix="a. ")
    add_body_p(doc, "- Về mặt kiến trúc và kỹ thuật: Xây dựng thành công hệ thống web hoàn chỉnh với mã nguồn PHP 8.x hướng cấu trúc rõ ràng, sử dụng kết nối cơ sở dữ liệu PDO an toàn tuyệt đối trước các nguy cơ tấn công SQL Injection. Cơ sở dữ liệu MySQL được thiết kế chuẩn mực 3NF với các ràng buộc khóa ngoại chặt chẽ và cơ chế Transaction đảm bảo toàn vẹn dữ liệu đơn hàng.")
    add_body_p(doc, "- Về mặt giao diện và trải nghiệm người dùng (UI/UX): Đạt chuẩn thiết kế hiện đại với bố cục hài hòa, màu sắc sang trọng, hệ thống icon FontAwesome 6 trực quan, hiệu ứng CSS Transitions và Keyframe Animations mượt mà, cùng bộ ảnh sản phẩm AI sắc nét và đồng bộ.")
    add_body_p(doc, "- Về mặt nghiệp vụ thương mại điện tử: Hoàn thiện đầy đủ quy trình bán hàng khép kín: Tìm kiếm -> Lọc sản phẩm -> Xem chi tiết -> Thêm giỏ hàng tương tác AJAX -> Áp dụng mã Voucher -> Đặt hàng thanh toán đa kênh -> Theo dõi trạng thái đơn hàng -> Đánh giá sao phản hồi.")
    add_body_p(doc, "- Về mặt quản trị: Xây dựng phân hệ Admin chuyên nghiệp giúp quản lý đơn hàng, theo dõi doanh thu và thêm sửa xóa sản phẩm nhanh chóng, thuận tiện.")

    add_body_p(doc, "2. Hạn chế còn tồn tại:", bold_prefix="b. ")
    add_body_p(doc, "- Hiện tại các phương thức thanh toán trực tuyến (VietQR, MoMo, VNPay) đang được xử lý dưới dạng mô phỏng quy trình xác nhận; chưa tích hợp trực tiếp Webhook phản hồi tức thời từ cổng thanh toán của ngân hàng thương mại.")
    add_body_p(doc, "- Hệ thống chưa tích hợp phân hệ tính toán phí vận chuyển tự động theo API địa lý của các đơn vị giao hàng như Giao Hàng Nhanh (GHN) hay Viettel Post.")

    add_body_p(doc, "3. Hướng phát triển và mở rộng trong tương lai:", bold_prefix="c. ")
    add_body_p(doc, "- Tích hợp trợ lý ảo AI Chatbot thông minh hỗ trợ tư vấn sản phẩm, so sánh thông số kỹ thuật gia dụng và giải đáp thắc mắc cho khách hàng 24/7.")
    add_body_p(doc, "- Kết nối API cổng thanh toán thực tế (VNPAY Sandbox / MoMo Payment Gateway) và tích hợp tổng đài gửi tin nhắn SMS / Email Brandname tự động thông báo trạng thái đơn hàng.")
    add_body_p(doc, "- Xây dựng ứng dụng di động DatCyber Mobile App (Flutter hoặc React Native) đồng bộ cơ sở dữ liệu với website để tăng cường trải nghiệm đa kênh (Omnichannel).")

    # ================= TÀI LIỆU THAM KHẢO =================
    doc.add_page_break()
    add_header_styled(doc, "TÀI LIỆU THAM KHẢO", level=1)
    
    refs = [
        "1. The PHP Documentation Group (2024), 'PHP: Hypertext Preprocessor Official Manual', truy cập tại https://www.php.net/docs.php.",
        "2. Oracle Corporation (2024), 'MySQL 8.0 Reference Manual', Oracle Documentation Library.",
        "3. Robin Nixon (2021), 'Learning PHP, MySQL & JavaScript: With jQuery, CSS & HTML5', 6th Edition, O'Reilly Media.",
        "4. Jon Duckett (2014), 'HTML and CSS: Design and Build Websites', John Wiley & Sons, Inc.",
        "5. Matt Zandstra (2020), 'PHP 8 Objects, Patterns, and Practice: Mastering OO Enhancements, Design Patterns, and Test-Driven Development', Apress.",
        "6. MDN Web Docs (2026), 'CSS Transitions and Animations Guide', Mozilla Developer Network, truy cập tại https://developer.mozilla.org.",
        "7. Bootstrap Team (2024), 'Bootstrap 5.3 Framework Documentation', truy cập tại https://getbootstrap.com."
    ]
    for r in refs:
        add_body_p(doc, r, space_after=8)

    # Save document
    import os
    outputPath = os.path.join(os.path.dirname(os.path.abspath(__file__)), "BaoCao.docx")
    doc.save(outputPath)
    print(f"Report generated successfully at: {outputPath}")

if __name__ == "__main__":
    generate_doc()
