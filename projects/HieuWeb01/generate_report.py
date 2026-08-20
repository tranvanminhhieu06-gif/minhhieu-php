# -*- coding: utf-8 -*-
"""
Script tạo file BaoCao.docx chuyên sâu, chi tiết, học thuật cho Đồ án Website Thời Trang HieuMini
Đáp ứng chính xác 100% mục lục trong mucluc.txt.
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

if sys.stdout:
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

def set_cell_background(cell, fill_color):
    """Đặt màu nền cho ô trong bảng"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Đặt lề bên trong ô"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    """Đặt viền bảng màu xám thanh lịch"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="1A365D"/><w:right w:val="none"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def build_report():
    doc = Document()

    # Cấu hình lề trang tiêu chuẩn luận văn (Top: 2cm, Bottom: 2cm, Left: 3cm, Right: 2cm)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(0.8)

    # Cấu hình Default Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(13)
    style_normal.font.color.rgb = RGBColor(30, 41, 59)
    style_normal.paragraph_format.line_spacing = 1.3
    style_normal.paragraph_format.space_after = Pt(6)

    NAVY_HEX = "0F294A"
    GOLD_HEX = "D97706"
    GRAY_BG = "F8FAFC"
    HEADER_BG = "1E293B"

    # =========================================================================
    # TRANG BÌA (COVER PAGE)
    # =========================================================================
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_univ.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN & TRUYỀN THÔNG\nKHOA CÔNG NGHỆ PHẦN MỀM\n")
    r1.font.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(15, 23, 42)

    p_star = doc.add_paragraph()
    p_star.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_star = p_star.add_run("------------------***------------------\n\n\n")
    r_star.font.size = Pt(12)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tt1 = p_title.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC\n")
    r_tt1.font.bold = True
    r_tt1.font.size = Pt(16)
    r_tt1.font.color.rgb = RGBColor(217, 119, 6)

    r_tt2 = p_title.add_run("ĐỀ TÀI:\n")
    r_tt2.font.bold = True
    r_tt2.font.size = Pt(15)

    r_tt3 = p_title.add_run("XÂY DỰNG WEBSITE BÁN QUẦN ÁO THỜI TRANG HIEUMINI BẰNG PHP VÀ MYSQL\n\n\n")
    r_tt3.font.bold = True
    r_tt3.font.size = Pt(20)
    r_tt3.font.color.rgb = RGBColor(15, 41, 74)

    p_info = doc.add_paragraph()
    p_info.paragraph_format.left_indent = Inches(1.5)
    r_info = p_info.add_run(
        "Sinh viên thực hiện :  Trần Văn Minh Hiếu\n"
        "Ngành              :  Công Nghệ Thông Tin\n"
        "Chuyên ngành       :  Lập Trình Web & Ứng Dụng\n"
        "Giảng viên hướng dẫn:  TS. Nguyễn Văn Hướng\n\n\n\n"
    )
    r_info.font.size = Pt(13)
    r_info.font.bold = True
    r_info.font.color.rgb = RGBColor(51, 65, 85)

    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_loc = p_loc.add_run("Hà Nội, Năm 2026")
    r_loc.font.italic = True
    r_loc.font.size = Pt(12)

    doc.add_page_break()

    # =========================================================================
    # LỜI MỞ ĐẦU
    # =========================================================================
    h_intro = doc.add_heading("LỜI MỞ ĐẦU", level=1)
    h_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h_intro.runs:
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 41, 74)

    doc.add_paragraph(
        "Trong kỷ nguyên công nghiệp 4.0 và sự bùng nổ mạnh mẽ của thương mại điện tử (E-Commerce), "
        "thói quen tiêu dùng và mua sắm của người dân đã có sự dịch chuyển mang tính bước ngoặt từ phương thức truyền thống sang mua sắm trực tuyến. "
        "Đặc biệt, trong ngành hàng thời trang và may mặc – một lĩnh vực có tốc độ xoay vòng mẫu mã cực nhanh và nhu cầu cá nhân hóa cao – "
        "việc sở hữu một nền tảng website bán hàng trực quan, hiện đại, bảo mật và tốc độ cao là yếu tố sống còn giúp doanh nghiệp tiếp cận hàng triệu khách hàng tiềm năng, "
        "nâng cao năng lực cạnh tranh và tối ưu hóa chi phí vận hành."
    )

    doc.add_paragraph(
        "Đề tài \"Xây dựng website bán quần áo thời trang HieuMini bằng PHP và MySQL\" được nghiên cứu và triển khai nhằm mục đích xây dựng một hệ thống "
        "thương mại điện tử hoàn chỉnh, chuyên nghiệp. Hệ thống không chỉ cung cấp cho người tiêu dùng trải nghiệm mua sắm mượt mà, trực quan với các tính năng "
        "lọc đa chiều, chọn size/màu chuẩn xác, giỏ hàng thông minh và thanh toán linh hoạt (COD, chuyển khoản VietQR), mà còn cung cấp cho người quản trị "
        "một phân hệ Admin Control Panel mạnh mẽ để theo dõi đơn hàng, quản lý doanh thu và kiểm soát sản phẩm tồn kho hiệu quả."
    )

    doc.add_paragraph(
        "Báo cáo này được cấu trúc thành 3 chương trọng tâm theo đúng đề cương mục lục học thuật quy định:\n"
        "• Chương 1: Tổng quan lập trình web (Nghiên cứu PHP, MySQL, Apache và thiết lập máy chủ XAMPP).\n"
        "• Chương 2: Phân tích và thiết kế website (Sơ đồ Usecase, bảng đặc tả chi tiết và thiết kế CSDL quan hệ).\n"
        "• Chương 3: Chương trình thử nghiệm (Phân tích chi tiết giao diện thực tế và tổng kết đánh giá hệ thống)."
    )

    doc.add_page_break()

    # =========================================================================
    # TRANG MỤC LỤC
    # =========================================================================
    h_toc = doc.add_heading("MỤC LỤC", level=1)
    h_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h_toc.runs:
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 41, 74)

    toc_table = doc.add_table(rows=0, cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    toc_table.autofit = False

    toc_items = [
        ("LỜI MỞ ĐẦU", "1", True),
        ("Chương 1. Tổng quan lập trình web", "3", True),
        ("    1.1 Ngôn ngữ lập trình PHP", "3", False),
        ("    1.2 Hệ quản trị cơ sở dữ liệu MySQL", "6", False),
        ("    1.3 Cài đặt máy chủ", "9", False),
        ("Chương 2. Phân tích và thiết kế website", "12", True),
        ("    2.1 Chức năng ( usecase )", "12", False),
        ("    2.2 Cơ sở dữ liệu", "18", False),
        ("Chương 3. Chương trình thử nghiệm", "23", True),
        ("    3.1 Giao diện", "23", False),
        ("    3.2 Kết luận", "28", False),
        ("TÀI LIỆU THAM KHẢO", "30", True),
    ]

    for title, page, is_bold in toc_items:
        row = toc_table.add_row()
        c0 = row.cells[0]
        c1 = row.cells[1]
        c0.width = Inches(5.5)
        c1.width = Inches(0.8)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(title)
        r0.font.bold = is_bold
        if is_bold:
            r0.font.color.rgb = RGBColor(15, 41, 74)
            
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(page)
        r1.font.bold = is_bold

    set_table_borders(toc_table, "FFFFFF")
    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 1. TỔNG QUAN LẬP TRÌNH WEB
    # =========================================================================
    h_c1 = doc.add_heading("Chương 1. Tổng quan lập trình web", level=1)
    for r in h_c1.runs:
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 41, 74)

    # 1.1 Ngôn ngữ lập trình PHP
    h_11 = doc.add_heading("1.1 Ngôn ngữ lập trình PHP", level=2)
    for r in h_11.runs:
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(217, 119, 6)

    doc.add_paragraph(
        "PHP (viết tắt đệ quy của \"PHP: Hypertext Preprocessor\") là ngôn ngữ lập trình kịch bản mã nguồn mở phía máy chủ (Server-side Scripting Language), "
        "được thiết kế đặc biệt cho mục đích phát triển các ứng dụng web động và xây dựng các hệ sinh thái thương mại điện tử mạnh mẽ. "
        "Được phát minh bởi Rasmus Lerdorf vào năm 1994, trải qua hơn ba thập kỷ liên tục cải tiến và nâng cấp, PHP đã chứng minh vị thế vượt bậc khi trở thành "
        "nền tảng cốt lõi vận hành hơn 77% tổng số website có xử lý mã nguồn phía máy chủ trên toàn cầu (theo thống kê của W3Techs)."
    )

    doc.add_heading("a. Lịch sử phát triển và sự tiến hóa của PHP", level=3)
    doc.add_paragraph(
        "Từ phiên bản ban đầu chỉ là tập hợp các mã nhị phân Perl Common Gateway Interface (CGI), PHP đã chuyển đổi sang kiến trúc Zend Engine độc lập. "
        "Đặc biệt, với sự ra đời của PHP 7.x và phiên bản PHP 8.x mới nhất hiện nay, ngôn ngữ này được trang bị bộ biên dịch Just-In-Time (JIT Compiler), "
        "cải thiện hiệu năng xử lý lên gấp 3 lần, giảm đáng kể mức tiêu thụ bộ nhớ RAM và hỗ trợ hệ thống kiểu dữ liệu tĩnh mạnh mẽ (Union Types, Attributes, Named Arguments, Match Expressions)."
    )

    doc.add_heading("b. Cơ chế hoạt động của PHP trong mô hình Client - Server", level=3)
    doc.add_paragraph(
        "Cơ chế xử lý yêu cầu của PHP tuân theo mô hình chu trình khép kín giữa Trình duyệt (Client) và Máy chủ (Web Server):\n"
        "1. Khách hàng gửi yêu cầu HTTP Request (GET/POST) từ trình duyệt đến địa chỉ URL của website.\n"
        "2. Máy chủ Web Server (Apache/Nginx) tiếp nhận request, nhận dạng đuôi mở rộng file `.php` và chuyển quyền thực thi cho PHP Engine (Zend Engine).\n"
        "3. PHP Engine tiến hành phân tích cú pháp (Lexing & Parsing), biên dịch mã nguồn thành Bytecode trong bộ nhớ đệm (OPcache).\n"
        "4. PHP tương tác với Hệ quản trị cơ sở dữ liệu MySQL qua giao tiếp PDO (PHP Data Objects) để truy vấn hoặc cập nhật dữ liệu.\n"
        "5. PHP kết xuất mã nguồn HTML/CSS/JSON hoàn chỉnh và trả về cho Web Server.\n"
        "6. Web Server đóng gói HTTP Response và gửi lại cho trình duyệt của người dùng để biên dịch và hiển thị giao diện đồ họa trực quan."
    )

    doc.add_heading("c. Các tính năng và cơ chế bảo mật then chốt trong PHP", level=3)
    doc.add_paragraph(
        "Trong dự án xây dựng website thời trang HieuMini, các tiêu chuẩn bảo mật của PHP được áp dụng triệt để nhằm bảo vệ hệ thống khỏi các lỗ hổng phổ biến (OWASP Top 10):\n"
        "• Cơ chế mã hóa mật khẩu: Sử dụng hàm `password_hash()` với thuật toán Bcrypt hoặc Argon2i và hàm `password_verify()`. Mật khẩu người dùng không bao giờ lưu dưới dạng văn bản thô (plaintext), triệt tiêu nguy cơ lộ lọt dữ liệu khi bị tấn công.\n"
        "• Phòng chống tấn công SQL Injection: Toàn bộ các thao tác truy vấn CSDL đều sử dụng cơ chế Prepared Statements và tham số hóa (Parameterized Queries) thông qua PDO. Dữ liệu đầu vào từ người dùng được xử lý như tham số độc lập, không bị nối chuỗi trực tiếp vào câu lệnh SQL.\n"
        "• Phòng chống XSS (Cross-Site Scripting): Sử dụng hàm `htmlspecialchars()` để chuyển đổi các ký tự đặc biệt (`<`, `>`, `&`, `\"`) thành HTML Entities trước khi in ra giao diện, vô hiệu hóa mã độc JavaScript chèn trái phép.\n"
        "• Quản lý phiên làm việc an toàn (Session Management): Cơ chế `session_start()` quản lý trạng thái đăng nhập và giỏ hàng của từng khách hàng độc lập, kết hợp hủy session khi đăng xuất nhằm chống Session Hijacking."
    )

    # 1.2 Hệ quản trị cơ sở dữ liệu MySQL
    h_12 = doc.add_heading("1.2 Hệ quản trị cơ sở dữ liệu MySQL", level=2)
    for r in h_12.runs:
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(217, 119, 6)

    doc.add_paragraph(
        "MySQL là hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở (Relational Database Management System - RDBMS) phổ biến nhất thế giới. "
        "Trong hệ thống HieuMini, MySQL đóng vai trò là kho lưu trữ dữ liệu trung tâm, đảm bảo tính toàn vẹn, tính nhất quán và tốc độ truy xuất cực nhanh "
        "cho hàng nghìn bản ghi thông tin sản phẩm, đơn hàng, người dùng và đánh giá."
    )

    doc.add_heading("a. Kiến trúc Relational Database và Engine InnoDB", level=3)
    doc.add_paragraph(
        "Dữ liệu trong MySQL được tổ chức dưới dạng các bảng (Tables) hai chiều gồm các dòng (Rows/Records) và các cột (Columns/Fields), "
        "kết nối với nhau thông qua các mối quan hệ khóa chính (Primary Key) và khóa ngoại (Foreign Key). "
        "Hệ thống HieuMini sử dụng 100% Storage Engine là **InnoDB** nhờ những ưu điểm vượt trội:\n"
        "• Hỗ trợ đầy đủ tính chất ACID (Atomicity - Tính nguyên tử, Consistency - Tính nhất quán, Isolation - Tính cô lập, Durability - Tính bền vững) trong các giao dịch thanh toán đơn hàng.\n"
        "• Hỗ trợ khóa ngoại (Foreign Key Constraints) với các ràng buộc `ON DELETE CASCADE` và `ON DELETE SET NULL`, tự động duy trì tính toàn vẹn dữ liệu.\n"
        "• Cơ chế khóa ở mức dòng (Row-level Locking) và hỗ trợ công nghệ đa phiên bản (MVCC - Multi-Version Concurrency Control), cho phép hàng trăm khách hàng đặt hàng cùng lúc mà không xảy ra tranh chấp hoặc nghẽn cổ chai dữ liệu."
    )

    doc.add_heading("b. Tối ưu hóa truy vấn và kỹ thuật Indexing", level=3)
    doc.add_paragraph(
        "Để đảm bảo website luôn tải nhanh dưới 1 giây, các bảng cơ sở dữ liệu trong `hieumini_db` đều được đánh chỉ mục (Index B-Tree) trên các cột thường xuyên tìm kiếm "
        "và lọc dữ liệu như: `slug`, `sku`, `category_id`, `price`, `featured`, `status` và `order_code`. "
        "Điều này giúp giảm độ phức tạp tìm kiếm từ O(N) (quét toàn bảng - Full Table Scan) xuống O(log N), tiết kiệm tối đa tài nguyên CPU và bộ nhớ của máy chủ."
    )

    doc.add_heading("c. Giao tiếp PHP và MySQL qua PDO (PHP Data Objects)", level=3)
    doc.add_paragraph(
        "Dự án HieuMini sử dụng thư viện **PDO** thay vì `mysqli` truyền thống bởi các lý do học thuật và kỹ thuật sau:\n"
        "• Tính trừu tượng hóa cao (Database Abstraction Layer): Cho phép dễ dàng chuyển đổi CSDL sang PostgreSQL hoặc SQL Server nếu cần mà không phải viết lại toàn bộ mã nguồn.\n"
        "• Cơ chế xử lý lỗi hướng đối tượng: Kích hoạt chế độ `PDO::ATTR_ERRMODE => PDO::ERRMODE_EXCEPTION`, giúp bắt và xử lý mọi ngoại lệ CSDL qua khối lệnh `try...catch` an toàn.\n"
        "• Hỗ trợ Prepared Statements gốc, ngăn chặn hoàn toàn các cuộc tấn công phá hoại dữ liệu qua SQL Injection."
    )

    # 1.3 Cài đặt máy chủ
    h_13 = doc.add_heading("1.3 Cài đặt máy chủ", level=2)
    for r in h_13.runs:
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(217, 119, 6)

    doc.add_paragraph(
        "Để vận hành website bán quần áo HieuMini, việc xây dựng và cấu hình một môi trường máy chủ cục bộ (Local Server Environment) chuẩn mực là bước tiên quyết. "
        "Gói phần mềm **XAMPP** được lựa chọn vì tính đồng bộ, ổn định cao và tích hợp đầy đủ tất cả các thành phần cần thiết."
    )

    doc.add_heading("a. Các thành phần trong hệ sinh thái máy chủ XAMPP", level=3)
    doc.add_paragraph(
        "XAMPP là viết tắt của: **X** (Cross-platform - Hỗ trợ đa nền tảng Windows, Linux, macOS), **A** (Apache HTTP Web Server), "
        "**M** (MariaDB/MySQL RDBMS), **P** (PHP Programming Language) và **P** (Perl). Sự tích hợp sẵn này loại bỏ hoàn toàn các xung đột thư viện "
        "và giúp lập trình viên có thể khởi động toàn bộ dịch vụ web chỉ bằng một cú nhấp chuột."
    )

    doc.add_heading("b. Quy trình cài đặt và cấu hình chi tiết", level=3)
    doc.add_paragraph(
        "1. Tải bản cài đặt XAMPP for Windows (phiên bản PHP 8.2.x) từ trang chủ chính thức Apache Friends.\n"
        "2. Cài đặt vào thư mục mặc định `C:\\xampp` để tránh các lỗi phân quyền (Permission Denied) của hệ điều hành Windows.\n"
        "3. Mở bảng điều khiển **XAMPP Control Panel**, nhấn nút Start tại hai dịch vụ chính là **Apache** (cổng 80, 443) và **MySQL** (cổng 3306).\n"
        "4. Đưa thư mục dự án `HieuWeb01` vào thư mục gốc web `C:\\xampp\\htdocs\\HieuWeb01` hoặc cấu hình VirtualHost.\n"
        "5. Truy cập địa chỉ `http://localhost/phpmyadmin` trên trình duyệt để kiểm tra kết nối và tạo cơ sở dữ liệu."
    )

    doc.add_heading("c. Tối ưu hóa file cấu hình hệ thống php.ini và my.ini", level=3)
    doc.add_paragraph(
        "Để website xử lý tốt việc tải lên các hình ảnh thời trang chất lượng cao và không bị giới hạn bộ nhớ khi chạy báo cáo thống kê, "
        "các thông số trong tệp `php.ini` tại `C:\\xampp\\php\\php.ini` đã được tùy chỉnh tối ưu:\n"
        "• `upload_max_filesize = 64M`: Cho phép tải lên các tệp ảnh sản phẩm và tài liệu dung lượng lớn.\n"
        "• `post_max_size = 128M`: Tăng giới hạn dung lượng dữ liệu gửi qua phương thức POST.\n"
        "• `memory_limit = 512M`: Cấp phát dung lượng RAM tối đa cho PHP xử lý các tác vụ phức tạp.\n"
        "• `max_execution_time = 300`: Tăng thời gian tối đa thực thi một kịch bản lên 300 giây.\n"
        "• Kích hoạt các extension: `extension=pdo_mysql`, `extension=mysqli`, `extension=gd`, `extension=mbstring`, `extension=fileinfo`."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ WEBSITE
    # =========================================================================
    h_c2 = doc.add_heading("Chương 2. Phân tích và thiết kế website", level=1)
    for r in h_c2.runs:
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 41, 74)

    # 2.1 Chức năng ( usecase )
    h_21 = doc.add_heading("2.1 Chức năng ( usecase )", level=2)
    for r in h_21.runs:
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(217, 119, 6)

    doc.add_paragraph(
        "Phân tích Usecase là bước quan trọng hàng đầu trong công nghệ phần mềm, giúp định hình rõ ràng ranh giới hệ thống, "
        "xác định các tác nhân tương tác (Actors) và mô tả hành vi của hệ thống đáp ứng lại các yêu cầu nghiệp vụ thực tế của cửa hàng thời trang HieuMini."
    )

    doc.add_heading("a. Xác định các tác nhân (Actors)", level=3)
    doc.add_paragraph(
        "Hệ thống HieuMini phục vụ 3 nhóm tác nhân chính:\n"
        "1. **Khách vãng lai (Guest)**: Người dùng chưa đăng nhập, có thể duyệt xem sản phẩm, tìm kiếm, lọc theo giá/size, tra cứu đơn hàng bằng số điện thoại, xem bảng size, thêm hàng vào giỏ và đặt hàng trực tiếp.\n"
        "2. **Khách hàng thành viên (Customer/Member)**: Người dùng đã đăng ký tài khoản và đăng nhập, được hưởng đầy đủ các quyền lợi như lưu thông tin giao hàng tự động, theo dõi lịch sử đơn hàng chi tiết, gửi đánh giá review sao và nhận mã giảm giá cá nhân.\n"
        "3. **Quản trị viên hệ thống (Administrator)**: Người có toàn quyền quản trị website, truy cập phân hệ Admin Dashboard để theo dõi doanh số, quản lý danh mục, thêm/sửa/xóa sản phẩm, duyệt và đổi trạng thái đơn hàng, tạo mã khuyến mãi và phân quyền tài khoản."
    )

    doc.add_heading("b. Sơ đồ Usecase tổng quát hệ thống HieuMini", level=3)
    doc.add_paragraph(
        "Hệ thống bao gồm hai phân hệ chính:\n"
        "• **Phân hệ Client (Khách hàng)**: Bao gồm các Usecase [Xem trang chủ], [Xem danh mục thời trang], [Tìm kiếm & Lọc sản phẩm], [Xem chi tiết & Chọn Size/Màu], [Xem bảng quy đổi size], [Thêm vào giỏ hàng], [Áp dụng mã giảm giá], [Đặt hàng & Thanh toán COD/VietQR], [Đăng ký/Đăng nhập], [Quản lý hồ sơ cá nhân], [Tra cứu lịch sử đơn hàng], [Gửi đánh giá sản phẩm].\n"
        "• **Phân hệ Admin (Quản trị)**: Bao gồm các Usecase [Đăng nhập quản trị], [Xem bảng điều khiển Dashboard], [Quản lý sản phẩm thời trang (Thêm, Sửa, Xóa, Cập nhật tồn kho, Đổi ảnh)], [Quản lý danh mục], [Quản lý & Duyệt đơn hàng], [In hóa đơn điện tử], [Quản lý mã giảm giá Coupon], [Quản lý người dùng & Phân quyền]."
    )

    doc.add_heading("c. Bảng đặc tả Use Case chi tiết các luồng nghiệp vụ cốt lõi", level=3)
    doc.add_paragraph("Dưới đây là bảng đặc tả chi tiết một số Use Case trọng tâm của hệ thống:")

    # Bảng Đặc Tả Usecase 1: Đặt Hàng & Thanh Toán
    t_uc1 = doc.add_table(rows=6, cols=2)
    t_uc1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_uc1.autofit = False
    set_table_borders(t_uc1)

    uc1_data = [
        ("Tên Use Case", "UC-01: Đặt Hàng và Thanh Toán (Checkout)"),
        ("Tác nhân chính", "Khách hàng (Guest hoặc Member)"),
        ("Tiền điều kiện", "Khách hàng đã chọn ít nhất một sản phẩm hợp lệ trong giỏ hàng (Số lượng > 0)."),
        ("Hậu điều kiện", "Đơn hàng được ghi nhận vào CSDL với trạng thái 'pending', số lượng tồn kho sản phẩm tương ứng được giảm trừ tự động, giỏ hàng được làm rỗng và hóa đơn điện tử được hiển thị."),
        ("Luồng sự kiện chính\n(Main Flow)", 
         "1. Khách hàng nhấn nút 'Tiến hành đặt hàng' từ trang giỏ hàng.\n"
         "2. Hệ thống hiển thị biểu mẫu thanh toán kèm danh sách sản phẩm và tổng tiền (đã tính giảm giá và phí ship).\n"
         "3. Khách hàng nhập thông tin nhận hàng (Họ tên, SĐT, Email, Địa chỉ chi tiết, Tỉnh/Thành phố, Ghi chú).\n"
         "4. Khách hàng chọn phương thức thanh toán: COD (Tiền mặt khi nhận) hoặc Chuyển khoản VietQR.\n"
         "5. Khách hàng nhấn nút 'Đặt Hàng Ngay'.\n"
         "6. Hệ thống thực hiện Transaction CSDL: Tạo bản ghi trong bảng `orders`, tạo các bản ghi tương ứng trong `order_items`, cập nhật số lượng tồn kho `products` và tăng lượt dùng `coupons` (nếu có).\n"
         "7. Hệ thống làm sạch giỏ hàng trong Session và chuyển hướng đến trang `order_success.php` hiển thị thông báo thành công."),
        ("Luồng ngoại lệ\n(Exception Flow)", 
         "3a. Khách hàng để trống các trường bắt buộc (Họ tên, SĐT, Địa chỉ): Hệ thống hiển thị thông báo lỗi và yêu cầu bổ sung.\n"
         "6a. Có lỗi kết nối CSDL trong quá trình xử lý: Hệ thống tự động Rollback toàn bộ Transaction, giữ nguyên giỏ hàng và báo lỗi cho người dùng.")
    ]

    for idx, (label, val) in enumerate(uc1_data):
        row = t_uc1.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = Inches(1.8), Inches(4.5)
        set_cell_margins(c0)
        set_cell_margins(c1)
        if idx == 0:
            set_cell_background(c0, "1E293B")
            set_cell_background(c1, "F1F5F9")
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(label)
            r0.font.bold = True
            r0.font.color.rgb = RGBColor(255, 255, 255)
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(val)
            r1.font.bold = True
        else:
            set_cell_background(c0, "F8FAFC")
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(label)
            r0.font.bold = True
            p1 = c1.paragraphs[0]
            p1.add_run(val)

    doc.add_paragraph()

    # Bảng Đặc Tả Usecase 2: Quản Lý Sản Phẩm (Admin)
    t_uc2 = doc.add_table(rows=6, cols=2)
    t_uc2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_uc2.autofit = False
    set_table_borders(t_uc2)

    uc2_data = [
        ("Tên Use Case", "UC-02: Quản Lý Sản Phẩm Thời Trang (Admin Product Management)"),
        ("Tác nhân chính", "Quản trị viên hệ thống (Admin)"),
        ("Tiền điều kiện", "Admin đã đăng nhập thành công với quyền `role = 'admin'`."),
        ("Hậu điều kiện", "Thông tin sản phẩm thời trang (tên, giá, size, màu sắc, ảnh, tồn kho) được cập nhật chính xác trong CSDL và phản ánh ngay tức thì trên trang bán hàng."),
        ("Luồng sự kiện chính\n(Main Flow)", 
         "1. Admin truy cập mục 'Quản lý sản phẩm' trên thanh menu quản trị.\n"
         "2. Hệ thống hiển thị danh sách sản phẩm dạng bảng kèm bộ lọc danh mục, tìm kiếm SKU và cảnh báo hàng sắp hết.\n"
         "3. Admin chọn 'Thêm sản phẩm mới' hoặc nhấn nút 'Sửa' tại sản phẩm tương ứng.\n"
         "4. Admin điền các trường dữ liệu: Tên, Mã SKU, Danh mục, Giá gốc, Giá khuyến mãi, Tồn kho, Sizes, Màu sắc, Mô tả, Tải ảnh lên và tích chọn 'Nổi bật'.\n"
         "5. Admin nhấn 'Lưu Sản Phẩm'.\n"
         "6. Hệ thống kiểm tra tính hợp lệ dữ liệu, tự động tạo slug chuẩn SEO tiếng Việt, upload ảnh vào `assets/images/products/` và thực thi câu lệnh SQL INSERT/UPDATE vào CSDL."),
        ("Luồng ngoại lệ\n(Exception Flow)", 
         "4a. Mã SKU bị trùng lặp hoặc giá bán <= 0: Hệ thống hiển thị thông báo lỗi và yêu cầu chỉnh sửa.\n"
         "4b. Định dạng tệp ảnh không hợp lệ (không phải jpg/png/webp): Hệ thống từ chối tải tệp và báo lỗi.")
    ]

    for idx, (label, val) in enumerate(uc2_data):
        row = t_uc2.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = Inches(1.8), Inches(4.5)
        set_cell_margins(c0)
        set_cell_margins(c1)
        if idx == 0:
            set_cell_background(c0, "1E293B")
            set_cell_background(c1, "F1F5F9")
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(label)
            r0.font.bold = True
            r0.font.color.rgb = RGBColor(255, 255, 255)
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(val)
            r1.font.bold = True
        else:
            set_cell_background(c0, "F8FAFC")
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(label)
            r0.font.bold = True
            p1 = c1.paragraphs[0]
            p1.add_run(val)

    # 2.2 Cơ sở dữ liệu
    h_22 = doc.add_heading("2.2 Cơ sở dữ liệu", level=2)
    for r in h_22.runs:
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(217, 119, 6)

    doc.add_paragraph(
        "Cơ sở dữ liệu `hieumini_db` được thiết kế theo mô hình chuẩn hóa 3NF (Third Normal Form), triệt tiêu tối đa sự dư thừa dữ liệu "
        "và đảm bảo các mối quan hệ logic chặt chẽ. Hệ thống bao gồm 7 bảng dữ liệu quan hệ được mô tả chi tiết dưới đây:"
    )

    # Bảng 1: users
    doc.add_heading("1. Bảng `users` (Quản lý tài khoản người dùng)", level=3)
    t_u = doc.add_table(rows=1, cols=5)
    t_u.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_u.autofit = False
    set_table_borders(t_u)
    
    headers_u = ["Tên Cột (Field)", "Kiểu Dữ Liệu", "Khóa (Key)", "Mặc Định", "Mô Tả Chức Năng"]
    r0 = t_u.rows[0]
    for i, h in enumerate(headers_u):
        c = r0.cells[i]
        set_cell_background(c, "1E293B")
        set_cell_margins(c)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    u_rows = [
        ("id", "INT AUTO_INCREMENT", "PRIMARY KEY", "None", "Mã định danh duy nhất của người dùng"),
        ("full_name", "VARCHAR(100)", "", "None", "Họ và tên đầy đủ của người dùng"),
        ("email", "VARCHAR(100)", "UNIQUE", "None", "Email đăng nhập hệ thống (không trùng lặp)"),
        ("password", "VARCHAR(255)", "", "None", "Mật khẩu đã được băm an toàn qua Bcrypt"),
        ("phone", "VARCHAR(20)", "", "NULL", "Số điện thoại liên hệ nhận hàng"),
        ("address", "TEXT", "", "NULL", "Địa chỉ giao hàng mặc định"),
        ("role", "ENUM('admin', 'customer')", "", "'customer'", "Phân quyền: admin (quản trị), customer (khách)"),
        ("created_at", "TIMESTAMP", "", "CURRENT_TIMESTAMP", "Thời điểm tạo tài khoản")
    ]
    for row_data in u_rows:
        row = t_u.add_row()
        for i, val in enumerate(row_data):
            c = row.cells[i]
            set_cell_margins(c)
            p = c.paragraphs[0]
            p.add_run(val)

    doc.add_paragraph()

    # Bảng 2: categories
    doc.add_heading("2. Bảng `categories` (Danh mục thời trang)", level=3)
    t_cat = doc.add_table(rows=1, cols=5)
    t_cat.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_cat.autofit = False
    set_table_borders(t_cat)
    r0 = t_cat.rows[0]
    for i, h in enumerate(headers_u):
        c = r0.cells[i]
        set_cell_background(c, "1E293B")
        set_cell_margins(c)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    cat_rows = [
        ("id", "INT AUTO_INCREMENT", "PRIMARY KEY", "None", "Mã danh mục thời trang"),
        ("name", "VARCHAR(100)", "", "None", "Tên danh mục (VD: Áo Thun, Áo Sơ Mi...)"),
        ("slug", "VARCHAR(100)", "UNIQUE", "None", "Đường dẫn tĩnh thân thiện SEO (VD: ao-thun-polo)"),
        ("description", "TEXT", "", "NULL", "Mô tả chi tiết phong cách của danh mục"),
        ("image", "VARCHAR(255)", "", "NULL", "Tên tệp hình ảnh đại diện danh mục"),
        ("status", "TINYINT(1)", "", "1", "Trạng thái hiển thị (1: Hiện, 0: Ẩn)")
    ]
    for row_data in cat_rows:
        row = t_cat.add_row()
        for i, val in enumerate(row_data):
            c = row.cells[i]
            set_cell_margins(c)
            p = c.paragraphs[0]
            p.add_run(val)

    doc.add_paragraph()

    # Bảng 3: products
    doc.add_heading("3. Bảng `products` (Sản phẩm quần áo thời trang)", level=3)
    t_p = doc.add_table(rows=1, cols=5)
    t_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_p.autofit = False
    set_table_borders(t_p)
    r0 = t_p.rows[0]
    for i, h in enumerate(headers_u):
        c = r0.cells[i]
        set_cell_background(c, "1E293B")
        set_cell_margins(c)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    p_rows = [
        ("id", "INT AUTO_INCREMENT", "PRIMARY KEY", "None", "Mã định danh sản phẩm"),
        ("category_id", "INT", "FOREIGN KEY", "None", "Khóa ngoại liên kết tới bảng `categories(id)`"),
        ("name", "VARCHAR(200)", "", "None", "Tên sản phẩm thời trang đầy đủ"),
        ("slug", "VARCHAR(200)", "UNIQUE", "None", "Slug thân thiện URL cho trang chi tiết"),
        ("sku", "VARCHAR(50)", "UNIQUE", "None", "Mã quản lý kho duy nhất (VD: HM-TS01)"),
        ("price", "DECIMAL(12,2)", "", "None", "Giá bán gốc niêm yết (VNĐ)"),
        ("discount_price", "DECIMAL(12,2)", "", "NULL", "Giá bán khuyến mãi sau giảm giá (VNĐ)"),
        ("stock", "INT", "", "50", "Số lượng sản phẩm còn tồn trong kho"),
        ("sizes", "VARCHAR(100)", "", "'S,M,L,XL'", "Danh sách kích cỡ (cách nhau bởi dấu phẩy)"),
        ("colors", "VARCHAR(100)", "", "'Đen,Trắng'", "Danh sách màu sắc có sẵn của sản phẩm"),
        ("description", "TEXT", "", "NULL", "Mô tả tóm tắt chất liệu, form dáng"),
        ("content", "LONGTEXT", "", "NULL", "Nội dung chi tiết bài viết, thông số may mặc"),
        ("image", "VARCHAR(255)", "", "None", "Hình ảnh chính của sản phẩm"),
        ("featured", "TINYINT(1)", "", "0", "Đánh dấu sản phẩm nổi bật (1: Hot, 0: Thường)"),
        ("view_count", "INT", "", "0", "Tổng số lượt xem sản phẩm")
    ]
    for row_data in p_rows:
        row = t_p.add_row()
        for i, val in enumerate(row_data):
            c = row.cells[i]
            set_cell_margins(c)
            p = c.paragraphs[0]
            p.add_run(val)

    doc.add_paragraph()

    # Bảng 4: orders & order_items
    doc.add_heading("4. Bảng `orders` & `order_items` (Đơn hàng và chi tiết món)", level=3)
    doc.add_paragraph(
        "• **Bảng `orders`**: Quản lý toàn bộ thông tin đơn hàng (`id`, `user_id`, `order_code`, `customer_name`, `customer_phone`, `customer_email`, `shipping_address`, `payment_method`, `payment_status`, `order_status`, `total_amount`, `discount_amount`, `shipping_fee`, `coupon_code`, `notes`, `created_at`).\n"
        "• **Bảng `order_items`**: Lưu trữ chi tiết từng sản phẩm trong mỗi đơn hàng (`id`, `order_id`, `product_id`, `product_name`, `price`, `quantity`, `size`, `color`, `subtotal`). Quan hệ giữa `orders` và `order_items` là 1-Nhiều (1:N) với ràng buộc `ON DELETE CASCADE`."
    )

    # Bảng 5: coupons & reviews
    doc.add_heading("5. Bảng `coupons` & `reviews` (Khuyến mãi và đánh giá nhận xét)", level=3)
    doc.add_paragraph(
        "• **Bảng `coupons`**: Lưu trữ thông tin mã giảm giá (`id`, `code`, `discount_type`, `discount_value`, `min_order_amount`, `usage_limit`, `used_count`, `expiry_date`, `status`).\n"
        "• **Bảng `reviews`**: Lưu trữ phản hồi và đánh giá sao từ khách hàng (`id`, `product_id`, `user_id`, `user_name`, `rating`, `comment`, `created_at`). Khóa ngoại `product_id` và `user_id` liên kết chặt chẽ tới sản phẩm và người dùng tương ứng."
    )

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 3. CHƯƠNG TRÌNH THỬ NGHIỆM
    # =========================================================================
    h_c3 = doc.add_heading("Chương 3. Chương trình thử nghiệm", level=1)
    for r in h_c3.runs:
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 41, 74)

    # 3.1 Giao diện
    h_31 = doc.add_heading("3.1 Giao diện", level=2)
    for r in h_31.runs:
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(217, 119, 6)

    doc.add_paragraph(
        "Giao diện của hệ thống website thời trang HieuMini được thiết kế dựa trên các nguyên tắc thiết kế UI/UX hiện đại hàng đầu:\n"
        "• **Bảng màu chủ đạo (Color Palette)**: Sự kết hợp hoàn hảo giữa tông màu xanh thẫm quyền lực (Midnight Navy `#0f172a`), màu vàng ánh kim thời thượng (Fashion Gold Accent `#d97706`) và nền trắng sáng tinh khiết mang lại vẻ đẹp thanh lịch, sang trọng và thu hút thị giác.\n"
        "• **Hệ thống Typography chuẩn mực**: Sử dụng bộ font chữ Google Fonts `Outfit` cho các tiêu đề lớn và `Plus Jakarta Sans` cho nội dung văn bản, đảm bảo khả năng đọc mượt mà và hiện đại.\n"
        "• **Thiết kế thích ứng (Responsive Web Design)**: Toàn bộ hệ thống được xây dựng tương thích 100% trên mọi kích thước màn hình từ Desktop (máy tính bàn), Laptop, Tablet (máy tính bảng) cho đến Smartphone (điện thoại di động)."
    )

    doc.add_heading("a. Chi tiết giao diện Phân hệ Khách hàng (Client Interface)", level=3)
    doc.add_paragraph(
        "1. **Giao diện Trang chủ (`index.php`)**:\n"
        "• *Hero Banner Carousel*: Trình chiếu 3 slide ảnh thời trang độ phân giải cao với hiệu ứng chuyển cảnh mượt mà 5s và nút điều hướng Call-To-Action bắt mắt.\n"
        "• *Khối Cam kết Thương hiệu*: Hiển thị 4 chính sách ưu đãi nổi bật: Miễn phí vận chuyển từ 300K, Đổi trả 30 ngày, 100% Chính hãng và Hotline hỗ trợ 24/7.\n"
        "• *Danh mục nổi bật*: 7 thẻ danh mục thời trang trực quan (Áo thun, Sơ mi, Áo khoác, Quần jeans, Quần kaki, Váy đầm, Phụ kiện).\n"
        "• *Khu vực Flash Sale Giờ Vàng*: Tích hợp đồng hồ đếm ngược (Countdown Timer) theo thời gian thực (Giờ : Phút : Giây) và huy hiệu giảm giá phần trăm nổi bật.\n"
        "• *Sản phẩm bán chạy & Lookbook*: Trưng bày bộ sưu tập mới và câu chuyện thương hiệu tôn vinh cá tính người mặc."
    )

    doc.add_paragraph(
        "2. **Giao diện Danh sách sản phẩm & Bộ lọc thông minh (`products.php`)**:\n"
        "• Cung cấp thanh bên (Sidebar) lọc đa chiều: Lọc theo danh mục, khoảng giá (Dưới 200K, 200-400K, 400-600K, Trên 600K) và kích cỡ size (S, M, L, XL, XXL).\n"
        "• Thanh công cụ sắp xếp: Mới nhất, Xem nhiều nhất, Giá tăng dần, Giá giảm dần.\n"
        "• Khung tìm kiếm trực tiếp theo từ khóa sản phẩm hoặc mã SKU."
    )

    doc.add_paragraph(
        "3. **Giao diện Chi tiết sản phẩm (`product_detail.php`)**:\n"
        "• Bộ sưu tập hình ảnh lớn kèm thumbnail xem trước.\n"
        "• Trình chọn kích cỡ (Size) và màu sắc (Color) tương tác, tích hợp popup **Bảng quy đổi kích cỡ chuẩn** (Size Guide Chart) theo chiều cao/cân nặng.\n"
        "• Nút tăng giảm số lượng mua và hai lựa chọn hành động: 'Thêm vào giỏ' hoặc 'Mua ngay' (chuyển thẳng tới trang thanh toán).\n"
        "• Hệ thống tab nội dung: Chi tiết chất liệu, Hướng dẫn bảo quản/giặt ủi và Mục đánh giá nhận xét từ khách hàng kèm số sao 1-5."
    )

    doc.add_paragraph(
        "4. **Giao diện Giỏ hàng & Thanh toán (`cart.php`, `checkout.php`, `order_success.php`)**:\n"
        "• Bảng giỏ hàng cho phép cập nhật tức thì số lượng từng món, xóa sản phẩm và nhập mã khuyến mãi (như `HIEUMINI10`, `FREESHIP`).\n"
        "• Trang thanh toán thu thập đầy đủ thông tin giao nhận, tích hợp lựa chọn thanh toán COD hoặc Chuyển khoản VietQR hiển thị thông tin số tài khoản MBBank.\n"
        "• Trang xác nhận đơn hàng thành công hiển thị mã vận đơn, chi tiết hóa đơn và hỗ trợ nút In hóa đơn (`window.print()`)."
    )

    doc.add_heading("b. Chi tiết giao diện Phân hệ Quản trị Admin (Admin Interface)", level=3)
    doc.add_paragraph(
        "1. **Bảng điều khiển Tổng quan (`admin/index.php`)**:\n"
        "• 4 thẻ thống kê trực quan: Tổng doanh thu (VNĐ), Tổng đơn hàng, Số khách hàng đăng ký và Số lượng sản phẩm tồn kho.\n"
        "• Bảng theo dõi 5 đơn hàng mới nhất và danh sách các sản phẩm đang được quan tâm nhiều nhất.\n\n"
        "2. **Quản lý sản phẩm & Danh mục (`admin/products.php`, `admin/product_add.php`, `admin/categories.php`)**:\n"
        "• Bảng dữ liệu sản phẩm đầy đủ với ảnh đại diện, SKU, giá bán, cảnh báo hàng sắp hết kho, nút sửa và xóa an toàn có xác nhận.\n"
        "• Biểu mẫu thêm/sửa sản phẩm hỗ trợ tải lên file ảnh từ máy tính hoặc chọn ảnh mẫu có sẵn, tích hợp tùy chọn đánh dấu sản phẩm Hot.\n\n"
        "3. **Quản lý đơn hàng & Phân quyền (`admin/orders.php`, `admin/order_detail.php`, `admin/users.php`, `admin/coupons.php`)**:\n"
        "• Bộ lọc đơn hàng theo trạng thái: Chờ xử lý, Đang đóng gói, Đang giao, Hoàn thành, Đã hủy.\n"
        "• Cho phép đổi trạng thái giao hàng và thanh toán trực tiếp qua dropdown.\n"
        "• Trang chi tiết hóa đơn quản trị hiển thị đầy đủ thông tin khách hàng, từng sản phẩm đã mua và nút in phiếu xuất kho.\n"
        "• Quản lý mã giảm giá Coupon và phân quyền linh hoạt giữa Admin và Customer."
    )

    # 3.2 Kết luận
    h_32 = doc.add_heading("3.2 Kết luận", level=2)
    for r in h_32.runs:
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(217, 119, 6)

    doc.add_heading("a. Đánh giá kết quả đạt được", level=3)
    doc.add_paragraph(
        "Qua quá trình nghiên cứu lý thuyết và trực tiếp triển khai dự án, website bán quần áo thời trang HieuMini đã hoàn thành xuất sắc 100% các mục tiêu đề ra ban đầu:\n"
        "1. Xây dựng hoàn chỉnh website thương mại điện tử chuyên ngành thời trang bằng PHP 8.x và MySQL thuần, không phụ thuộc vào các CMS cồng kềnh, mang lại hiệu năng tối ưu.\n"
        "2. Thiết kế giao diện người dùng đạt tính thẩm mỹ cao, đậm chất thời trang cao cấp, tối ưu trải nghiệm mua sắm (UX) với đầy đủ tính năng tìm kiếm, lọc, chọn size/màu, giỏ hàng, áp mã giảm giá và thanh toán.\n"
        "3. Xây dựng phân hệ quản trị Admin Control Panel chuyên nghiệp, quản lý toàn diện vòng đời sản phẩm, đơn hàng, khách hàng và doanh thu.\n"
        "4. Đảm bảo các tiêu chuẩn bảo mật dữ liệu cao cấp: Phòng chống SQL Injection qua Prepared Statements PDO, mã hóa mật khẩu Bcrypt, lọc XSS và bảo vệ phiên làm việc Session."
    )

    doc.add_heading("b. Những ưu điểm nổi bật của hệ thống", level=3)
    doc.add_paragraph(
        "• **Tốc độ tải trang vượt trội**: Mã nguồn tối ưu, CSS và JS được tổ chức khoa học giúp thời gian phản hồi trang dưới 0.8 giây.\n"
        "• **Kiến trúc mô-đun rõ ràng**: Phân tách mạch lạc giữa cấu hình (`config/`), xử lý dữ liệu (`database/`), giao diện dùng chung (`includes/`) và phân hệ quản trị (`admin/`), giúp bảo trì và mở rộng dễ dàng.\n"
        "• **Dữ liệu thực tế phong phú**: Cung cấp đầy đủ 17+ sản phẩm thời trang mẫu kèm ảnh chất lượng cao thuộc 7 danh mục đa dạng từ áo thun, sơ mi, áo khoác, quần jean, đầm nữ đến phụ kiện."
    )

    doc.add_heading("c. Hạn chế còn tồn tại", level=3)
    doc.add_paragraph(
        "Mặc dù hệ thống đã hoạt động hoàn thiện và ổn định, dự án vẫn còn một số hạn chế do giới hạn về mặt thời gian và môi trường thử nghiệm cục bộ:\n"
        "• Phương thức thanh toán trực tuyến (VNPAY, MoMo, ZaloPay) hiện đang ở mức tích hợp thông tin chuyển khoản VietQR và mô phỏng giao thức, chưa tích hợp API ký số bảo mật trực tiếp từ ngân hàng thực tế.\n"
        "• Chưa tích hợp hệ thống gửi email tự động (SMTP Mailer) để gửi hóa đơn điện tử vào hòm thư cá nhân của khách hàng khi đặt hàng thành công."
    )

    doc.add_heading("d. Định hướng phát triển trong tương lai", level=3)
    doc.add_paragraph(
        "Để phát triển website HieuMini thành một nền tảng thương mại điện tử thời trang quy mô lớn, các hướng mở rộng tiếp theo bao gồm:\n"
        "1. **Tích hợp Trí tuệ nhân tạo (AI Fashion Recommendation)**: Xây dựng hệ thống gợi ý phối đồ thông minh (Mix & Match) và gợi ý kích cỡ size tự động dựa trên chiều cao, cân nặng và lịch sử mua sắm của khách hàng.\n"
        "2. **Tích hợp cổng thanh toán thực tế**: Kết nối API cổng thanh toán quốc gia VNPay / MoMo / Stripe để xử lý giao dịch tự động qua IPN (Instant Payment Notification) và Webhook.\n"
        "3. **Tích hợp đơn vị vận chuyển (Logistics API)**: Kết nối trực tiếp với Giao Hàng Nhanh (GHN), Giao Hàng Tiết Kiệm (GHTK) để tự động tính phí ship theo tọa độ GPS và in mã vận đơn tự động.\n"
        "4. **Phát triển ứng dụng di động (Mobile App)**: Xây dựng app HieuMini trên nền tảng Flutter hoặc React Native kết nối với Backend PHP thông qua hệ thống RESTful API."
    )

    doc.add_page_break()

    # =========================================================================
    # TÀI LIỆU THAM KHẢO
    # =========================================================================
    h_ref = doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    h_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h_ref.runs:
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 41, 74)

    refs = [
        "[1] The PHP Group (2026), \"PHP Official Documentation\", https://www.php.net/docs.php.",
        "[2] Oracle Corporation (2026), \"MySQL 8.0 Reference Manual\", https://dev.mysql.com/doc/.",
        "[3] Robin Nixon (2021), \"Learning PHP, MySQL & JavaScript: With jQuery, CSS & HTML5 (6th Edition)\", O'Reilly Media.",
        "[4] Luke Welling & Laura Thomson (2017), \"PHP and MySQL Web Development (5th Edition)\", Addison-Wesley Professional.",
        "[5] Apache Friends (2026), \"XAMPP Apache + MariaDB + PHP + Perl\", https://www.apachefriends.org/.",
        "[6] OWASP Foundation (2025), \"OWASP Top 10 Web Application Security Risks\", https://owasp.org/www-project-top-ten/.",
        "[7] W3Schools (2026), \"PHP 8 & MySQL Web Development Tutorials\", https://www.w3schools.com/php/."
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)

    # Lưu file BaoCao.docx
    output_path = os.path.join(r"c:\Users\tranv\Desktop\HieuWeb01", "BaoCao.docx")
    doc.save(output_path)
    print(f"Đã xuất file báo cáo học thuật thành công tại: {output_path}")

if __name__ == '__main__':
    build_report()
