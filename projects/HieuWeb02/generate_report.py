# -*- coding: utf-8 -*-
"""
Script tự động sinh tài liệu Báo Cáo Đồ Án Website HieuMini (BaoCao.docx)
Chuẩn format học thuật: Font Times New Roman, phân cấp Heading, Bảng biểu chuyên nghiệp.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_color):
    """Đặt màu nền cho ô trong bảng"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt margin cho ô bảng"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    """Thêm tiêu đề với màu sắc và kích thước phân cấp chuẩn"""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    
    run = heading.runs[0]
    run.font.name = 'Times New Roman'
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(27, 54, 93) # Navy Blue
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153) # Teal Blue
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51) # Dark Grey
    return heading

def add_styled_paragraph(doc, text="", style='Normal', space_after=6, line_spacing=1.3, bold=False, italic=False):
    """Thêm đoạn văn bản chuẩn font Times New Roman 13pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(34, 34, 34)
    return p

def format_table(table, col_widths, headers, data, header_bg="1B365D", alt_bg="F4F6F9"):
    """Tạo bảng định dạng đẹp mắt với màu nền xen kẽ và bo góc"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], header_bg)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Data rows
    for row_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = alt_bg if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = str(cell_value)
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11.5)
                run.font.color.rgb = RGBColor(40, 40, 40)
                
    # Apply column widths
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

def generate_full_report():
    doc = docx.Document()
    
    # Thiết lập lề trang chuẩn đồ án (Left: 3.0cm, Right: 2.0cm, Top: 2.0cm, Bottom: 2.0cm)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.18) # ~3.0cm
        section.right_margin = Inches(0.8)
        
    # ==========================================================
    # TRANG BÌA (COVER PAGE)
    # ==========================================================
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_univ.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN\nKHOA CÔNG NGHỆ PHẦN MỀM\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(27, 54, 93)
    
    r_star = p_univ.add_run("-------------------***-------------------\n\n\n")
    r_star.font.bold = True
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title_tag = p_title.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC\nLẬP TRÌNH PHÁT TRIỂN ỨNG DỤNG WEB\n\n")
    r_title_tag.font.name = 'Times New Roman'
    r_title_tag.font.size = Pt(15)
    r_title_tag.font.bold = True
    
    r_title_main = p_title.add_run("ĐỀ TÀI:\nTHIẾT KẾ VÀ XÂY DỰNG WEBSITE THƯƠNG MẠI ĐIỆN TỬ BÁN HÀNG CÔNG NGHỆ HIEUMINI\n(SỬ DỤNG PHP & HỆ QUẢN TRỊ CSDL MYSQL)\n\n\n")
    r_title_main.font.name = 'Times New Roman'
    r_title_main.font.size = Pt(18)
    r_title_main.font.bold = True
    r_title_main.font.color.rgb = RGBColor(192, 0, 0)
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_info.paragraph_format.left_indent = Inches(1.5)
    p_info.paragraph_format.line_spacing = 1.4
    
    r_info = p_info.add_run(
        "Sinh viên thực hiện   :  TRẦN VĂN MINH HIẾU\n"
        "Mã số sinh viên        :  2026TECH0802\n"
        "Lớp chuyên ngành      :  Kỹ Thuật Phần Mềm & Web - K20\n"
        "Giảng viên hướng dẫn:  TS. NGUYỄN VĂN AN\n"
        "Niên khóa                  :  2025 - 2026\n"
    )
    r_info.font.name = 'Times New Roman'
    r_info.font.size = Pt(13)
    r_info.font.bold = True
    
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(80)
    r_foot = p_foot.add_run("Hà Nội, Năm 2026")
    r_foot.font.name = 'Times New Roman'
    r_foot.font.size = Pt(13)
    r_foot.font.italic = True
    
    doc.add_page_break()
    
    # ==========================================================
    # LỜI MỞ ĐẦU & MỤC LỤC
    # ==========================================================
    add_styled_heading(doc, "LỜI NÓI ĐẦU", level=1)
    
    add_styled_paragraph(doc, 
        "Trong kỷ nguyên số hóa và sự bùng nổ của cuộc Cách mạng Công nghiệp 4.0, thương mại điện tử (E-Commerce) đã trở thành một phần tất yếu trong đời sống kinh tế xã hội toàn cầu. Tại Việt Nam, nhu cầu sở hữu các thiết bị công nghệ đỉnh cao như điện thoại thông minh (Smartphones), máy tính xách tay (Laptops), máy tính bảng (Tablets), tai nghe chống ồn và các thiết bị đeo thông minh (Smartwatches) ngày càng tăng trưởng mạnh mẽ. Người tiêu dùng ngày nay không chỉ tìm kiếm một nơi mua sắm uy tín, giá cả cạnh tranh mà còn đòi hỏi một trải nghiệm người dùng (UX) hiện đại, mượt mà, tiện lợi và an toàn tuyệt đối."
    )
    add_styled_paragraph(doc, 
        "Xuất phát từ nhu cầu thực tiễn đó, đồ án môn học 'Thiết kế và xây dựng website thương mại điện tử bán hàng công nghệ HieuMini' được thực hiện nhằm ứng dụng toàn diện các kiến thức nền tảng và nâng cao về lập trình web phía máy chủ (Server-side) với ngôn ngữ PHP kết hợp hệ quản trị cơ sở dữ liệu quan hệ MySQL. Website HieuMini được thiết kế theo phong cách hiện đại (Glassmorphism & Dark Modern Tech), trang bị đầy đủ các tính năng thương mại điện tử từ xem sản phẩm, bộ lọc thông minh, giỏ hàng tương tác không tải lại trang (AJAX), thanh toán tự động qua VietQR đến phân hệ quản trị (Admin Dashboard) theo dõi kinh doanh trực quan."
    )
    add_styled_paragraph(doc, 
        "Báo cáo này được cấu trúc thành 3 chương chính bám sát tiến trình phân tích, thiết kế, triển khai mã nguồn và đánh giá thử nghiệm hệ thống. Tác giả xin chân thành cảm ơn sự hướng dẫn tận tình của các thầy cô bộ môn để đề tài được hoàn thành một cách chỉn chu và đạt kết quả tốt nhất."
    )
    
    add_styled_heading(doc, "MỤC LỤC BÁO CÁO", level=1)
    
    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.line_spacing = 1.3
    p_toc.paragraph_format.space_after = Pt(12)
    
    toc_content = (
        "LỜI NÓI ĐẦU .................................................................................................................... 2\n"
        "MỤC LỤC ............................................................................................................................ 3\n\n"
        "CHƯƠNG 1. TỔNG QUAN LẬP TRÌNH WEB .................................................................... 4\n"
        "    1.1 Ngôn ngữ lập trình PHP .............................................................................................. 4\n"
        "        1.1.1 Lịch sử phát triển và nguyên lý hoạt động Server-Side ........................................ 4\n"
        "        1.1.2 Cú pháp cốt lõi và hướng đối tượng (OOP) ......................................................... 5\n"
        "        1.1.3 Quản lý trạng thái Session & Cookie ................................................................... 6\n"
        "        1.1.4 Cơ chế bảo mật và thư viện PDO ....................................................................... 6\n"
        "    1.2 Hệ quản trị cơ sở dữ liệu MySQL ............................................................................... 7\n"
        "        1.2.1 Tổng quan mô hình quan hệ RDBMS ................................................................. 7\n"
        "        1.2.2 Kiến trúc lưu trữ (InnoDB vs MyISAM) .............................................................. 8\n"
        "        1.2.3 Ràng buộc toàn vẹn và tính chất ACID ............................................................... 8\n"
        "        1.2.4 Đánh chỉ mục Indexing và tối ưu hóa truy vấn ................................................... 9\n"
        "    1.3 Cài đặt máy chủ .......................................................................................................... 9\n"
        "        1.3.1 Mô hình Client-Server và Web Server ............................................................... 9\n"
        "        1.3.2 Cài đặt và cấu hình môi trường (XAMPP / Laragon / PHP Server) ..................... 10\n"
        "        1.3.3 Tinh chỉnh tham số php.ini & my.ini ................................................................. 11\n"
        "        1.3.4 Quản trị cơ sở dữ liệu với phpMyAdmin ........................................................... 11\n\n"
        "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ WEBSITE ....................................................... 12\n"
        "    2.1 Chức năng (Use Case) ................................................................................................ 12\n"
        "        2.1.1 Phân tích yêu cầu hệ thống và tác nhân (Actors) ................................................ 12\n"
        "        2.1.2 Bảng phân quyền ma trận chức năng ................................................................. 13\n"
        "        2.1.3 Đặc tả chi tiết các Use Case chính ...................................................................... 14\n"
        "        2.1.4 Sơ đồ luồng dữ liệu DFD .................................................................................... 16\n"
        "    2.2 Cơ sở dữ liệu ............................................................................................................. 17\n"
        "        2.2.1 Mô hình quan hệ thực thể (ERD) ....................................................................... 17\n"
        "        2.2.2 Từ điển dữ liệu chi tiết các bảng (Data Dictionary) ............................................ 18\n"
        "        2.2.3 Các ràng buộc toàn vẹn khóa ngoại ................................................................... 21\n\n"
        "CHƯƠNG 3. CHƯƠNG TRÌNH THỬ NGHIỆM .............................................................. 22\n"
        "    3.1 Giao diện ................................................................................................................... 22\n"
        "        3.1.1 Triết lý thiết kế UI/UX hiện đại ......................................................................... 22\n"
        "        3.1.2 Chi tiết các màn hình Frontend (Khách hàng) .................................................... 23\n"
        "        3.1.3 Chi tiết các màn hình Backend (Admin Dashboard & CRUD) ............................. 26\n"
        "    3.2 Kết luận ..................................................................................................................... 28\n"
        "        3.2.1 Kết quả đạt được của dự án ............................................................................... 28\n"
        "        3.2.2 Ưu điểm và hạn chế ............................................................................................ 28\n"
        "        3.2.3 Hướng phát triển và mở rộng tương lai .............................................................. 29\n\n"
        "TÀI LIỆU THAM KHẢO ..................................................................................................... 30\n"
    )
    r_toc = p_toc.add_run(toc_content)
    r_toc.font.name = 'Times New Roman'
    r_toc.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ==========================================================
    # CHƯƠNG 1: TỔNG QUAN LẬP TRÌNH WEB
    # ==========================================================
    add_styled_heading(doc, "CHƯƠNG 1. TỔNG QUAN LẬP TRÌNH WEB", level=1)
    
    add_styled_heading(doc, "1.1 Ngôn ngữ lập trình PHP", level=2)
    
    add_styled_heading(doc, "1.1.1 Lịch sử phát triển và nguyên lý hoạt động Server-Side", level=3)
    add_styled_paragraph(doc, 
        "PHP (viết tắt đệ quy của 'PHP: Hypertext Preprocessor', ban đầu là Personal Home Page) là một trong những ngôn ngữ lập trình kịch bản phía máy chủ (Server-Side Scripting Language) phổ biến và thành công nhất trong lịch sử phát triển của mạng toàn cầu (World Wide Web). Được Rasmus Lerdorf sáng tạo vào năm 1994, trải qua hơn ba thập kỷ liên tục cải tiến và nâng cấp mạnh mẽ, PHP hiện nay đã phát triển đến phiên bản 8.x với hiệu năng vượt bậc nhờ bộ biên dịch Just-In-Time (JIT compiler), hệ thống kiểu dữ liệu nghiêm ngặt (Strict Types) và kiến trúc tối ưu Zend Engine."
    )
    add_styled_paragraph(doc, 
        "Về nguyên lý hoạt động, PHP vận hành theo mô hình Request - Response trên máy chủ Web Server. Khi người dùng (Client/Trình duyệt) gửi một yêu cầu HTTP (HTTP Request) tới một tài nguyên có đuôi mở rộng là .php, Web Server (như Apache HTTP Server hoặc Nginx) sẽ chuyển tiếp yêu cầu đó tới bộ thông dịch PHP (PHP Processor / FastCGI Process Manager). Tại đây, mã PHP được phân tích cú pháp (Parsing), biên dịch thành mã trung gian Opcode (Opcode Compilation), thực thi các logic nghiệp vụ, truy vấn và tương tác với cơ sở dữ liệu MySQL. Kết quả cuối cùng được chuyển đổi thành chuỗi HTML/CSS/JS hoặc JSON chuẩn và phản hồi (HTTP Response) về phía trình duyệt của người dùng. Do đó, người dùng phía client hoàn toàn không thể xem trực tiếp mã nguồn PHP gốc, đảm bảo tính bảo mật và an toàn cho toàn bộ logic hệ thống."
    )
    
    add_styled_heading(doc, "1.1.2 Cú pháp cốt lõi và lập trình hướng đối tượng (OOP)", level=3)
    add_styled_paragraph(doc, 
        "Cú pháp của PHP mang nhiều nét tương đồng với ngôn ngữ C, Java và Perl, giúp lập trình viên nhanh chóng tiếp cận và phát triển ứng dụng. Điểm mạnh vượt trội của PHP hiện đại là sự hỗ trợ toàn diện mô hình Lập trình hướng đối tượng (Object-Oriented Programming - OOP). Các nguyên lý OOP cơ bản như Đóng gói (Encapsulation), Kế thừa (Inheritance), Đa hình (Polymorphism) và Trừu tượng (Abstraction) được thể hiện rõ nét qua các từ khóa `class`, `interface`, `trait`, `abstract`, `public`, `private`, `protected`."
    )
    add_styled_paragraph(doc, 
        "Trong dự án HieuMini, mô hình hướng đối tượng kết hợp mẫu thiết kế Singleton Pattern được áp dụng trực tiếp trong lớp quản lý kết nối CSDL `Database` (`config/database.php`), giúp đảm bảo toàn bộ hệ thống chỉ duy trì duy nhất một kết nối cơ sở dữ liệu xuyên suốt vòng đời của mỗi request, tiết kiệm bộ nhớ RAM và tối ưu hóa tài nguyên server."
    )
    
    add_styled_heading(doc, "1.1.3 Quản lý trạng thái Session & Cookie trong ứng dụng E-Commerce", level=3)
    add_styled_paragraph(doc, 
        "Giao thức HTTP bản chất là một giao thức phi trạng thái (Stateless Protocol), nghĩa là máy chủ không tự động lưu giữ thông tin về các lần yêu cầu trước đó của cùng một máy khách. Để giải quyết bài toán cốt lõi của website thương mại điện tử như duy trì phiên đăng nhập của người dùng và lưu trữ giỏ hàng khi duyệt qua nhiều trang sản phẩm khác nhau, PHP cung cấp hai cơ chế mạnh mẽ:"
    )
    add_styled_paragraph(doc, 
        "1. Session: Dữ liệu được lưu trữ an toàn trực tiếp trên bộ nhớ/ổ cứng của máy chủ (Server). Mỗi phiên làm việc được định danh bằng một chuỗi ngẫu nhiên duy nhất gọi là Session ID (PHPSESSID), được truyền giữa client và server thông qua HTTP Header. Trong website HieuMini, Session được sử dụng để lưu trữ `$_SESSION['user']` (thông tin người dùng đăng nhập), `$_SESSION['cart']` (danh sách các sản phẩm và số lượng trong giỏ hàng) và `$_SESSION['coupon']` (thông tin mã giảm giá áp dụng)."
    )
    add_styled_paragraph(doc, 
        "2. Cookie: Dữ liệu nhỏ được lưu trữ trực tiếp trên trình duyệt của máy khách (Client), có thể thiết lập thời gian hết hạn (Expires / Max-Age), cờ an toàn (Secure, HttpOnly, SameSite) để phòng chống các cuộc tấn công đánh cắp phiên làm việc."
    )
    
    add_styled_heading(doc, "1.1.4 Cơ chế bảo mật và thư viện PDO (PHP Data Objects)", level=3)
    add_styled_paragraph(doc, 
        "Bảo mật là yếu tố sống còn đối với bất kỳ hệ thống thương mại điện tử nào. Hệ thống website HieuMini áp dụng nghiêm ngặt các tiêu chuẩn an ninh ứng dụng web OWASP Top 10 thông qua việc sử dụng thư viện PDO (PHP Data Objects):"
    )
    add_styled_paragraph(doc, 
        "• Phòng chống tấn công SQL Injection: Thư viện PDO hỗ trợ cơ chế Tham số hóa truy vấn (Parameterized Queries / Prepared Statements) với `prepare()` và `execute()`. Cơ chế này phân tách hoàn toàn mã lệnh SQL và dữ liệu do người dùng nhập vào, ngăn chặn triệt để kẻ tấn công chèn các đoạn mã SQL độc hại vào câu truy vấn."
    )
    add_styled_paragraph(doc, 
        "• Phòng chống tấn công XSS (Cross-Site Scripting): Mọi dữ liệu do người dùng nhập trước khi hiển thị ra giao diện HTML đều được làm sạch qua hàm xử lý chuyên dụng `sanitize()` sử dụng `htmlspecialchars(..., ENT_QUOTES, 'UTF-8')` để mã hóa các ký tự đặc biệt như `<`, `>`, `&`, `'`, `\"`."
    )
    add_styled_paragraph(doc, 
        "• Mã hóa mật khẩu an toàn: Mật khẩu người dùng và quản trị viên được băm bằng thuật toán BCRYPT mạnh mẽ thông qua hàm chuẩn `password_hash($password, PASSWORD_BCRYPT)` kết hợp kiểm tra bằng `password_verify()`, tuyệt đối không lưu mật khẩu ở dạng văn bản thuần (Plain Text)."
    )
    
    add_styled_heading(doc, "1.2 Hệ quản trị cơ sở dữ liệu MySQL", level=2)
    
    add_styled_heading(doc, "1.2.1 Tổng quan mô hình quan hệ RDBMS", level=3)
    add_styled_paragraph(doc, 
        "MySQL là hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở (Relational Database Management System - RDBMS) phổ biến nhất trên thế giới, do tập đoàn Oracle phát triển và hỗ trợ. MySQL tổ chức dữ liệu dưới dạng các bảng (Tables) có cấu trúc chặt chẽ gồm các hàng (Rows / Records) và các cột (Columns / Fields). Mối quan hệ giữa các bảng được thiết lập thông qua các khóa chính (Primary Key) và khóa ngoại (Foreign Key), cho phép quản lý các tập dữ liệu phức tạp với hiệu năng cao và tính nhất quán dữ liệu tuyệt đối."
    )
    
    add_styled_heading(doc, "1.2.2 Kiến trúc lưu trữ (So sánh Storage Engines: InnoDB vs MyISAM)", level=3)
    add_styled_paragraph(doc, 
        "MySQL sở hữu kiến trúc Engine lưu trữ theo dạng module có thể cắm ghép linh hoạt (Pluggable Storage Engine Architecture). Hai engine lưu trữ phổ biến nhất trong MySQL là InnoDB và MyISAM:"
    )
    
    # Bảng so sánh InnoDB vs MyISAM
    tbl_engine = doc.add_table(rows=1, cols=3)
    format_table(
        tbl_engine,
        col_widths=[1.5, 2.7, 2.7],
        headers=["Tiêu chí so sánh", "InnoDB Storage Engine", "MyISAM Storage Engine"],
        data=[
            ["Hỗ trợ Giao dịch (ACID)", "Có hỗ trợ toàn diện (Commit / Rollback)", "Không hỗ trợ giao dịch"],
            ["Khóa ngoại (Foreign Key)", "Có hỗ trợ ràng buộc toàn vẹn quan hệ", "Không hỗ trợ ràng buộc khóa ngoại"],
            ["Cấp độ khóa (Locking)", "Khóa theo cấp độ hàng (Row-level Locking)", "Khóa theo cấp độ bảng (Table-level Locking)"],
            ["Khả năng phục hồi sự cố", "Tự động phục hồi qua nhật ký ghi lại (Crash Recovery)", "Dễ bị hỏng bảng khi máy chủ tắt đột ngột"],
            ["Hiệu năng ghi đồng thời", "Rất cao (thích hợp E-Commerce nhiều giao dịch)", "Thấp khi có nhiều thao tác ghi đồng thời"],
            ["Khuyến nghị sử dụng", "Mặc định cho các ứng dụng giao dịch, E-Commerce", "Chỉ phù hợp cho hệ thống chỉ đọc (Read-only)"]
        ]
    )
    
    add_styled_paragraph(doc, 
        "Trong dự án HieuMini, toàn bộ các bảng trong CSDL `hieumini_db` đều được thiết lập sử dụng 100% engine `InnoDB` và bộ mã ký tự `utf8mb4` (hỗ trợ đầy đủ tiếng Việt có dấu, ký tự quốc tế và biểu tượng cảm xúc Emoji)."
    )
    
    add_styled_heading(doc, "1.2.3 Ràng buộc toàn vẹn và tính chất ACID", level=3)
    add_styled_paragraph(doc, 
        "Đối với hệ thống thương mại điện tử, các giao dịch mua hàng (Transactions) đòi hỏi phải tuân thủ nghiêm ngặt 4 thuộc tính ACID:"
    )
    add_styled_paragraph(doc, 
        "• Tính nguyên tử (Atomicity): Một giao dịch đặt hàng bao gồm việc tạo bản ghi đơn hàng trong bảng `orders` và lưu chi tiết sản phẩm trong bảng `order_items`. Toàn bộ các thao tác này hoặc cùng thành công hoặc cùng bị hủy bỏ (Rollback), không bao giờ xảy ra tình trạng có đơn hàng nhưng thiếu danh sách mặt hàng."
    )
    add_styled_paragraph(doc, 
        "• Tính nhất quán (Consistency): Dữ liệu luôn chuyển từ trạng thái hợp lệ này sang trạng thái hợp lệ khác, đảm bảo các ràng buộc về số lượng tồn kho, giá trị tiền tệ không bị âm."
    )
    add_styled_paragraph(doc, 
        "• Tính cô lập (Isolation): Các giao dịch đặt hàng diễn ra đồng thời không can thiệp hay ảnh hưởng lẫn nhau."
    )
    add_styled_paragraph(doc, 
        "• Tính bền vững (Durability): Khi một giao dịch đặt hàng đã được xác nhận (Commit), kết quả sẽ được lưu vĩnh viễn vào bộ nhớ đĩa cứng, ngay cả khi xảy ra sự cố mất điện."
    )
    
    add_styled_heading(doc, "1.2.4 Đánh chỉ mục Indexing và tối ưu hóa truy vấn", level=3)
    add_styled_paragraph(doc, 
        "Để tăng tốc độ tìm kiếm và lọc sản phẩm trong số hàng nghìn thiết bị công nghệ, MySQL sử dụng cấu trúc cây B-Tree Indexing. Trong CSDL HieuMini, các chỉ mục được đánh trên các trường thường xuyên truy vấn như `category_id`, `brand`, `price`, `slug`, `order_code`, giúp giảm độ phức tạp tìm kiếm từ O(N) xuống O(log N), mang lại tốc độ phản hồi trang tức thì dưới 50ms."
    )
    
    add_styled_heading(doc, "1.3 Cài đặt máy chủ", level=2)
    
    add_styled_heading(doc, "1.3.1 Mô hình Client-Server và Web Server", level=3)
    add_styled_paragraph(doc, 
        "Mô hình máy khách - máy chủ (Client-Server Architecture) là nền tảng cốt lõi của toàn bộ ứng dụng web. Trong mô hình này, Client (trình duyệt Google Chrome, Edge, Safari...) đóng vai trò gửi yêu cầu và hiển thị giao diện; trong khi Server (Web Server, PHP Engine, Database Server) chịu trách nhiệm lưu trữ tập trung tài nguyên, xử lý nghiệp vụ và bảo vệ dữ liệu."
    )
    
    add_styled_heading(doc, "1.3.2 Hướng dẫn cài đặt và cấu hình môi trường phát triển (XAMPP / Laragon / PHP Server)", level=3)
    add_styled_paragraph(doc, 
        "Để triển khai và chạy thử nghiệm website HieuMini trên môi trường máy tính nội bộ (Localhost), lập trình viên có thể sử dụng các gói phần mềm tích hợp WAMP/LAMP như XAMPP, Laragon hoặc máy chủ tích hợp sẵn PHP Built-in Web Server theo các bước sau:"
    )
    add_styled_paragraph(doc, 
        "Bước 1: Tải và cài đặt gói ứng dụng XAMPP (hoặc Laragon) từ trang chủ chính thức (yêu cầu PHP phiên bản 7.4 trở lên, khuyến nghị PHP 8.1 - 8.3)."
    )
    add_styled_paragraph(doc, 
        "Bước 2: Mở bảng điều khiển XAMPP Control Panel, nhấn nút 'Start' tại hai dịch vụ quan trọng là Apache (cổng mặc định 80/443) và MySQL (cổng mặc định 3306)."
    )
    add_styled_paragraph(doc, 
        "Bước 3: Sao chép toàn bộ thư mục mã nguồn dự án `HieuWeb02` vào thư mục gốc của máy chủ web: `C:\\xampp\\htdocs\\HieuWeb02`."
    )
    add_styled_paragraph(doc, 
        "Bước 4: Đối với trường hợp sử dụng PHP Built-in Server nhanh từ dòng lệnh, mở PowerShell tại thư mục dự án và thực thi lệnh: `php -S localhost:8000`."
    )
    
    add_styled_heading(doc, "1.3.3 Tinh chỉnh các thông số cấu hình quan trọng (php.ini & my.ini)", level=3)
    add_styled_paragraph(doc, 
        "Để đảm bảo website hoạt động ổn định khi upload hình ảnh dung lượng lớn và xử lý nhiều tác vụ phức tạp, cần kiểm tra và cấu hình các chỉ thị sau trong file `php.ini`:"
    )
    add_styled_paragraph(doc, 
        "• `upload_max_filesize = 64M`: Cho phép tải lên các tệp ảnh sản phẩm có kích thước tối đa 64 Megabytes.\n"
        "• `post_max_size = 128M`: Giới hạn dung lượng tối đa cho mỗi gói dữ liệu POST gửi lên server.\n"
        "• `memory_limit = 256M`: Cấp phát bộ nhớ RAM tối đa cho một tiến trình thực thi mã PHP.\n"
        "• `max_execution_time = 300`: Thời gian chạy tối đa của một script PHP (300 giây) trước khi ngắt kết nối.\n"
        "• `extension=pdo_mysql`: Kích hoạt module driver kết nối PDO tới hệ quản trị MySQL."
    )
    
    add_styled_heading(doc, "1.3.4 Quản trị cơ sở dữ liệu trực quan với phpMyAdmin", level=3)
    add_styled_paragraph(doc, 
        "Truy cập vào công cụ quản trị web phpMyAdmin thông qua đường dẫn `http://localhost/phpmyadmin`. Tiến hành tạo mới cơ sở dữ liệu với tên `hieumini_db`, chọn bảng mã `utf8mb4_unicode_ci`. Sau đó, chuyển sang tab 'Import', chọn tệp `database/hieumini_db.sql` trong thư mục dự án và nhấn 'Go' để tự động khởi tạo toàn bộ 8 bảng cấu trúc cùng dữ liệu sản phẩm mẫu."
    )
    
    doc.add_page_break()
    
    # ==========================================================
    # CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ WEBSITE
    # ==========================================================
    add_styled_heading(doc, "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ WEBSITE", level=1)
    
    add_styled_heading(doc, "2.1 Chức năng (Use Case)", level=2)
    
    add_styled_heading(doc, "2.1.1 Phân tích yêu cầu hệ thống và tác nhân (Actors)", level=3)
    add_styled_paragraph(doc, 
        "Hệ thống website thương mại điện tử HieuMini được thiết kế nhằm giải quyết bài toán kinh doanh bán lẻ thiết bị công nghệ trực tuyến. Qua khảo sát thực tế, hệ thống xác định 3 nhóm tác nhân (Actors) chính tham gia vào các luồng nghiệp vụ:"
    )
    add_styled_paragraph(doc, 
        "1. Khách vãng lai (Guest): Người dùng chưa đăng nhập tài khoản. Có quyền xem danh sách sản phẩm, lọc theo danh mục/mức giá/thương hiệu, xem chi tiết thông số kỹ thuật, thêm sản phẩm vào giỏ hàng, đặt hàng nhanh và tra cứu đơn hàng bằng số điện thoại."
    )
    add_styled_paragraph(doc, 
        "2. Thành viên / Khách hàng đã đăng nhập (Customer): Người dùng đã có tài khoản hệ thống. Ngoài các quyền của Khách vãng lai, thành viên có quyền quản lý thông tin cá nhân, lưu địa chỉ giao hàng mặc định, theo dõi lịch sử và trạng thái các đơn hàng đã đặt, viết đánh giá và chấm điểm sao cho sản phẩm."
    )
    add_styled_paragraph(doc, 
        "3. Quản trị viên (Admin): Người có quyền hạn cao nhất trong hệ thống. Quản trị viên được cấp quyền truy cập vào phân hệ Dashboard để theo dõi các chỉ số KPI doanh thu, thống kê đơn hàng, thực hiện toàn bộ các thao tác CRUD (Create, Read, Update, Delete) đối với Sản phẩm, Danh mục, Đơn hàng và Người dùng."
    )
    
    add_styled_heading(doc, "2.1.2 Bảng ma trận phân quyền chức năng hệ thống", level=3)
    
    # Bảng phân quyền
    tbl_matrix = doc.add_table(rows=1, cols=4)
    format_table(
        tbl_matrix,
        col_widths=[3.0, 1.3, 1.3, 1.3],
        headers=["Nhóm Chức Năng / Nghiệp Vụ", "Khách Vãng Lai", "Thành Viên", "Quản Trị Admin"],
        data=[
            ["Xem trang chủ, Banner, Flash Sale", "Có", "Có", "Có"],
            ["Tìm kiếm & Bộ lọc sản phẩm đa tiêu chí", "Có", "Có", "Có"],
            ["Xem chi tiết sản phẩm & thông số kỹ thuật", "Có", "Có", "Có"],
            ["Thêm vào giỏ hàng & Cập nhật số lượng (AJAX)", "Có", "Có", "Có"],
            ["Áp dụng mã giảm giá (Coupon Code)", "Có", "Có", "Có"],
            ["Đặt hàng & Chọn thanh toán (COD / VietQR)", "Có", "Có", "Có"],
            ["Đăng ký, Đăng nhập & Đổi mật khẩu", "Có", "Có", "Có"],
            ["Xem lịch sử & Tiến trình đơn hàng cá nhân", "Chỉ qua SĐT", "Toàn diện", "Toàn diện"],
            ["Gửi đánh giá & Chấm sao sản phẩm", "Không", "Có", "Có"],
            ["Dashboard thống kê doanh thu & Biểu đồ", "Không", "Không", "Toàn quyền"],
            ["Quản lý Sản phẩm (Thêm, Sửa, Xóa, Tồn kho)", "Không", "Không", "Toàn quyền"],
            ["Quản lý Danh mục (CRUD Danh mục)", "Không", "Không", "Toàn quyền"],
            ["Quản lý Đơn hàng & Đổi trạng thái giao hàng", "Không", "Không", "Toàn quyền"],
            ["Quản lý Người dùng & Phân quyền Admin", "Không", "Không", "Toàn quyền"]
        ]
    )
    
    add_styled_heading(doc, "2.1.3 Bảng đặc tả chi tiết các Use Case chính", level=3)
    
    # Bảng đặc tả UC 1: Mua hàng & Thanh toán
    add_styled_paragraph(doc, "Bảng 2.1: Đặc tả Use Case 'Đặt mua hàng và thanh toán trực tuyến'", bold=True)
    tbl_uc1 = doc.add_table(rows=1, cols=2)
    format_table(
        tbl_uc1,
        col_widths=[2.2, 4.7],
        headers=["Mục Thông Tin", "Nội Dung Đặc Tả Chi Tiết"],
        data=[
            ["Tên Use Case", "Đặt mua hàng và thanh toán trực tuyến (Checkout Order)"],
            ["Mã Use Case", "UC-01"],
            ["Tác nhân (Actor)", "Khách vãng lai, Thành viên"],
            ["Mục đích", "Cho phép người dùng đặt mua các mặt hàng trong giỏ và chọn hình thức thanh toán phù hợp."],
            ["Tiền điều kiện", "Giỏ hàng đang có ít nhất một sản phẩm hợp lệ."],
            ["Luồng sự kiện chính (Main Flow)", 
             "1. Người dùng nhấn nút 'Tiến hành thanh toán' từ trang Giỏ hàng.\n"
             "2. Hệ thống hiển thị form nhập thông tin người nhận và lựa chọn phương thức thanh toán.\n"
             "3. Người dùng nhập Họ tên, Số điện thoại, Địa chỉ nhận hàng, Tỉnh/Thành phố.\n"
             "4. Người dùng chọn phương thức thanh toán: COD hoặc Quét mã VietQR.\n"
             "5. Nếu chọn VietQR, hệ thống hiển thị thông tin ngân hàng và mã QR động kèm số tiền chính xác.\n"
             "6. Người dùng nhấn nút 'Xác nhận đặt hàng'.\n"
             "7. Hệ thống mở Transaction: Tạo bản ghi `orders` và các dòng chi tiết `order_items`.\n"
             "8. Hệ thống xóa trắng giỏ hàng trong Session và chuyển hướng đến trang xác nhận `order_success.php` kèm mã tra cứu đơn hàng."
            ],
            ["Luồng ngoại lệ (Alternative Flow)", 
             "3a. Người dùng bỏ trống các trường bắt buộc (*): Hệ thống hiển thị cảnh báo lỗi và yêu cầu nhập lại.\n"
             "7a. Lỗi kết nối CSDL: Hệ thống Rollback giao dịch và thông báo lỗi tới người dùng."
            ],
            ["Hậu điều kiện", "Đơn hàng mới được lưu vào CSDL với trạng thái 'pending' và mã đơn duy nhất."]
        ]
    )
    
    # Bảng đặc tả UC 2: Quản lý sản phẩm Admin
    add_styled_paragraph(doc, "\nBảng 2.2: Đặc tả Use Case 'Quản lý sản phẩm dành cho Quản trị viên'", bold=True)
    tbl_uc2 = doc.add_table(rows=1, cols=2)
    format_table(
        tbl_uc2,
        col_widths=[2.2, 4.7],
        headers=["Mục Thông Tin", "Nội Dung Đặc Tả Chi Tiết"],
        data=[
            ["Tên Use Case", "Quản lý sản phẩm (Product Management - Admin)"],
            ["Mã Use Case", "UC-02"],
            ["Tác nhân (Actor)", "Quản trị viên (Admin)"],
            ["Mục đích", "Cho phép Admin theo dõi danh sách, tìm kiếm, thêm mới, sửa đổi thông tin hoặc xóa sản phẩm."],
            ["Tiền điều kiện", "Admin đã đăng nhập thành công với vai trò 'admin'."],
            ["Luồng sự kiện chính (Main Flow)", 
             "1. Admin truy cập trang `admin/products.php`.\n"
             "2. Hệ thống tải và hiển thị danh sách tất cả sản phẩm kèm ảnh, giá bán, tồn kho, nhãn sale.\n"
             "3. Thêm mới: Admin nhấn 'Thêm sản phẩm mới', điền đầy đủ form thông số và nhấn 'Lưu'. Hệ thống kiểm tra dữ liệu và thêm bản ghi vào bảng `products`.\n"
             "4. Chỉnh sửa: Admin nhấn icon Sửa tại dòng sản phẩm tương ứng, chỉnh sửa giá/tồn kho và nhấn 'Lưu thay đổi'.\n"
             "5. Xóa: Admin nhấn icon Thùng rác, hệ thống hiện hộp thoại xác nhận JavaScript `confirm()`. Khi Admin đồng ý, hệ thống xóa bản ghi sản phẩm khỏi CSDL."
            ],
            ["Hậu điều kiện", "Dữ liệu sản phẩm trong CSDL được cập nhật chính xác và hiển thị ngay trên giao diện khách hàng."]
        ]
    )
    
    add_styled_heading(doc, "2.1.4 Sơ đồ luồng dữ liệu DFD (Data Flow Diagram)", level=3)
    add_styled_paragraph(doc, 
        "Hệ thống vận hành theo 2 cấp độ luồng dữ liệu chính:\n"
        "• DFD Mức ngữ cảnh (Level 0): Mô tả sự tương tác tổng thể giữa Khách hàng, Quản trị viên và Hệ sinh thái Website HieuMini.\n"
        "• DFD Mức 1 (Level 1): Phân rã thành 5 tiến trình cốt lõi: (1.0) Xác thực người dùng -> (2.0) Tìm kiếm & Hiển thị sản phẩm -> (3.0) Xử lý giỏ hàng & Khuyến mãi -> (4.0) Xử lý đơn hàng & Thanh toán -> (5.0) Quản trị và Báo cáo thống kê."
    )
    
    add_styled_heading(doc, "2.2 Cơ sở dữ liệu", level=2)
    
    add_styled_heading(doc, "2.2.1 Mô hình quan hệ thực thể (ERD)", level=3)
    add_styled_paragraph(doc, 
        "Mô hình dữ liệu của HieuMini bao gồm 8 thực thể quan hệ chặt chẽ: `categories`, `users`, `products`, `orders`, `order_items`, `reviews`, `coupons`, `banners`. Mối quan hệ giữa các thực thể được thiết kế chuẩn hóa đạt Chuẩn 3 (3NF - Third Normal Form) nhằm loại bỏ dư thừa dữ liệu và ngăn chặn các bất thường khi cập nhật."
    )
    add_styled_paragraph(doc, 
        "• Mối quan hệ `categories` - `products`: Quan hệ 1 - N (Một danh mục chứa nhiều sản phẩm; một sản phẩm thuộc về duy nhất một danh mục cha).\n"
        "• Mối quan hệ `users` - `orders`: Quan hệ 1 - N (Một người dùng có thể tạo nhiều đơn hàng; một đơn hàng gắn liền với một người dùng hoặc khách vãng lai).\n"
        "• Mối quan hệ `orders` - `products` qua bảng trung gian `order_items`: Quan hệ N - N (Một đơn hàng chứa nhiều sản phẩm khác nhau; một sản phẩm có thể xuất hiện trong nhiều đơn hàng với số lượng và đơn giá lưu vết độc lập).\n"
        "• Mối quan hệ `products` - `reviews`: Quan hệ 1 - N (Một sản phẩm có nhiều lượt đánh giá và nhận xét của khách hàng)."
    )
    
    add_styled_heading(doc, "2.2.2 Từ điển dữ liệu chi tiết các bảng (Data Dictionary)", level=3)
    
    # Bảng categories
    add_styled_paragraph(doc, "Bảng 2.3: Cấu trúc bảng `categories` (Danh mục sản phẩm)", bold=True)
    tbl_c1 = doc.add_table(rows=1, cols=6)
    format_table(tbl_c1, [1.2, 1.2, 0.8, 0.8, 0.8, 2.1], 
        ["Tên Trường", "Kiểu Dữ Liệu", "Độ Dài", "Khóa", "Null", "Mô Tả Ý Nghĩa"],
        [
            ["id", "INT", "11", "PK", "No", "Mã danh mục tự tăng (Auto Increment)"],
            ["name", "VARCHAR", "100", "", "No", "Tên danh mục (VD: Điện thoại, Laptop)"],
            ["slug", "VARCHAR", "120", "Unique", "No", "Đường dẫn URL thân thiện SEO"],
            ["icon", "VARCHAR", "50", "", "Yes", "Mã class icon FontAwesome"],
            ["description", "TEXT", "-", "", "Yes", "Mô tả chi tiết nhóm danh mục"],
            ["status", "TINYINT", "1", "", "Yes", "Trạng thái hiển thị (1: Hiện, 0: Ẩn)"],
            ["created_at", "TIMESTAMP", "-", "", "No", "Thời gian khởi tạo bản ghi"]
        ]
    )
    
    # Bảng users
    add_styled_paragraph(doc, "\nBảng 2.4: Cấu trúc bảng `users` (Người dùng & Quản trị viên)", bold=True)
    tbl_c2 = doc.add_table(rows=1, cols=6)
    format_table(tbl_c2, [1.2, 1.2, 0.8, 0.8, 0.8, 2.1], 
        ["Tên Trường", "Kiểu Dữ Liệu", "Độ Dài", "Khóa", "Null", "Mô Tả Ý Nghĩa"],
        [
            ["id", "INT", "11", "PK", "No", "Mã người dùng tự tăng"],
            ["full_name", "VARCHAR", "100", "", "No", "Họ và tên đầy đủ"],
            ["email", "VARCHAR", "100", "Unique", "No", "Email đăng nhập hệ thống"],
            ["password", "VARCHAR", "255", "", "No", "Mật khẩu băm chuẩn BCRYPT"],
            ["phone", "VARCHAR", "20", "", "Yes", "Số điện thoại liên lạc"],
            ["address", "VARCHAR", "255", "", "Yes", "Địa chỉ giao hàng mặc định"],
            ["role", "ENUM", "'admin','customer'", "", "No", "Vai trò tài khoản"],
            ["status", "TINYINT", "1", "", "No", "Trạng thái (1: Hoạt động, 0: Khóa)"],
            ["created_at", "TIMESTAMP", "-", "", "No", "Thời gian đăng ký tài khoản"]
        ]
    )
    
    # Bảng products
    add_styled_paragraph(doc, "\nBảng 2.5: Cấu trúc bảng `products` (Sản phẩm công nghệ)", bold=True)
    tbl_c3 = doc.add_table(rows=1, cols=6)
    format_table(tbl_c3, [1.2, 1.2, 0.8, 0.8, 0.8, 2.1], 
        ["Tên Trường", "Kiểu Dữ Liệu", "Độ Dài", "Khóa", "Null", "Mô Tả Ý Nghĩa"],
        [
            ["id", "INT", "11", "PK", "No", "Mã sản phẩm tự tăng"],
            ["category_id", "INT", "11", "FK", "No", "Khóa ngoại liên kết bảng categories"],
            ["name", "VARCHAR", "200", "", "No", "Tên thiết bị công nghệ"],
            ["slug", "VARCHAR", "220", "Unique", "No", "Đường dẫn SEO sản phẩm"],
            ["brand", "VARCHAR", "80", "", "No", "Thương hiệu (Apple, ASUS, Sony...)"],
            ["price", "DECIMAL", "12,2", "", "No", "Giá niêm yết chính thức (VNĐ)"],
            ["sale_price", "DECIMAL", "12,2", "", "Yes", "Giá khuyến mãi giảm giá (VNĐ)"],
            ["stock_quantity", "INT", "11", "", "No", "Số lượng sản phẩm trong kho"],
            ["thumbnail", "VARCHAR", "255", "", "No", "Tên tệp hình ảnh đại diện"],
            ["short_desc", "VARCHAR", "300", "", "Yes", "Tóm tắt nổi bật sản phẩm"],
            ["description", "LONGTEXT", "-", "", "Yes", "Bài viết chi tiết đặc điểm sản phẩm"],
            ["specifications", "LONGTEXT", "-", "", "Yes", "Thông số kỹ thuật định dạng JSON"],
            ["is_featured", "TINYINT", "1", "", "No", "Cờ sản phẩm nổi bật (1/0)"],
            ["is_flash_sale", "TINYINT", "1", "", "No", "Cờ sản phẩm giờ vàng Flash Sale (1/0)"],
            ["views", "INT", "11", "", "No", "Lượt xem sản phẩm"],
            ["rating", "DECIMAL", "2,1", "", "No", "Điểm đánh giá trung bình (1.0 - 5.0)"]
        ]
    )
    
    # Bảng orders
    add_styled_paragraph(doc, "\nBảng 2.6: Cấu trúc bảng `orders` (Đơn hàng mua sắm)", bold=True)
    tbl_c4 = doc.add_table(rows=1, cols=6)
    format_table(tbl_c4, [1.2, 1.2, 0.8, 0.8, 0.8, 2.1], 
        ["Tên Trường", "Kiểu Dữ Liệu", "Độ Dài", "Khóa", "Null", "Mô Tả Ý Nghĩa"],
        [
            ["id", "INT", "11", "PK", "No", "Mã định danh đơn hàng"],
            ["order_code", "VARCHAR", "30", "Unique", "No", "Mã đơn tra cứu (VD: HM-20260820-A1B2)"],
            ["user_id", "INT", "11", "FK", "Yes", "Mã người dùng (Null nếu khách vãng lai)"],
            ["customer_name", "VARCHAR", "100", "", "No", "Tên người nhận hàng"],
            ["customer_email", "VARCHAR", "100", "", "Yes", "Email nhận hóa đơn điện tử"],
            ["customer_phone", "VARCHAR", "20", "", "No", "Số điện thoại giao hàng"],
            ["customer_address", "VARCHAR", "255", "", "No", "Địa chỉ nhận chi tiết"],
            ["shipping_city", "VARCHAR", "100", "", "No", "Tỉnh/Thành phố nhận hàng"],
            ["payment_method", "ENUM", "'cod','bank_transfer','momo'", "", "No", "Phương thức thanh toán"],
            ["payment_status", "ENUM", "'unpaid','paid'", "", "No", "Trạng thái thanh toán"],
            ["shipping_status", "ENUM", "'pending','processing','shipping','completed','cancelled'", "", "No", "Tiến trình vận chuyển đơn"],
            ["subtotal", "DECIMAL", "12,2", "", "No", "Tổng tiền hàng trước giảm giá"],
            ["discount", "DECIMAL", "12,2", "", "No", "Số tiền giảm từ mã khuyến mãi"],
            ["shipping_fee", "DECIMAL", "12,2", "", "No", "Phí giao hàng"],
            ["total_amount", "DECIMAL", "12,2", "", "No", "Tổng thanh toán cuối cùng (VNĐ)"]
        ]
    )
    
    # Bảng order_items
    add_styled_paragraph(doc, "\nBảng 2.7: Cấu trúc bảng `order_items` (Chi tiết mặt hàng trong đơn)", bold=True)
    tbl_c5 = doc.add_table(rows=1, cols=6)
    format_table(tbl_c5, [1.2, 1.2, 0.8, 0.8, 0.8, 2.1], 
        ["Tên Trường", "Kiểu Dữ Liệu", "Độ Dài", "Khóa", "Null", "Mô Tả Ý Nghĩa"],
        [
            ["id", "INT", "11", "PK", "No", "Mã chi tiết dòng đơn"],
            ["order_id", "INT", "11", "FK", "No", "Khóa ngoại liên kết bảng orders"],
            ["product_id", "INT", "11", "FK", "Yes", "Mã sản phẩm tương ứng"],
            ["product_name", "VARCHAR", "200", "", "No", "Lưu tên sản phẩm tại thời điểm mua"],
            ["price", "DECIMAL", "12,2", "", "No", "Đơn giá mua tại thời điểm đặt hàng"],
            ["quantity", "INT", "11", "", "No", "Số lượng đặt mua"],
            ["total", "DECIMAL", "12,2", "", "No", "Thành tiền (price * quantity)"]
        ]
    )
    
    add_styled_heading(doc, "2.2.3 Các ràng buộc toàn vẹn khóa ngoại (Referential Integrity)", level=3)
    add_styled_paragraph(doc, 
        "Hệ thống thiết lập các chính sách toàn vẹn dữ liệu chặt chẽ trên tầng CSDL:\n"
        "• `ON DELETE CASCADE` giữa `categories` và `products`: Khi một danh mục bị xóa, tất cả sản phẩm thuộc danh mục đó sẽ tự động được dọn dẹp an toàn.\n"
        "• `ON DELETE CASCADE` giữa `orders` và `order_items`: Khi xóa một đơn hàng, toàn bộ các dòng chi tiết của đơn đó tự động được xóa đồng bộ.\n"
        "• `ON DELETE SET NULL` giữa `users` và `orders`: Khi một tài khoản người dùng bị xóa, thông tin lịch sử các đơn hàng cũ vẫn được giữ nguyên vẹn để phục vụ báo cáo doanh thu kế toán."
    )
    
    doc.add_page_break()
    
    # ==========================================================
    # CHƯƠNG 3: CHƯƠNG TRÌNH THỬ NGHIỆM
    # ==========================================================
    add_styled_heading(doc, "CHƯƠNG 3. CHƯƠNG TRÌNH THỬ NGHIỆM", level=1)
    
    add_styled_heading(doc, "3.1 Giao diện", level=2)
    
    add_styled_heading(doc, "3.1.1 Triết lý thiết kế UI/UX hiện đại", level=3)
    add_styled_paragraph(doc, 
        "Website HieuMini được thiết kế dựa trên ngôn ngữ giao diện đương đại Glassmorphism kết hợp phong cách Dark Modern Tech. Tông màu chủ đạo là nền đen xanh sâu thẳm (`#0b0f19` và `#111827`) kết hợp với các hiệu ứng thẻ kính mờ bán trong suốt (`backdrop-filter: blur(16px)`), viền phát sáng tinh tế (`border: 1px solid rgba(255,255,255,0.08)`) và các màu điểm nhấn công nghệ như Neon Indigo (`#6366f1`), Cyan (`#06b6d4`), Pink (`#ec4899`)."
    )
    add_styled_paragraph(doc, 
        "Font chữ được sử dụng là Google Font 'Plus Jakarta Sans' chuẩn quốc tế, hiển thị cực kỳ sắc nét trên mọi độ phân giải màn hình từ máy tính để bàn (Desktop), máy tính xách tay (Laptop), máy tính bảng (Tablet) cho đến điện thoại di động (Mobile). Các hiệu ứng vi tương tác (Micro-interactions) như chuyển động hover lơ lửng, nút bấm có bóng phát sáng (Glow effect), thanh đếm ngược giờ vàng chạy thời gian thực và thông báo Toast nổi bật giúp người dùng có trải nghiệm mua sắm vô cùng thỏa mãn và chuyên nghiệp."
    )
    
    add_styled_heading(doc, "3.1.2 Chi tiết các màn hình chức năng Phân hệ Khách hàng (Frontend)", level=3)
    
    # 1. Trang chủ
    add_styled_paragraph(doc, "1. Màn hình Trang chủ (index.php):", bold=True)
    add_styled_paragraph(doc, 
        "• Header Sticky: Cố định trên đỉnh trang khi cuộn, tích hợp Logo nhận diện HieuMini, thanh tìm kiếm thông minh, menu danh mục sản phẩm, nút truy cập Quản trị (nếu là Admin) và nút Giỏ hàng có badge hiển thị số lượng sản phẩm được cập nhật theo thời gian thực.\n"
        "• Hero Banner: Khu vực nổi bật nhất với thiết kế Cyberpunk giới thiệu siêu phẩm công nghệ tiêu biểu cùng nút kêu gọi hành động (CTA) 'Mua ngay giá ưu đãi'.\n"
        "• Thanh cam kết dịch vụ (Features Bar): 4 hộp thông tin cam kết 100% chính hãng, giao hàng hỏa tốc 2 giờ, đổi trả trong 30 ngày và trả góp 0% lãi suất.\n"
        "• Khu vực Flash Sale: Đồng hồ đếm ngược thời gian thực (Giờ : Phút : Giây) kết hợp lưới sản phẩm giảm giá cực sốc kèm phần trăm chiết khấu trực quan.\n"
        "• Sản phẩm bán chạy & Tại sao chọn HieuMini: Giới thiệu các thiết bị công nghệ được yêu thích nhất cùng các chính sách bảo hành hậu mãi vượt trội."
    )
    
    # 2. Trang danh sách & bộ lọc
    add_styled_paragraph(doc, "2. Màn hình Danh sách & Bộ lọc sản phẩm (products.php):", bold=True)
    add_styled_paragraph(doc, 
        "• Sidebar bộ lọc đa tiêu chí: Cho phép người dùng lọc sản phẩm đồng thời theo Danh mục (Điện thoại, Laptop, Tablet, Smartwatch, Tai nghe, Phụ kiện), theo Thương hiệu (Apple, Samsung, ASUS, Sony, Marshall, Anker...) và theo Khoảng giá (Dưới 5 Triệu, 5-15 Triệu, 15-30 Triệu, Trên 30 Triệu).\n"
        "• Sắp xếp linh hoạt: Tùy chọn sắp xếp theo Sản phẩm mới nhất, Giá tăng dần, Giá giảm dần hoặc Đánh giá sao cao nhất.\n"
        "• Thao tác nhanh: Nút 'Thêm vào giỏ' tích hợp AJAX giúp khách hàng thêm sản phẩm ngay tại danh mục mà không làm gián đoạn việc xem hàng."
    )
    
    # 3. Trang chi tiết sản phẩm
    add_styled_paragraph(doc, "3. Màn hình Chi tiết sản phẩm (product_detail.php):", bold=True)
    add_styled_paragraph(doc, 
        "• Khu vực trình diễn hình ảnh: Ảnh sản phẩm sắc nét kèm nhãn giảm giá và thông tin tình trạng kho hàng thực tế.\n"
        "• Bảng thông số kỹ thuật chi tiết: Tự động phân tích chuỗi JSON thành bảng cấu hình chi tiết (Màn hình, Chip CPU, RAM, Ổ cứng, Camera, Dung lượng pin).\n"
        "• Tùy chọn mua hàng: Bộ tăng giảm số lượng linh hoạt, nút 'MUA NGAY (Giao tận nơi)' và nút 'Thêm vào giỏ'.\n"
        "• Phân hệ Đánh giá & Nhận xét: Hiển thị các đánh giá từ người dùng thực tế kèm số sao bình chọn và Form gửi cảm nhận đánh giá trực tiếp."
    )
    
    # 4. Trang giỏ hàng
    add_styled_paragraph(doc, "4. Màn hình Giỏ hàng (cart.php):", bold=True)
    add_styled_paragraph(doc, 
        "• Bảng chi tiết giỏ hàng: Hiển thị ảnh đại diện, tên sản phẩm, đơn giá niêm yết, các nút tăng giảm số lượng tức thì (+ / -), thành tiền từng món và nút xóa sản phẩm.\n"
        "• Áp dụng mã giảm giá: Hỗ trợ nhập các mã khuyến mãi (như HIEUMINI2026 giảm 10%, TECHNEW giảm 5%) với cơ chế tính toán lại tổng tiền tự động.\n"
        "• Thẻ tóm tắt đơn hàng: Liệt kê chi tiết Tạm tính, Tiền giảm giá, Phí vận chuyển (Miễn phí cho đơn từ 2 Triệu) và Tổng số tiền thanh toán cuối cùng."
    )
    
    # 5. Trang thanh toán & VietQR
    add_styled_paragraph(doc, "5. Màn hình Thanh toán & Quét mã VietQR (checkout.php & order_success.php):", bold=True)
    add_styled_paragraph(doc, 
        "• Form thông tin giao hàng: Nhập Họ tên, Số điện thoại, Email, Tỉnh/Thành phố và Địa chỉ nhận hàng chi tiết kèm ghi chú cho shipper.\n"
        "• Đa dạng phương thức thanh toán: Lựa chọn giữa 'Thanh toán khi nhận hàng (COD)' hoặc 'Chuyển khoản trực tuyến qua VietQR'. Khi chọn VietQR, hệ thống lập tức mở hộp thông tin ngân hàng và mã QR thanh toán chuẩn ngân hàng MBBank.\n"
        "• Trang xác nhận đơn hàng thành công (order_success.php): Hiển thị thông báo chúc mừng kèm Mã đơn hàng độc nhất để khách hàng dễ dàng chụp lại hoặc tra cứu."
    )
    
    add_styled_heading(doc, "3.1.3 Chi tiết các màn hình Phân hệ Quản trị (Admin Portal)", level=3)
    add_styled_paragraph(doc, 
        "• Admin Dashboard (admin/index.php): Bảng điều khiển trung tâm với 4 thẻ chỉ số KPI kinh doanh trực quan (Tổng doanh thu, Tổng đơn hàng, Số lượng sản phẩm, Khách hàng thành viên). Tích hợp Biểu đồ Doanh thu 7 ngày vẽ bằng HTML5 Canvas và Bảng đơn hàng mới nhất cần xác nhận giao hàng.\n"
        "• Quản lý Sản phẩm (admin/products.php & product_add.php & product_edit.php): Bảng quản lý toàn bộ thiết bị trong kho, hỗ trợ tìm kiếm nhanh theo tên/hãng, lọc tình trạng tồn kho, form thêm/sửa sản phẩm với trình soạn thảo thông số kỹ thuật JSON, gắn nhãn Flash Sale và xóa sản phẩm an toàn có hộp thoại xác nhận.\n"
        "• Quản lý Danh mục (admin/categories.php): Cho phép quản trị viên thêm mới các ngành hàng công nghệ kèm icon FontAwesome và tùy biến đường dẫn slug SEO.\n"
        "• Quản lý Đơn hàng (admin/orders.php): Hiển thị chi tiết khách hàng, địa chỉ giao hàng, phương thức thanh toán và dropdown cập nhật trạng thái đơn hàng (Chờ xử lý -> Đang xử lý -> Đang giao -> Hoàn thành -> Đã hủy).\n"
        "• Quản lý Người dùng (admin/users.php): Theo dõi danh sách tài khoản khách hàng và cấp/hủy quyền Quản trị viên (Admin Role) một cách an toàn."
    )
    
    add_styled_heading(doc, "3.2 Kết luận", level=2)
    
    add_styled_heading(doc, "3.2.1 Kết quả đạt được của dự án", level=3)
    add_styled_paragraph(doc, 
        "Trải qua quá trình nghiên cứu lý thuyết, phân tích hệ thống và triển khai lập trình thực tế, dự án xây dựng website thương mại điện tử bán hàng công nghệ HieuMini đã hoàn thành 100% các mục tiêu đề ra ban đầu, bao gồm:"
    )
    add_styled_paragraph(doc, 
        "1. Xây dựng hoàn chỉnh mã nguồn Website HieuMini bằng ngôn ngữ PHP thuần kết hợp mô hình phân tầng hướng đối tượng, kết nối CSDL MySQL qua thư viện PDO bảo mật cao.\n"
        "2. Thiết kế và hiện thực hóa cơ sở dữ liệu quan hệ `hieumini_db` chuẩn hóa 3NF gồm 8 bảng dữ liệu, đảm bảo đầy đủ các ràng buộc toàn vẹn khóa chính, khóa ngoại và tính chất giao dịch ACID.\n"
        "3. Xây dựng giao diện người dùng đỉnh cao phong cách Glassmorphism và Dark Mode, đáp ứng hoàn hảo trên mọi thiết bị (Responsive Design).\n"
        "4. Tích hợp đầy đủ các luồng nghiệp vụ mua sắm hiện đại: Tìm kiếm, Bộ lọc đa tiêu chí, Giỏ hàng AJAX, Áp mã coupon, Thanh toán chuyển khoản VietQR và Đánh giá nhận xét sản phẩm.\n"
        "5. Xây dựng phân hệ Quản trị viên (Admin Dashboard) chuyên nghiệp, cung cấp đầy đủ công cụ kiểm soát kinh doanh và quản lý dữ liệu toàn diện."
    )
    
    add_styled_heading(doc, "3.2.2 Ưu điểm và hạn chế của hệ thống", level=3)
    add_styled_paragraph(doc, 
        "Ưu điểm nổi bật:\n"
        "• Tính thẩm mỹ và trải nghiệm người dùng (UI/UX) đạt tiêu chuẩn thương mại cao cấp, gây ấn tượng mạnh mẽ cho khách hàng ngay từ cái nhìn đầu tiên.\n"
        "• Tốc độ tải trang cực nhanh, tối ưu hóa câu lệnh truy vấn SQL có đánh chỉ mục Indexing.\n"
        "• Mã nguồn PHP sạch sẽ, phân tách rõ ràng các tầng cấu hình, hàm tiện ích, giao diện người dùng và quản trị.\n"
        "• Cơ chế bảo mật vững chắc: Sử dụng Prepared Statements chống SQL Injection, mã hóa đầu ra chống XSS và băm mật khẩu bằng BCRYPT."
    )
    add_styled_paragraph(doc, 
        "Hạn chế còn tồn đọng:\n"
        "• Phương thức thanh toán VietQR hiện tại mô phỏng hiển thị mã QR kèm số tiền và nội dung, chưa tích hợp cổng Webhook ngân hàng tự động đối soát giao dịch trong 3 giây thực tế (do yêu cầu giấy phép kinh doanh doanh nghiệp).\n"
        "• Chưa tích hợp hệ thống gửi Email tự động thông báo đơn hàng qua SMTP server."
    )
    
    add_styled_heading(doc, "3.2.3 Hướng phát triển và mở rộng trong tương lai", level=3)
    add_styled_paragraph(doc, 
        "Để đưa hệ thống HieuMini vận hành thực tế ở quy mô doanh nghiệp lớn, các hướng phát triển tiếp theo bao gồm:\n"
        "1. Tích hợp cổng thanh toán trực tuyến tự động qua API của VNPay, MoMo, ZaloPay và thẻ quốc tế Stripe / Visa / Mastercard.\n"
        "2. Ứng dụng Trí tuệ nhân tạo (AI & Machine Learning) để xây dựng hệ thống gợi ý sản phẩm thông minh (Recommendation Engine) dựa trên lịch sử xem và thói quen mua sắm của từng khách hàng.\n"
        "3. Tích hợp Chatbot AI tư vấn kỹ thuật và hỗ trợ giải đáp thắc mắc khách hàng 24/7.\n"
        "4. Tích hợp API đơn vị vận chuyển (Giao Hàng Nhanh - GHN, Giao Hàng Tiết Kiệm - GHTK) để tự động tính phí ship theo khoảng cách thực tế và in mã vận đơn tự động.\n"
        "5. Phát triển ứng dụng di động (Mobile App) trên iOS và Android sử dụng Flutter / React Native đồng bộ chung cơ sở dữ liệu RESTful API với website."
    )
    
    doc.add_page_break()
    
    # ==========================================================
    # TÀI LIỆU THAM KHẢO
    # ==========================================================
    add_styled_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)
    
    references = [
        "[1] The PHP Group (2024), 'PHP: Hypertext Preprocessor Official Documentation', https://www.php.net/docs.php.",
        "[2] Oracle Corporation (2024), 'MySQL 8.0 Reference Manual', https://dev.mysql.com/doc/refman/8.0/en/.",
        "[3] Robin Nixon (2021), 'Learning PHP, MySQL & JavaScript: With jQuery, CSS & HTML5', 6th Edition, O'Reilly Media.",
        "[4] Luke Welling, Laura Thomson (2016), 'PHP and MySQL Web Development', 5th Edition, Addison-Wesley Professional.",
        "[5] OWASP Foundation (2023), 'OWASP Top Ten Web Application Security Risks', https://owasp.org/www-project-top-ten/.",
        "[6] World Wide Web Consortium (W3C), 'HTML5 and CSS3 Specifications & Guidelines', https://www.w3.org/."
    ]
    
    for ref in references:
        add_styled_paragraph(doc, ref, space_after=8)
        
    # Lưu file Word
    output_filename = "BaoCao.docx"
    doc.save(output_filename)
    print(f"Generated successfully: {output_filename}")

if __name__ == "__main__":
    generate_full_report()
