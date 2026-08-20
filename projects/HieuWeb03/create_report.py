import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_report():
    doc = Document()

    # Page setup - A4 with 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    # Color Palette
    PRIMARY_NAVY = RGBColor(30, 41, 59)     # #1E293B
    ACCENT_INDIGO = RGBColor(79, 70, 229)   # #4F46E5
    TEXT_DARK = RGBColor(15, 23, 42)        # #0F172A
    MUTED_GRAY = RGBColor(100, 116, 139)    # #64748B

    # Helper: Title Formatting
    def add_custom_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = ACCENT_INDIGO
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(20)
        return p

    def add_custom_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_NAVY
        return p

    def add_custom_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = ACCENT_INDIGO
        return p

    def add_custom_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_NAVY
        return p

    def add_body_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.25
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Times New Roman'
            r_bold.font.size = Pt(12)
            r_bold.font.bold = True
            r_bold.font.color.rgb = TEXT_DARK
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = italic
        run.font.color.rgb = TEXT_DARK
        return p

    def add_bullet_p(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Times New Roman'
            r_bold.font.size = Pt(12)
            r_bold.font.bold = True
            r_bold.font.color.rgb = TEXT_DARK
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = TEXT_DARK
        return p

    def add_callout_box(text, title="LƯU Ý QUAN TRỌNG"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "EEF2FF")
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r_title = p.add_run(f"📌 {title}: ")
        r_title.font.name = 'Times New Roman'
        r_title.font.size = Pt(11.5)
        r_title.font.bold = True
        r_title.font.color.rgb = ACCENT_INDIGO

        r_text = p.add_run(text)
        r_text.font.name = 'Times New Roman'
        r_text.font.size = Pt(11.5)
        r_text.font.color.rgb = TEXT_DARK

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # =========================================================================
    # 1. TRANG BÌA HỌC THUẬT (COVER PAGE)
    # =========================================================================
    p_cov_top = doc.add_paragraph()
    p_cov_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_cov_top.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC KHOA HỌC VÀ CÔNG NGHỆ\nKHOA CÔNG NGHỆ THÔNG TIN\n---------------------------------------")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = PRIMARY_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rep = p_title.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC\n")
    r_rep.font.name = 'Times New Roman'
    r_rep.font.size = Pt(16)
    r_rep.font.bold = True
    r_rep.font.color.rgb = MUTED_GRAY

    r_main = p_title.add_run("PHÂN TÍCH, THIẾT KẾ VÀ XÂY DỰNG WEBSITE BÁN ĐỒ DÙNG HỌC TẬP VÀ VĂN PHÒNG PHẨM TRỰC TUYẾN “HIEUMINI”")
    r_main.font.name = 'Times New Roman'
    r_main.font.size = Pt(20)
    r_main.font.bold = True
    r_main.font.color.rgb = ACCENT_INDIGO

    p_tech = doc.add_paragraph()
    p_tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_tech.add_run("\nCông nghệ ứng dụng: Ngôn ngữ lập trình PHP & Cơ sở dữ liệu MySQL\nThiết kế UI/UX theo tiêu chuẩn Hiện đại, Tương tác động Transitions & Animations")
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(13)
    r_t.font.italic = True
    r_t.font.color.rgb = PRIMARY_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(60)

    # Student Info Table on Cover
    cov_table = doc.add_table(rows=4, cols=2)
    cov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Sinh viên thực hiện:", "Trần Văn Minh Hiếu"),
        ("Mã số sinh viên:", "2026-IT-HIEU06"),
        ("Chuyên ngành:", "Công nghệ Phần mềm & Lập trình Web"),
        ("Giảng viên hướng dẫn:", "TS. Nguyễn Văn Hướng")
    ]
    for idx, (label, val) in enumerate(info_data):
        row = cov_table.rows[idx]
        cell_l = row.cells[0]
        cell_r = row.cells[1]
        
        p_l = cell_l.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_l = p_l.add_run(label)
        r_l.font.name = 'Times New Roman'
        r_l.font.size = Pt(13)
        r_l.font.bold = True
        r_l.font.color.rgb = PRIMARY_NAVY

        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_r = p_r.add_run(val)
        r_r.font.name = 'Times New Roman'
        r_r.font.size = Pt(13)
        r_r.font.bold = False
        r_r.font.color.rgb = TEXT_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(50)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_d = p_date.add_run("Hà Nội, Năm 2026")
    r_d.font.name = 'Times New Roman'
    r_d.font.size = Pt(12)
    r_d.font.bold = True

    doc.add_page_break()

    # =========================================================================
    # 2. MỤC LỤC CHI TIẾT (TABLE OF CONTENTS)
    # =========================================================================
    p_ml = doc.add_paragraph()
    p_ml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ml = p_ml.add_run("MỤC LỤC NỘI DUNG BÁO CÁO")
    r_ml.font.name = 'Times New Roman'
    r_ml.font.size = Pt(18)
    r_ml.font.bold = True
    r_ml.font.color.rgb = PRIMARY_NAVY
    p_ml.paragraph_format.space_after = Pt(18)

    toc_items = [
        ("LỜI NÓI ĐẦU VÀ TÍNH CẤP THIẾT CỦA ĐỀ TÀI", "Trang 3"),
        ("CHƯƠNG 1. TỔNG QUAN LẬP TRÌNH WEB", "Trang 4"),
        ("   1.1 Ngôn ngữ lập trình PHP", "Trang 4"),
        ("       1.1.1 Lịch sử phát triển và kiến trúc PHP", "Trang 4"),
        ("       1.1.2 Cơ chế hoạt động Client - Server trong ứng dụng web", "Trang 5"),
        ("       1.1.3 Ưu điểm vượt trội và các tính năng cốt lõi của PHP 8.x", "Trang 6"),
        ("       1.1.4 Cơ chế bảo mật và giao tiếp cơ sở dữ liệu qua PDO", "Trang 7"),
        ("   1.2 Hệ quản trị cơ sở dữ liệu MySQL", "Trang 8"),
        ("       1.2.1 Khái niệm và kiến trúc cơ sở dữ liệu quan hệ (RDBMS)", "Trang 8"),
        ("       1.2.2 Hệ thống lưu trữ InnoDB, tính toàn vẹn và giao dịch ACID", "Trang 9"),
        ("       1.2.3 Khóa chính, khóa ngoại và kỹ thuật tối ưu hóa truy vấn SQL", "Trang 10"),
        ("   1.3 Cài đặt máy chủ", "Trang 11"),
        ("       1.3.1 Tổng quan môi trường phát triển XAMPP / Apache / MySQL", "Trang 11"),
        ("       1.3.2 Cấu hình tệp tin php.ini và my.ini chuyên sâu", "Trang 12"),
        ("       1.3.3 Hướng dẫn thiết lập và quản lý dữ liệu qua phpMyAdmin", "Trang 13"),
        ("CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ WEBSITE", "Trang 14"),
        ("   2.1 Chức năng (Use case)", "Trang 14"),
        ("       2.1.1 Khảo sát nghiệp vụ cửa hàng đồ dùng học tập HieuMini", "Trang 14"),
        ("       2.1.2 Phân tích tác nhân (Actors) và biểu đồ Use Case tổng quát", "Trang 15"),
        ("       2.1.3 Đặc tả chi tiết các Use Case phân hệ Người dùng (Khách hàng)", "Trang 16"),
        ("       2.1.4 Đặc tả chi tiết các Use Case phân hệ Quản trị viên (Admin)", "Trang 18"),
        ("   2.2 Cơ sở dữ liệu", "Trang 20"),
        ("       2.2.1 Sơ đồ quan hệ thực thể ERD (Entity Relationship Diagram)", "Trang 20"),
        ("       2.2.2 Thiết kế cấu trúc chi tiết 8 bảng cơ sở dữ liệu", "Trang 21"),
        ("       2.2.3 Ràng buộc toàn vẹn, khóa ngoại và kịch bản dữ liệu 30 sản phẩm", "Trang 24"),
        ("CHƯƠNG 3. CHƯƠNG TRÌNH THỬ NGHIỆM", "Trang 25"),
        ("   3.1 Giao diện", "Trang 25"),
        ("       3.1.1 Ngôn ngữ thiết kế UI/UX hiện đại và hệ thống Design Tokens", "Trang 25"),
        ("       3.1.2 Hệ thống Transitions và Micro-Animations tương tác sống động", "Trang 26"),
        ("       3.1.3 Giao diện phân hệ người dùng (Trang chủ, Catalog, Chi tiết, Giỏ hàng, Checkout, Success)", "Trang 27"),
        ("       3.1.4 Giao diện phân hệ Quản trị Admin (Dashboard, CRUD Sản phẩm, Đơn hàng)", "Trang 30"),
        ("   3.2 Kết luận", "Trang 32"),
        ("       3.2.1 Đánh giá kết quả đạt được đối chiếu mục tiêu", "Trang 32"),
        ("       3.2.2 Ưu điểm nổi bật của website HieuMini", "Trang 33"),
        ("       3.2.3 Hạn chế còn tồn đọng và hướng phát triển tương lai", "Trang 34"),
        ("TÀI LIỆU THAM KHẢO", "Trang 35")
    ]

    for title_text, page_str in toc_items:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(1)
        p_t.paragraph_format.space_after = Pt(2)
        r_left = p_t.add_run(title_text)
        r_left.font.name = 'Times New Roman'
        r_left.font.size = Pt(11.5)
        if "CHƯƠNG" in title_text or "LỜI NÓI ĐẦU" in title_text or "TÀI LIỆU" in title_text:
            r_left.font.bold = True
            r_left.font.color.rgb = PRIMARY_NAVY
        else:
            r_left.font.color.rgb = TEXT_DARK
        
        # Dots filler
        dots_count = max(2, 85 - len(title_text))
        r_dots = p_t.add_run(" " + "." * dots_count + " ")
        r_dots.font.color.rgb = MUTED_GRAY

        r_page = p_t.add_run(page_str)
        r_page.font.bold = True
        r_page.font.color.rgb = ACCENT_INDIGO

    doc.add_page_break()

    # =========================================================================
    # LỜI NÓI ĐẦU
    # =========================================================================
    add_custom_h1("LỜI NÓI ĐẦU VÀ TÍNH CẤP THIẾT CỦA ĐỀ TÀI")
    
    add_body_p("Trong kỷ nguyên chuyển đổi số và thương mại điện tử phát triển bùng nổ, việc ứng dụng công nghệ thông tin vào mô hình kinh doanh bán lẻ không chỉ là một giải pháp gia tăng doanh số mà đã trở thành yếu tố sống còn đối với mọi doanh nghiệp. Thị trường đồ dùng học tập, văn phòng phẩm và dụng cụ mỹ thuật sáng tạo tại Việt Nam là một thị trường có quy mô rất lớn, với đối tượng khách hàng mục tiêu là hàng triệu học sinh, sinh viên, giáo viên, giới văn phòng và những người yêu thích nghệ thuật ghi chép (Journaling, Bullet Journal, Calligraphy).")
    
    add_body_p("Tuy nhiên, phần lớn các cửa hàng văn phòng phẩm truyền thống hiện nay vẫn phụ thuộc vào việc bán hàng trực tiếp tại điểm bán, hoặc kinh doanh nhỏ lẻ trên mạng xã hội với quy trình quản lý đơn hàng thủ công, dễ thất thoát và thiếu tính chuyên nghiệp. Nhằm giải quyết các bài toán trên, đồ án ", "Tính cấp thiết của đề tài: ")
    add_body_p("“Xây dựng website bán đồ dùng học tập và văn phòng phẩm HieuMini sử dụng PHP và MySQL” được nghiên cứu và hiện thực hóa. Dự án hướng tới mục tiêu xây dựng một nền tảng thương mại điện tử hoàn chỉnh, tích hợp đầy đủ các quy trình nghiệp vụ mua bán, quản lý kho hàng với 30 sản phẩm thực tế sắc nét, trải nghiệm giao diện người dùng (UI/UX) đỉnh cao với hiệu ứng chuyển động mượt mà Transitions & Animations, cùng hệ thống quản trị nội dung (Admin Dashboard) mạnh mẽ, an toàn và bảo mật.")

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 1. TỔNG QUAN LẬP TRÌNH WEB
    # =========================================================================
    add_custom_h1("Chương 1. Tổng quan lập trình web")
    
    add_body_p("Lập trình web là một lĩnh vực trọng yếu trong công nghệ phần mềm, bao gồm các công nghệ xây dựng giao diện người dùng (Front-end), xử lý logic nghiệp vụ phía máy chủ (Back-end) và quản lý cơ sở dữ liệu (Database). Trong chương này, báo cáo sẽ trình bày chi tiết về ngôn ngữ lập trình PHP, hệ quản trị cơ sở dữ liệu quan hệ MySQL và quy trình cài đặt, cấu hình môi trường máy chủ cục bộ phục vụ việc vận hành hệ thống website HieuMini.")

    # 1.1 PHP
    add_custom_h2("1.1 Ngôn ngữ lập trình PHP")
    
    add_body_p("PHP (viết tắt đệ quy của Hypertext Preprocessor) là một ngôn ngữ lập trình kịch bản mã nguồn mở chạy ở phía máy chủ (Server-side Scripting Language), được thiết kế chuyên biệt để phát triển các ứng dụng web động. Được Rasmus Lerdorf sáng tạo vào năm 1994, trải qua hơn 3 thập kỷ không ngừng hoàn thiện, PHP đã phát triển mạnh mẽ từ các tập lệnh đơn giản thành một ngôn ngữ hướng đối tượng hoàn chỉnh, đa năng và có mặt trên hơn 75% số website sử dụng mã nguồn phía máy chủ trên toàn cầu (theo thống kê W3Techs).")

    add_custom_h3("1.1.1 Lịch sử phát triển và kiến trúc PHP")
    add_body_p("Sự tiến hóa của PHP được đánh dấu qua các mốc son công nghệ quan trọng:", "Các giai đoạn phát triển chính: ")
    add_bullet_p("PHP 1.0 & 2.0 (1994 - 1997): Khởi nguồn từ bộ công cụ Personal Home Page Tools viết bằng ngôn ngữ C dùng để theo dõi lượt truy cập hồ sơ cá nhân.", "• ")
    add_bullet_p("PHP 3.0 & 4.0 (1998 - 2000): Zeev Suraski và Andi Gutmans viết lại hoàn toàn lõi phân tích cú pháp, khai sinh Zend Engine 1.0, đưa PHP trở thành nền tảng lập trình web chuẩn mực.", "• ")
    add_bullet_p("PHP 5.x (2004 - 2014): Giới thiệu Zend Engine 2.0 với mô hình Hướng đối tượng (OOP) hoàn thiện, hỗ trợ ngoại lệ (Exceptions), thư viện PDO (PHP Data Objects) và tích hợp XML.", "• ")
    add_bullet_p("PHP 7.x (2015 - 2020): Bước nhảy vọt về hiệu năng nhờ Zend Engine 3.0, giảm 50% mức tiêu thụ bộ nhớ RAM và tăng gấp đôi tốc độ thực thi lệnh.", "• ")
    add_bullet_p("PHP 8.x (2020 - nay - Phiên bản ứng dụng trong dự án HieuMini): Giới thiệu trình biên dịch JIT (Just-In-Time Compiler), Union Types, Named Arguments, Constructor Property Promotion, Match Expression và Fiber hỗ trợ xử lý bất đồng bộ.", "• ")

    add_custom_h3("1.1.2 Cơ chế hoạt động Client - Server trong ứng dụng web")
    add_body_p("Kiến trúc hoạt động của ứng dụng PHP tuân thủ chặt chẽ mô hình Client-Server Request-Response không trạng thái (Stateless):")
    add_bullet_p("Bước 1: Trình duyệt web (Client) gửi yêu cầu HTTP Request (GET/POST) tới Web Server (Apache/Nginx) thông qua giao thức TCP/IP.", "1. ")
    add_bullet_p("Bước 2: Máy chủ web tiếp nhận yêu cầu, phân tích định dạng tệp tin. Nếu là tệp tĩnh (.html, .css, .js, .png), máy chủ trả về ngay lập tức. Nếu là tệp .php, máy chủ chuyển giao cho trình thông dịch PHP Module (Zend Engine).", "2. ")
    add_bullet_p("Bước 3: Trình thông dịch PHP phân tích cú pháp, nạp cấu hình, kết nối tới Cơ sở dữ liệu MySQL để truy vấn hoặc cập nhật dữ liệu hàng hóa, tài khoản, giỏ hàng.", "3. ")
    add_bullet_p("Bước 4: PHP xử lý toàn bộ logic nghiệp vụ (tính toán giảm giá, kiểm tra session người dùng, sinh mã đơn hàng), biên dịch mã nguồn thành tài liệu HTML/CSS/JSON thuần túy.", "4. ")
    add_bullet_p("Bước 5: Web Server đóng gói tài liệu đã biên dịch thành HTTP Response trả về cho trình duyệt Client hiển thị cho người dùng.", "5. ")

    add_custom_h3("1.1.3 Ưu điểm vượt trội của PHP trong xây dựng sàn thương mại điện tử")
    add_body_p("Ngôn ngữ PHP sở hữu những ưu điểm nổi bật giúp nó trở thành lựa chọn tối ưu cho website HieuMini:")
    add_bullet_p("Mã nguồn mở và chi phí triển khai 0 đồng: PHP hoàn toàn miễn phí, có thể chạy mượt mà trên nhiều hệ điều hành (Windows, Linux, macOS, Unix).", "• ")
    add_bullet_p("Cộng đồng hỗ trợ khổng lồ: Tài liệu hướng dẫn phong phú, thư viện mở rộng đa dạng đáp ứng mọi nhu cầu xử lý hình ảnh, mã hóa dữ liệu, thanh toán điện tử.", "• ")
    add_bullet_p("Dễ dàng nhúng vào HTML: Cho phép lập trình viên kết hợp mã PHP trực tiếp bên trong cấu trúc HTML, tạo điều kiện thuận lợi cho việc kết xuất dữ liệu động (Server-Side Rendering).", "• ")
    add_bullet_p("Khả năng tương thích cơ sở dữ liệu cao: Hỗ trợ sẵn các trình điều khiển kết nối native tới hầu hết các RDBMS phổ biến như MySQL, PostgreSQL, SQLite, Oracle, SQL Server.", "• ")

    add_custom_h3("1.1.4 Cơ chế bảo mật và giao tiếp CSDL qua PDO (PHP Data Objects)")
    add_body_p("Trong dự án HieuMini, thay vì sử dụng thư viện cũ `mysql_*` đã bị loại bỏ hoặc `mysqli_*` phụ thuộc vào MySQL, hệ thống sử dụng độc quyền ", "Chuẩn giao tiếp PDO: ")
    add_body_p("thư viện trừu tượng hóa cơ sở dữ liệu PDO (PHP Data Objects). PDO mang lại tính nhất quán cao và hỗ trợ kỹ thuật Prepared Statements (tham số hóa truy vấn) với cơ chế bảo vệ 100% chống lại các cuộc tấn công tiêm mã độc SQL Injection nguy hiểm.")

    add_callout_box("Dự án HieuMini sử dụng Prepared Statements của PDO cho 100% các câu lệnh truy vấn có chứa tham số đầu vào từ người dùng, kết hợp mã hóa mật khẩu một chiều an toàn bằng thuật toán bcrypt (hàm password_hash() và password_verify()).", "BẢO MẬT HỆ THỐNG");

    # 1.2 MySQL
    add_custom_h2("1.2 Hệ quản trị cơ sở dữ liệu MySQL")
    
    add_body_p("MySQL là hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở (RDBMS) phổ biến nhất thế giới, được sở hữu và phát triển bởi tập đoàn Oracle. MySQL sử dụng ngôn ngữ truy vấn có cấu trúc chuẩn SQL (Structured Query Language) để định nghĩa (DDL), thao tác (DML) và kiểm soát dữ liệu (DCL).")

    add_custom_h3("1.2.1 Khái niệm và kiến trúc cơ sở dữ liệu quan hệ")
    add_body_p("Cơ sở dữ liệu quan hệ tổ chức dữ liệu dưới dạng các bảng (Tables) gồm các hàng (Rows / Records) đại diện cho các đối tượng thực thể và các cột (Columns / Fields) đại diện cho các thuộc tính của thực thể đó. Mối liên hệ giữa các bảng được thiết lập thông qua các cặp Khóa chính (Primary Key) và Khóa ngoại (Foreign Key), đảm bảo loại bỏ tình trạng dư thừa dữ liệu (Data Redundancy) và tuân thủ các dạng chuẩn hóa cơ sở dữ liệu (1NF, 2NF, 3NF).")

    add_custom_h3("1.2.2 Engine lưu trữ InnoDB và nguyên tắc giao dịch ACID")
    add_body_p("Hệ thống website HieuMini sử dụng Storage Engine mặc định là ", "Công nghệ InnoDB: ")
    add_body_p("InnoDB thay vì MyISAM. InnoDB là công cụ lưu trữ tiêu chuẩn doanh nghiệp, hỗ trợ đầy đủ các tính năng quan trọng:")
    add_bullet_p("Khóa mức hàng (Row-level Locking): Tăng cường khả năng xử lý đồng thời (Concurrency) khi nhiều khách hàng cùng đặt hàng một lúc mà không bị nghẽn bảng.", "• ")
    add_bullet_p("Ràng buộc Khóa ngoại (Foreign Key Constraints): Tự động kiểm tra tính toàn vẹn tham chiếu, ngăn chặn xóa các danh mục đang chứa sản phẩm hoặc các đơn hàng mồ côi.", "• ")
    add_bullet_p("Nguyên tắc ACID trong quản lý giao dịch (Transactions):", "• ")
    add_bullet_p("   + Atomicity (Tính nguyên tử): Toàn bộ thao tác tạo đơn hàng và trừ số lượng tồn kho sản phẩm hoặc cùng thành công hoặc cùng bị hủy bỏ (Rollback).", "")
    add_bullet_p("   + Consistency (Tính nhất quán): Dữ liệu luôn chuyển đổi từ trạng thái hợp lệ này sang trạng thái hợp lệ khác tuân theo mọi ràng buộc.", "")
    add_bullet_p("   + Isolation (Tính độc lập): Các giao dịch thực hiện đồng thời không can thiệp hay nhìn thấy trạng thái trung gian của nhau.", "")
    add_bullet_p("   + Durability (Tính bền vững): Dữ liệu sau khi xác nhận (Commit) được ghi bền vững vào đĩa cứng kể cả khi mất điện đột ngột.", "")

    add_custom_h3("1.2.3 Khóa chính, khóa ngoại và tối ưu hóa truy vấn")
    add_body_p("Các bảng dữ liệu của website HieuMini (`products`, `categories`, `orders`, `order_items`, `users`, `reviews`, `coupons`, `contacts`) đều được thiết lập khóa chính tự tăng (AUTO_INCREMENT INT PRIMARY KEY). Các trường tìm kiếm thường xuyên như `slug`, `sku`, `email`, `order_code` đều được đánh chỉ mục UNIQUE INDEX giúp tốc độ tìm kiếm đạt độ phức tạp O(log N) thay vì O(N), giảm thiểu tối đa thời gian phản hồi máy chủ.")

    # 1.3 Cài đặt máy chủ
    add_custom_h2("1.3 Cài đặt máy chủ")
    
    add_body_p("Để triển khai và thử nghiệm website HieuMini, một môi trường máy chủ web đầy đủ (Web Server Stack) bao gồm hệ điều hành, máy chủ HTTP Apache, trình thông dịch PHP và hệ quản trị MySQL là điều kiện tiên quyết.")

    add_custom_h3("1.3.1 Tổng quan môi trường phát triển XAMPP")
    add_body_p("XAMPP là một gói phần mềm mã nguồn mở hoàn toàn miễn phí và cực kỳ phổ biến của Apache Friends. Tên gọi XAMPP là từ viết tắt của: X (Cross-platform: Đa nền tảng), A (Apache HTTP Server), M (MariaDB/MySQL), P (PHP) và P (Perl). XAMPP cung cấp bảng điều khiển XAMPP Control Panel giúp người quản trị dễ dàng khởi động (Start), dừng (Stop) và cấu hình các dịch vụ máy chủ chỉ bằng một cú nhấp chuột.")

    add_custom_h3("1.3.2 Cấu hình tệp tin php.ini và my.ini chuyên sâu")
    add_body_p("Để hệ thống website bán đồ dùng học tập vận hành trơn tru, hỗ trợ xử lý ảnh sản phẩm dung lượng cao và phiên làm việc giỏ hàng ổn định, các tham số cấu hình trong `php.ini` và `my.ini` được tinh chỉnh như sau:")
    add_bullet_p("max_execution_time = 300: Cho phép thời gian thực thi tối đa của script là 300 giây, tránh timeout khi kết xuất báo cáo hoặc tạo ảnh hàng loạt.", "• ")
    add_bullet_p("memory_limit = 256M: Phân bổ dung lượng bộ nhớ RAM tối đa cho PHP xử lý đồ họa ảnh và truy vấn CSDL lớn.", "• ")
    add_bullet_p("upload_max_filesize = 64M & post_max_size = 64M: Hỗ trợ người quản trị tải lên các hình ảnh đồ dùng học tập chất lượng cao sắc nét.", "• ")
    add_bullet_p("session.gc_maxlifetime = 14400: Duy trì phiên đăng nhập và giỏ hàng của khách hàng trong 4 giờ liên tục không bị mất.", "• ")
    add_bullet_p("default_character_set = 'UTF-8': Đảm bảo hiển thị chuẩn xác tiếng Việt có dấu trong toàn bộ hệ thống.", "• ")

    add_custom_h3("1.3.3 Quản trị CSDL trực quan với phpMyAdmin và PHP Built-in Server")
    add_body_p("Dự án HieuMini cung cấp sẵn tệp kịch bản CSDL chuẩn `database.sql`. Người quản trị có thể dễ dàng khởi tạo cơ sở dữ liệu `hieumini_db` thông qua giao diện đồ họa phpMyAdmin (tại đường dẫn `http://localhost/phpmyadmin`) hoặc chạy máy chủ phát triển tích hợp sẵn của PHP bằng câu lệnh cực kỳ tinh gọn:")
    
    add_callout_box("Khởi chạy máy chủ thử nghiệm nhanh tại thư mục gốc dự án:\n> php -S localhost:8000\nTruy cập ngay trên trình duyệt: http://localhost:8000\nĐăng nhập Admin: admin@hieumini.vn / admin123\nĐăng nhập Khách hàng: user@hieumini.vn / user123", "HƯỚNG DẪN CHẠY DỰ ÁN");

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ WEBSITE
    # =========================================================================
    add_custom_h1("Chương 2. Phân tích và thiết kế website")
    
    add_body_p("Giai đoạn phân tích và thiết kế hệ thống đóng vai trò nền tảng quyết định sự thành công và khả năng mở rộng lâu dài của một ứng dụng web thương mại điện tử. Chương 2 tập trung phân tích chi tiết các yêu cầu chức năng thông qua biểu đồ ca sử dụng (Use Case), mô tả các tác nhân (Actors) và xây dựng lược đồ cơ sở dữ liệu quan hệ (ERD) cùng đặc tả chi tiết các bảng dữ liệu của website HieuMini.")

    # 2.1 Use case
    add_custom_h2("2.1 Chức năng ( usecase )")
    
    add_body_p("Cửa hàng văn phòng phẩm và đồ dùng học tập HieuMini phục vụ nhu cầu mua sắm trực tuyến của học sinh, sinh viên, giáo viên và các bậc phụ huynh. Hệ thống yêu cầu sự tiện lợi, tìm kiếm sản phẩm nhanh chóng, thanh toán trực quan và giao diện thân thiện.")

    add_custom_h3("2.1.1 Xác định tác nhân hệ thống (Actors)")
    add_body_p("Hệ thống website HieuMini phân định 3 tác nhân chính tham gia tương tác với các quyền hạn và vai trò rõ ràng:")
    add_bullet_p("Khách vãng lai (Guest User): Người dùng chưa đăng nhập tài khoản. Có quyền xem danh mục, tìm kiếm đồ dùng học tập, lọc theo mức giá, xem chi tiết sản phẩm, thêm hàng vào giỏ và tiến hành đặt hàng thanh toán nhanh không bắt buộc đăng ký.", "1. ")
    add_bullet_p("Khách hàng thành viên (Registered Customer): Người dùng đã đăng ký và đăng nhập tài khoản. Kế thừa toàn bộ quyền của khách vãng lai, đồng thời được quản lý thông tin cá nhân, xem lại lịch sử các đơn hàng đã đặt, theo dõi tiến độ vận chuyển và gửi đánh giá nhận xét về sản phẩm.", "2. ")
    add_bullet_p("Quản trị viên (Administrator): Người nắm quyền kiểm soát toàn bộ hệ thống. Có quyền truy cập Bảng điều khiển Admin Dashboard, quản lý danh mục đồ dùng học tập, thực hiện thêm/sửa/xóa sản phẩm, theo dõi và cập nhật trạng thái xử lý đơn hàng, quản trị người dùng và giải đáp tin nhắn thắc mắc của khách hàng.", "3. ")

    add_custom_h3("2.1.2 Ma trận phân quyền chức năng Use Case")
    
    # Use case table
    uc_table = doc.add_table(rows=1, cols=4)
    uc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = uc_table.rows[0].cells
    hdr_titles = ["Nhóm Chức Năng", "Chi Tiết Use Case", "Khách Hàng", "Quản Trị (Admin)"]
    for i, title in enumerate(hdr_titles):
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(title)
        r.font.name = 'Times New Roman'
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(11)

    uc_data = [
        ("Tài khoản & Xác thực", "Đăng ký tài khoản mới", "✔ Có", "✔ Quản lý"),
        ("Tài khoản & Xác thực", "Đăng nhập / Đăng xuất", "✔ Có", "✔ Có"),
        ("Tài khoản & Xác thực", "Xem & cập nhật hồ sơ cá nhân", "✔ Có", "✔ Xem tất cả"),
        ("Duyệt & Tìm kiếm", "Xem danh sách 30 sản phẩm", "✔ Có", "✔ Quản lý"),
        ("Duyệt & Tìm kiếm", "Tìm kiếm trực tiếp gợi ý AJAX", "✔ Có", "✔ Có"),
        ("Duyệt & Tìm kiếm", "Lọc theo danh mục / mức giá / hot", "✔ Có", "✔ Lọc quản trị"),
        ("Duyệt & Tìm kiếm", "Xem chi tiết & thông số kỹ thuật", "✔ Có", "✔ Có"),
        ("Giỏ hàng & Đặt hàng", "Thêm sản phẩm vào giỏ (AJAX)", "✔ Có", "―"),
        ("Giỏ hàng & Đặt hàng", "Tăng/giảm số lượng, xóa giỏ hàng", "✔ Có", "―"),
        ("Giỏ hàng & Đặt hàng", "Áp dụng mã giảm giá (Voucher)", "✔ Có", "✔ Quản lý mã"),
        ("Giỏ hàng & Đặt hàng", "Thanh toán (COD / Chuyển khoản)", "✔ Có", "―"),
        ("Giỏ hàng & Đặt hàng", "Theo dõi trạng thái đơn hàng", "✔ Có", "✔ Cập nhật trạng thái"),
        ("Tương tác & Đánh giá", "Gửi nhận xét đánh giá sao", "✔ Có", "✔ Kiểm duyệt"),
        ("Tương tác & Đánh giá", "Gửi liên hệ & góp ý", "✔ Có", "✔ Trả lời"),
        ("Quản trị hệ thống", "Xem Dashboard thống kê doanh thu", "―", "✔ Toàn quyền"),
        ("Quản trị hệ thống", "Thêm, sửa, xóa sản phẩm (CRUD)", "―", "✔ Toàn quyền"),
        ("Quản trị hệ thống", "Thêm, sửa, xóa danh mục", "―", "✔ Toàn quyền")
    ]

    for row_idx, data in enumerate(uc_data):
        row = uc_table.add_row()
        cells = row.cells
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(data):
            set_cell_background(cells[c_idx], bg_col)
            set_cell_margins(cells[c_idx], top=80, bottom=80, left=120, right=120)
            p = cells[c_idx].paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10.5)
            if c_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_custom_h3("2.1.3 Đặc tả chi tiết các Use Case chính")
    add_body_p("1. Use Case 'Thêm vào giỏ hàng và Đặt hàng trực tuyến':", "• ")
    add_body_p("   - Tác nhân: Khách hàng (Vãng lai hoặc Thành viên).")
    add_body_p("   - Điều kiện tiên quyết: Sản phẩm còn tồn kho (> 0).")
    add_body_p("   - Luồng sự kiện chính: Khách hàng click 'Thêm vào giỏ' tại danh mục hoặc trang chi tiết -> Javascript kích hoạt yêu cầu AJAX gửi tới `api/cart.php` -> PHP kiểm tra sản phẩm, cập nhật giỏ hàng trong Session -> Trả về JSON cập nhật số lượng badge và hiển thị Toast thông báo thành công mà không phải tải lại trang -> Khách hàng chuyển sang trang `cart.php`, nhập voucher giảm giá -> Đi tới `checkout.php`, nhập địa chỉ nhận hàng, chọn phương thức COD/Chuyển khoản -> Nhấn 'Đặt hàng ngay' -> Hệ thống khởi tạo Transaction, ghi thông tin vào bảng `orders` và `order_items`, giảm trừ tồn kho sản phẩm, làm trống giỏ hàng và chuyển hướng sang `order-success.php` kèm pháo hoa chúc mừng.")
    
    add_body_p("2. Use Case 'Quản lý đơn hàng và Cập nhật trạng thái (Admin)':", "• ")
    add_body_p("   - Tác nhân: Quản trị viên (Admin).")
    add_body_p("   - Điều kiện tiên quyết: Đã đăng nhập bằng tài khoản có quyền `role = admin`.")
    add_body_p("   - Luồng sự kiện chính: Admin truy cập `admin/orders.php` -> Hệ thống liệt kê toàn bộ danh sách đơn hàng kèm bộ lọc trạng thái -> Admin click xem chi tiết `admin/order-detail.php` -> Kiểm tra thông tin người nhận, danh sách hàng hóa và tổng tiền -> Lựa chọn trạng thái mới (Đang xử lý -> Đang giao -> Hoàn thành) -> Bấm cập nhật -> Hệ thống lưu vào CSDL và đồng bộ thời gian thực.")

    # 2.2 Cơ sở dữ liệu
    add_custom_h2("2.2 Cơ sở dữ liệu")
    
    add_body_p("Cơ sở dữ liệu của website HieuMini mang tên `hieumini_db`, được thiết kế theo chuẩn dạng chuẩn 3 (3NF), đảm bảo tính toàn vẹn dữ liệu, tối ưu hóa kích thước lưu trữ và tốc độ truy vấn cao.")

    add_custom_h3("2.2.1 Sơ đồ thực thể liên kết (ERD)")
    add_body_p("Mối quan hệ giữa các thực thể trong cơ sở dữ liệu `hieumini_db` được thiết kế chặt chẽ:")
    add_bullet_p("Quan hệ Categories - Products (1 - N): Một danh mục có thể chứa nhiều sản phẩm đồ dùng học tập. Khi danh mục bị xóa, các sản phẩm liên quan được bảo vệ hoặc xử lý theo ràng buộc CASCADE/RESTRICT.", "• ")
    add_bullet_p("Quan hệ Users - Orders (1 - N): Một khách hàng đã đăng ký có thể tạo nhiều đơn hàng trong lịch sử. Nếu tài khoản người dùng bị xóa, trường `user_id` trong đơn hàng được SET NULL để bảo toàn chứng từ kế toán doanh thu.", "• ")
    add_bullet_p("Quan hệ Orders - Order_Items (1 - N): Một đơn hàng bao gồm nhiều mặt hàng chi tiết. Quan hệ ràng buộc ON DELETE CASCADE đảm bảo tính đồng bộ khi xóa đơn.", "• ")
    add_bullet_p("Quan hệ Products - Reviews (1 - N): Mỗi sản phẩm có thể nhận được nhiều đánh giá và nhận xét số sao từ khách hàng.", "• ")

    add_custom_h3("2.2.2 Cấu trúc chi tiết 8 bảng dữ liệu")
    
    tables_desc = [
        ("Bảng 1: categories (Danh mục sản phẩm)", [
            ("id", "INT", "Khóa chính (PK), tự động tăng (AUTO_INCREMENT)"),
            ("name", "VARCHAR(100)", "Tên danh mục đồ dùng học tập (Bút viết, Sổ tay, v.v.)"),
            ("slug", "VARCHAR(100)", "Đường dẫn tĩnh thân thiện SEO, chỉ mục UNIQUE"),
            ("description", "TEXT", "Mô tả chi tiết về nhóm đồ dùng học tập"),
            ("icon", "VARCHAR(50)", "Mã biểu tượng Bootstrap Icons đại diện"),
            ("badge", "VARCHAR(50)", "Nhãn hiển thị nổi bật (Bán chạy, Mới về, Hot)"),
            ("created_at", "TIMESTAMP", "Thời điểm tạo danh mục trong hệ thống")
        ]),
        ("Bảng 2: products (30 Sản phẩm đồ dùng học tập)", [
            ("id", "INT", "Khóa chính (PK), tự động tăng"),
            ("category_id", "INT", "Khóa ngoại (FK) liên kết tới categories.id"),
            ("name", "VARCHAR(255)", "Tên đầy đủ của sản phẩm đồ dùng học tập"),
            ("slug", "VARCHAR(255)", "Đường dẫn URL thân thiện, chỉ mục UNIQUE"),
            ("sku", "VARCHAR(50)", "Mã quản lý kho hàng độc nhất (vd: PEN-MOR-01)"),
            ("price", "DECIMAL(12,2)", "Giá bán niêm yết chính thức (VNĐ)"),
            ("sale_price", "DECIMAL(12,2)", "Giá khuyến mãi giờ vàng (VNĐ), có thể NULL"),
            ("image", "VARCHAR(255)", "Tên tệp hình ảnh sản phẩm (vd: p1.png... p30.png)"),
            ("description", "TEXT", "Bài viết mô tả công năng, ưu điểm sản phẩm"),
            ("specification", "TEXT", "Thông số kỹ thuật chi tiết (ngòi bút, kích thước, số trang)"),
            ("stock_quantity", "INT", "Số lượng hàng còn tồn trong kho"),
            ("is_featured", "TINYINT(1)", "Cờ đánh dấu sản phẩm nổi bật trang chủ (0/1)"),
            ("is_hot", "TINYINT(1)", "Cờ đánh dấu sản phẩm bán chạy Hot (0/1)"),
            ("is_new", "TINYINT(1)", "Cờ đánh dấu sản phẩm mới về (0/1)"),
            ("rating", "DECIMAL(3,2)", "Điểm đánh giá trung bình (1.0 đến 5.0 sao)"),
            ("review_count", "INT", "Tổng số lượt đánh giá từ khách hàng")
        ]),
        ("Bảng 3: users (Người dùng và Quản trị viên)", [
            ("id", "INT", "Khóa chính (PK), tự động tăng"),
            ("fullname", "VARCHAR(100)", "Họ và tên đầy đủ của người dùng"),
            ("email", "VARCHAR(150)", "Địa chỉ Email đăng nhập, chỉ mục UNIQUE"),
            ("password", "VARCHAR(255)", "Mật khẩu đã băm bảo mật bằng thuật toán bcrypt"),
            ("phone", "VARCHAR(20)", "Số điện thoại liên hệ nhận hàng"),
            ("address", "TEXT", "Địa chỉ giao hàng mặc định của khách"),
            ("role", "ENUM", "Phân quyền: 'admin' (Quản trị) hoặc 'customer' (Khách)"),
            ("avatar", "VARCHAR(255)", "Ảnh đại diện người dùng")
        ]),
        ("Bảng 4: orders (Đơn đặt hàng)", [
            ("id", "INT", "Khóa chính (PK), tự động tăng"),
            ("order_code", "VARCHAR(30)", "Mã đơn hàng độc nhất (vd: HM-20260801-ABCD)"),
            ("user_id", "INT", "Khóa ngoại (FK) tới users.id (NULL nếu khách vãng lai)"),
            ("customer_name", "VARCHAR(100)", "Họ tên người nhận hàng"),
            ("customer_email", "VARCHAR(150)", "Email nhận thông báo trạng thái đơn"),
            ("customer_phone", "VARCHAR(20)", "Số điện thoại người nhận hàng"),
            ("shipping_address", "TEXT", "Địa chỉ nhận hàng chi tiết"),
            ("subtotal", "DECIMAL(12,2)", "Tổng tiền hàng trước giảm giá"),
            ("discount_amount", "DECIMAL(12,2)", "Số tiền được giảm từ Voucher khuyến mãi"),
            ("shipping_fee", "DECIMAL(12,2)", "Phí giao hàng (0đ nếu >= 250k)"),
            ("total_amount", "DECIMAL(12,2)", "Tổng số tiền thực tế khách phải thanh toán"),
            ("payment_method", "ENUM", "Phương thức: 'cod', 'bank_transfer', 'momo'"),
            ("payment_status", "ENUM", "Trạng thái thanh toán: 'unpaid', 'paid', 'refunded'"),
            ("status", "ENUM", "Tiến độ đơn: 'pending', 'processing', 'shipping', 'completed', 'cancelled'")
        ]),
        ("Bảng 5: order_items (Chi tiết mặt hàng trong đơn)", [
            ("id", "INT", "Khóa chính (PK), tự động tăng"),
            ("order_id", "INT", "Khóa ngoại (FK) tới orders.id (ON DELETE CASCADE)"),
            ("product_id", "INT", "Mã sản phẩm tương ứng"),
            ("product_name", "VARCHAR(255)", "Tên sản phẩm tại thời điểm đặt mua"),
            ("product_image", "VARCHAR(255)", "Hình ảnh sản phẩm tại thời điểm mua"),
            ("price", "DECIMAL(12,2)", "Đơn giá mua tại thời điểm đặt hàng"),
            ("quantity", "INT", "Số lượng đặt mua"),
            ("total_price", "DECIMAL(12,2)", "Thành tiền (price * quantity)")
        ]),
        ("Bảng 6: reviews (Đánh giá sản phẩm)", [
            ("id", "INT", "Khóa chính (PK), tự động tăng"),
            ("product_id", "INT", "Khóa ngoại (FK) tới products.id"),
            ("user_name", "VARCHAR(100)", "Tên người đánh giá"),
            ("rating", "INT", "Số sao đánh giá (từ 1 đến 5 sao)"),
            ("comment", "TEXT", "Nội dung nhận xét chi tiết")
        ]),
        ("Bảng 7: coupons (Mã giảm giá Voucher)", [
            ("id", "INT", "Khóa chính (PK), tự động tăng"),
            ("code", "VARCHAR(50)", "Mã voucher (vd: HIEUMINI10, FREESHIP, BACK2SCHOOL)"),
            ("discount_type", "ENUM", "Loại giảm: 'percentage' (%) hoặc 'fixed' (tiền mặt)"),
            ("discount_value", "DECIMAL(10,2)", "Giá trị chiết khấu tương ứng"),
            ("min_order_value", "DECIMAL(12,2)", "Giá trị đơn hàng tối thiểu để được áp dụng"),
            ("is_active", "TINYINT(1)", "Trạng thái kích hoạt của mã (0/1)")
        ]),
        ("Bảng 8: contacts (Hòm thư liên hệ)", [
            ("id", "INT", "Khóa chính (PK), tự động tăng"),
            ("fullname", "VARCHAR(100)", "Họ tên người gửi thư liên hệ"),
            ("email", "VARCHAR(150)", "Email phản hồi"),
            ("phone", "VARCHAR(20)", "Số điện thoại liên hệ"),
            ("subject", "VARCHAR(255)", "Tiêu đề yêu cầu tư vấn"),
            ("message", "TEXT", "Nội dung tin nhắn khách hàng gửi"),
            ("status", "ENUM", "Trạng thái xử lý: 'new' (Mới) hoặc 'replied' (Đã trả lời)")
        ])
    ]

    for tbl_name, tbl_cols in tables_desc:
        add_custom_h3(tbl_name)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        h_cells = tbl.rows[0].cells
        for idx, col_h in enumerate(["Tên Trường", "Kiểu Dữ Liệu", "Mô Tả & Ràng Buộc"]):
            set_cell_background(h_cells[idx], "334155")
            set_cell_margins(h_cells[idx], top=80, bottom=80, left=100, right=100)
            p = h_cells[idx].paragraphs[0]
            r = p.add_run(col_h)
            r.font.name = 'Times New Roman'
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)
        
        for r_idx, (cname, ctype, cdesc) in enumerate(tbl_cols):
            row = tbl.add_row()
            cells = row.cells
            bg = "F1F5F9" if r_idx % 2 == 0 else "FFFFFF"
            for c_i, c_val in enumerate([cname, ctype, cdesc]):
                set_cell_background(cells[c_i], bg)
                set_cell_margins(cells[c_i], top=60, bottom=60, left=100, right=100)
                p = cells[c_i].paragraphs[0]
                r = p.add_run(c_val)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                if c_i == 0:
                    r.font.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # =========================================================================
    # CHƯƠNG 3. CHƯƠNG TRÌNH THỬ NGHIỆM
    # =========================================================================
    add_custom_h1("Chương 3. Chương trình thử nghiệm")
    
    add_body_p("Chương 3 trình bày toàn diện về kết quả thử nghiệm thực tế của hệ thống website HieuMini, bao gồm kiến trúc giao diện người dùng, hệ thống hiệu ứng chuyển động Transitions & Animations, phân tích chi tiết từng màn hình chức năng và đưa ra các kết luận, đánh giá chuyên sâu.")

    # 3.1 Giao diện
    add_custom_h2("3.1 Giao diện")
    
    add_body_p("Giao diện người dùng (UI) và trải nghiệm người dùng (UX) của website HieuMini được thiết kế dựa trên triết lý Hiện đại, Tinh tế, Trẻ trung và Đậm chất nghệ thuật học đường, tuân thủ nghiêm ngặt các nguyên tắc thiết kế giao diện cao cấp.")

    add_custom_h3("3.1.1 Ngôn ngữ thiết kế UI/UX và hệ thống Design Tokens")
    add_bullet_p("Bảng màu chủ đạo (Color Palette): Tone màu chính là sắc Indigo (#4F46E5) tượng trưng cho tri thức và công nghệ hiện đại, kết hợp hài hòa với màu hồng Magenta pastel (#EC4899), xanh ngọc Cyan (#06B6D4) và vàng hổ phách Amber (#F59E0B) tạo cảm giác ngọt ngào, tươi vui và đầy cảm hứng cho lứa tuổi học sinh, sinh viên.", "• ")
    add_bullet_p("Hệ thống Typography chuẩn quốc tế: Sử dụng bộ phông chữ Google Fonts 'Plus Jakarta Sans' với khả năng hiển thị tiếng Việt sắc nét, bo tròn tinh tế và độ tương phản quang học xuất sắc trên mọi loại màn hình từ Retina đến màn hình điện thoại.", "• ")
    add_bullet_p("Hiệu ứng kính mờ thời thượng (Glassmorphism): Thanh điều hướng Navbar dính (Sticky Navbar) ứng dụng kỹ thuật backdrop-filter: blur(16px) với nền mờ trong suốt 92%, tạo cảm giác nổi bật và sang trọng khi cuộn trang.", "• ")
    add_bullet_p("30 Hình ảnh sản phẩm trực quan chất lượng cao: Toàn bộ 30 sản phẩm từ p1.png đến p30.png đều được thiết kế và render đồ họa ở độ phân giải cao 800x800, thể hiện chính xác từng chi tiết ngòi bút, thỏi màu, sổ còng và ba lô chống gù.", "• ")

    add_custom_h3("3.1.2 Hệ thống Transitions và Micro-Animations tương tác sống động")
    add_body_p("Website HieuMini được trang bị hệ sinh thái hiệu ứng chuyển động phong phú (Animations & Transitions) nhằm nâng tầm trải nghiệm người dùng:")
    add_bullet_p("@keyframes floatHero & floatUp: Tạo hiệu ứng trôi nổi bồng bềnh nhẹ nhàng cho các huy hiệu ưu đãi và banner quảng cáo trên Hero Section, thu hút sự chú ý của khách hàng ngay từ cái nhìn đầu tiên.", "1. ")
    add_bullet_p("@keyframes pulseBadge: Hiệu ứng đập nhịp tim nhẹ nhàng cho vòng tròn số lượng giỏ hàng trên thanh Menu, tạo điểm nhấn hành động.", "2. ")
    add_bullet_p("Card 3D Hover & Elevation: Thẻ sản phẩm khi di chuột vào (hover) sẽ nổi lên mượt mà theo đường cong cubic-bezier(0.16, 1, 0.3, 1), phóng to nhẹ hình ảnh (+8%) và trượt thanh thao tác nhanh 'Xem chi tiết / Mua ngay' từ dưới lên.", "3. ")
    add_bullet_p("Hiệu ứng rung giỏ hàng (.shake) & Toast trượt ngang: Khi khách hàng thêm sản phẩm vào giỏ, biểu tượng giỏ hàng sẽ lắc nhẹ 3 nhịp và một hộp thông báo Toast bóng bẩy sẽ trượt từ góc phải màn hình vào với thanh đếm thời gian tự tắt.", "4. ")
    add_bullet_p("Pháo hoa chúc mừng đặt hàng thành công (Confetti Animation): Tại trang `order-success.php`, hệ thống tự động kích hoạt 60 hạt pháo hoa nhiều màu sắc rơi tự do xoay chuyển khắp màn hình, mang lại cảm xúc phấn khởi và hài lòng tuyệt đối cho khách hàng.", "5. ")

    add_custom_h3("3.1.3 Giới thiệu chi tiết các phân hệ giao diện website")
    add_body_p("1. Phân hệ Trang chủ (index.php):", "• ")
    add_body_p("   - Banner Hero hoành tráng với thông điệp chào mừng tựu trường 2026, nút kêu gọi hành động CTA nổi bật.")
    add_body_p("   - Thanh cam kết 4 giá trị vàng: Freeship đơn từ 250k, 100% Chính hãng, Đổi trả 7 ngày, Hỗ trợ 24/7.")
    add_body_p("   - Lưới 6 danh mục đồ dùng học tập với icon biểu tượng sinh động và hiệu ứng đổi màu khi hover.")
    add_body_p("   - Khung Flash Sale giờ vàng với đồng hồ đếm ngược thời gian thực (Giờ : Phút : Giây).")
    add_body_p("   - Khối sản phẩm 'Bán chạy nhất' và 'Hàng mới về' tích hợp nút thêm nhanh vào giỏ hàng.")
    add_body_p("   - 3 Banner quảng cáo bộ sưu tập bút pastel, sổ bullet journal và dụng cụ mỹ thuật.")

    add_body_p("2. Phân hệ Danh mục sản phẩm & Bộ lọc thông minh (products.php):", "• ")
    add_body_p("   - Thanh tìm kiếm trực tiếp với gợi ý tức thì (Live Search Suggestions) thông qua AJAX.")
    add_body_p("   - Bộ lọc đa tiêu chí phía bên trái: Lọc theo 6 danh mục, lọc theo 4 khoảng giá tiền (<50k, 50k-100k, 100k-200k, >200k), lọc theo sản phẩm Hot / Giảm giá / Mới.")
    add_body_p("   - Menu sắp xếp linh hoạt: Mới nhất, Giá tăng dần, Giá giảm dần, Đánh giá sao cao nhất, Tên A-Z.")

    add_body_p("3. Phân hệ Chi tiết sản phẩm & Đánh giá (product-detail.php):", "• ")
    add_body_p("   - Bộ sưu tập ảnh sản phẩm với chức năng chuyển đổi nhanh thumbnail và xem ảnh phóng to.")
    add_body_p("   - Thông tin giá bán, tỷ lệ giảm giá, mã SKU, trạng thái còn hàng.")
    add_body_p("   - Bộ điều chỉnh tăng/giảm số lượng với nút bấm +/- có giới hạn tồn kho.")
    add_body_p("   - Bảng thông số kỹ thuật chi tiết trình bày khoa học.")
    add_body_p("   - Khu vực Đánh giá nhận xét của khách hàng kèm biểu mẫu gửi đánh giá sao trực tiếp lưu vào MySQL.")
    add_body_p("   - Khối sản phẩm gợi ý cùng danh mục.")

    add_body_p("4. Phân hệ Giỏ hàng & Thanh toán (cart.php, checkout.php, order-success.php):", "• ")
    add_body_p("   - Bảng giỏ hàng hiển thị trực quan sản phẩm, đơn giá, số lượng thay đổi thời gian thực.")
    add_body_p("   - Thanh tiến trình thông minh: 'Mua thêm... để được FREESHIP toàn quốc'.")
    add_body_p("   - Ô áp dụng mã giảm giá voucher (HIEUMINI10, FREESHIP, BACK2SCHOOL) tự động tính toán số tiền giảm.")
    add_body_p("   - Biểu mẫu thanh toán bảo mật, hỗ trợ các hình thức: COD, Quét mã VietQR và Ví MoMo.")
    add_body_p("   - Hóa đơn điện tử chi tiết và hiệu ứng pháo hoa ăn mừng tại trang xác nhận đơn hàng.")

    add_body_p("5. Phân hệ Quản trị Admin Dashboard (admin/):", "• ")
    add_body_p("   - Bảng điều khiển tổng quan với 4 thẻ chỉ số KPI: Tổng doanh thu, Tổng đơn hàng, Tổng sản phẩm, Khách hàng.")
    add_body_p("   - Bảng quản lý đơn hàng mới với bộ chọn đổi trạng thái nhanh (Chờ xác nhận, Đang xử lý, Đang giao, Hoàn thành, Hủy).")
    add_body_p("   - Quản lý danh sách sản phẩm với chức năng Thêm mới, Sửa thông tin, Xóa và cập nhật số lượng kho.")
    add_body_p("   - Quản lý danh mục đồ dùng học tập, quản lý tài khoản người dùng và hòm thư liên hệ.")

    # 3.2 Kết luận
    add_custom_h2("3.2 Kết luận")
    
    add_custom_h3("3.2.1 Đánh giá kết quả đạt được đối chiếu mục tiêu")
    add_body_p("Sau quá trình nghiên cứu, phân tích và triển khai thực tế, dự án xây dựng website bán đồ dùng học tập HieuMini đã hoàn thành xuất sắc và vượt mức 100% các mục tiêu đề ra ban đầu:")
    add_bullet_p("Hoàn thiện hệ thống Backend mạnh mẽ với PHP 8.x và CSDL MySQL, tuân thủ chuẩn PDO Prepared Statements bảo mật tuyệt đối.", "✔ ")
    add_bullet_p("Tích hợp đầy đủ 30 sản phẩm đồ dùng học tập thực tế sắc nét, chia đều trên 6 danh mục phong phú.", "✔ ")
    add_bullet_p("Xây dựng giao diện UI/UX đỉnh cao với hệ thống Design Tokens, Glassmorphism và hơn 10 hiệu ứng Transitions & Animations sống động.", "✔ ")
    add_bullet_p("Hoàn thiện toàn diện 2 phân hệ: Khách hàng mua sắm và Quản trị viên Admin Dashboard điều hành.", "✔ ")
    add_bullet_p("Tốc độ tải trang siêu nhanh, cấu trúc mã nguồn module hóa sạch sẽ, dễ bảo trì và mở rộng.", "✔ ")

    add_custom_h3("3.2.2 Ưu điểm nổi bật của hệ thống HieuMini")
    add_bullet_p("Tính thực tiễn cao: Đáp ứng trọn vẹn nghiệp vụ bán lẻ văn phòng phẩm với giỏ hàng AJAX, voucher khuyến mãi, thanh toán COD/VietQR và quản lý kho hàng.", "• ")
    add_bullet_p("Trải nghiệm mượt mà: Nhờ ứng dụng kỹ thuật AJAX trong tìm kiếm và giỏ hàng, người dùng không bị gián đoạn hay phải tải lại trang nhiều lần.", "• ")
    add_bullet_p("An toàn bảo mật: Ngăn chặn triệt để SQL Injection qua PDO, mã hóa mật khẩu bcrypt và lọc dữ liệu đầu vào chống XSS.", "• ")
    add_bullet_p("Khả năng tương thích thiết bị (Responsive): Giao diện tự động co giãn và hiển thị tối ưu trên máy tính để bàn, laptop, máy tính bảng và điện thoại thông minh.", "• ")

    add_custom_h3("3.2.3 Hạn chế còn tồn đọng và hướng phát triển trong tương lai")
    add_body_p("Bên cạnh những thành tựu đạt được, hệ thống vẫn còn một số điểm có thể tiếp tục nâng cấp trong các phiên bản tiếp theo:")
    add_bullet_p("Tích hợp cổng thanh toán trực tuyến tự động qua API chính thức của VNPay, MoMo và ZaloPay có IPN Callback xác thực tức thì.", "1. Hướng phát triển 1: ")
    add_bullet_p("Ứng dụng Trí tuệ Nhân tạo (AI Recommendation Engine) để phân tích hành vi người dùng và tự động gợi ý các combo dụng cụ học tập phù hợp theo từng khối lớp và môn học.", "2. Hướng phát triển 2: ")
    add_bullet_p("Phát triển ứng dụng di động (Mobile App) trên nền tảng Flutter / React Native đồng bộ dữ liệu qua RESTful API với máy chủ PHP hiện tại.", "3. Hướng phát triển 3: ")

    doc.add_page_break()

    # =========================================================================
    # TÀI LIỆU THAM KHẢO
    # =========================================================================
    add_custom_h1("TÀI LIỆU THAM KHẢO")
    
    references = [
        ("1. ", "The PHP Group. (2024). PHP Manual - Documentation and Language Reference. Truy cập tại: https://www.php.net/docs.php"),
        ("2. ", "Oracle Corporation. (2024). MySQL 8.0 Reference Manual. Truy cập tại: https://dev.mysql.com/doc/"),
        ("3. ", "W3Schools. (2024). PHP & MySQL Web Development Tutorials and Best Practices."),
        ("4. ", "Mozilla Developer Network (MDN). (2024). CSS Transitions and Keyframe Animations Guide."),
        ("5. ", "Robin Nixon. (2021). Learning PHP, MySQL & JavaScript: With jQuery, CSS & HTML5 (6th Edition). O'Reilly Media."),
        ("6. ", "Matt Stauffer. (2019). Laravel: Up & Running - A Framework for Building Modern PHP Apps. O'Reilly Media.")
    ]

    for num, ref_text in references:
        add_body_p(ref_text, bold_prefix=num)

    # Save document
    out_docx_path = r"c:\Users\tranv\Desktop\HieuWeb03\BaoCao.docx"
    doc.save(out_docx_path)
    print(f"Report generated successfully at: {out_docx_path}")

create_report()
